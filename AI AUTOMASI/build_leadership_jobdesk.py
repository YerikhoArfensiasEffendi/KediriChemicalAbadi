import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

BASE_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/KCA DOKUMEN"
MGT_DIR = os.path.join(BASE_DIR, "09_DOKUMEN_DIREKSI_DAN_MANAJEMEN_PUNCAK")
os.makedirs(MGT_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PT KCA CHEMICAL\nSTRUKTUR MANAJEMEN\nSISTEM ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI - DIREKSI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block_3(doc, t1="Komisaris Utama / Owner", n1="( Ayah / Owner )", t2="Manager Utama / Direktur", n2="( Anda / Manager Utama )", t3="Wakil Manager Operasional", n3="( Teman Anda / Wakil )"):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    col_widths = [Cm(5.0), Cm(5.0), Cm(5.5)]
    
    titles = [(t1, n1), (t2, n2), (t3, n3)]
    for idx, (title, name) in enumerate(titles):
        cell = table.cell(0, idx)
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(title + "\n\n\n\n\n" + name + "\nTanggal: .....................")

def create_leadership_doc():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Struktur Organisasi Kepemimpinan & Uraian Tugas (Owner, Manager Utama & Wakil Manager)", "MGT-DIR-007", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. STRUKTUR KEPEMIMPINAN INTI PERUSAHAAN")
    add_p(doc, "Untuk menjamin tata kelola perusahaan yang profesional, transparan, dan terarah (Good Corporate Governance & ISO 9001:2015 Klausul 5), struktur kepemimpinan PT KCA Chemical dibagi ke dalam 3 Tingkatan Utama:")
    
    headers_org = ["Jabatan", "Peran Strategis", "Kedudukan / Lini", "Penanggung Jawab"]
    data_org = [
        ["Komisaris Utama / Owner", "Pemilik Perusahaan, Investor & Pengawas Strategis", "Tingkat Puncak (Top Governance)", "Ayah"],
        ["Manager Utama / Direktur", "Pengelola Utama Bisnis, Finansial & Operasional", "Pimpinan Eksekutif Tertinggi", "Anda"],
        ["Wakil Manager Operasional", "Tangan Kanan Eksekutif & Pengawas Lapangan Pabrik", "Wakil Pimpinan Operasional", "Teman Anda"]
    ]
    add_styled_table(doc, headers_org, data_org, [Cm(3.8), Cm(5.2), Cm(3.8), Cm(2.7)])

    add_heading_1(doc, "2. RUANG LINGKUP & JOB DESCRIPTION DETAIL")
    
    add_heading_2(doc, "A. AYAH (KOMISARIS UTAMA / OWNER / PEMILIK PERUSAHAAN)")
    add_p(doc, "Fokus Utama: Pengawasan investasi modal, perlindungan aset jangka panjang, dan persetujuan kebijakan strategis.")
    add_bullet(doc, "Mengawasi arah kebijakan umum perusahaan dan kesinambungan bisnis jangka panjang.", "1. Pengawasan Strategis: ");
    add_bullet(doc, "Memberikan persetujuan atas rencana investasi besar (pembelian tanah, ekspansi gedung pabrik, atau pembelian mesin skala besar).", "2. Otoritas Finansial Tertinggi: ");
    add_bullet(doc, "Menerima dan mengevaluasi Laporan Kinerja Bisnis & Laporan Keuangan Bulanan dari Manager Utama.", "3. Evaluasi Kinerja: ");
    add_bullet(doc, "Membantu membina hubungan tingkat tinggi (relasi perbankan, instansi, dan pemangku kepentingan senior).", "4. Hubungan Eksternal: ");
    add_bullet(doc, "Bebas dari beban operasional harian teknis pabrik (tidak perlu mengurusi mixing, logistik, atau penanganan komplain harian).", "5. Batas Wewenang: ");

    add_heading_2(doc, "B. ANDA (MANAGER UTAMA / DIREKTUR PENGELOLA UTAMA)")
    add_p(doc, "Fokus Utama: Mengendalikan roda bisnis, strategi komersial makloon, arus kas (cash flow), kepatuhan regulasi Kemenkes, dan kepemimpinan seluruh tim.")
    add_bullet(doc, "Memimpin negosiasi kerja sama makloon dengan PT Mitra/Distributor, menentukan harga jual jasa (Quotation), dan menandatangani kontrak hukum bermaterai.", "1. Hubungan Klien & Komersial: ");
    add_bullet(doc, "Mengatur arus kas operasional (Cash Flow), penerimaan pembayaran DP/Pelunasan makloon, pembayaran supplier bahan kimia, dan penetapan gaji karyawan/PJT.", "2. Otoritas Keuangan Operasional: ");
    add_bullet(doc, "Mengarahkan dan mengawasi kinerja PJT dalam hal pendaftaran Izin Edar Kemenkes RI PKD, pengujian laboratorium KAN, dan pelaporan e-Report.", "3. Pengawasan Regulasi & Mutu: ");
    add_bullet(doc, "Memberikan arahan kerja, target jadwal penyelesaian pesanan, dan pendelegasian tugas lapangan kepada Wakil Manager.", "4. Manajemen Tim Pabrik: ");
    add_bullet(doc, "Menyusun dan menyampaikan Laporan Eksekutif Bulanan kepada Ayah (Owner).", "5. Pelaporan: ");

    add_heading_2(doc, "C. TEMAN ANDA (WAKIL MANAGER OPERASIONAL / DEPUTY GENERAL MANAGER)")
    add_p(doc, "Fokus Utama: Memastikan eksekusi teknis di lantai pabrik berjalan 100% tepat waktu, disiplin, efisien, dan mengawal alur dari bahan masuk hingga pengiriman.")
    add_bullet(doc, "Mengontrol kedisiplinan jam kerja operator, kepatuhan pemakaian APD, dan standar kebersihan ruangan pabrik sesuai CPKRTB.", "1. Pengawasan Lantai Produksi: ");
    add_bullet(doc, "Memastikan proses penimbangan, mixing, filling, dan pelabelan botol berjalan lancar sesuai Surat Perintah Kerja (SPK) dari Manager Utama.", "2. Eksekusi Jadwal Produksi: ");
    add_bullet(doc, "Memantau ketersediaan stok bahan kimia & kemasan di gudang (mengingatkan Manager Utama saat bahan mulai menipis untuk re-order).", "3. Manajemen Inventaris Gudang: ");
    add_bullet(doc, "Mengawal proses pengepakan karton, uji kebocoran tutup botol, hingga proses pemuatan barang ke truk ekspedisi bersama Surat Jalan & CoA.", "4. Kontrol Kualitas & Ekspedisi: ");
    add_bullet(doc, "Mengatasi kendala teknis harian di lapangan dan melaporkan status harian kepada Manager Utama.", "5. Pelaporan & Koordinasi: ");
    add_bullet(doc, "Mewakili Manager Utama di pabrik saat Manager Utama sedang bertugas di luar (meeting dengan klien makloon / instansi).", "6. Representasi Internal: ");

    add_heading_1(doc, "3. MATRIKS OTORITAS & BATAS PENGAMBILAN KEPUTUSAN")
    headers_mat = ["Jenis Keputusan / Tindakan", "Wakil Manager", "Manager Utama (Anda)", "Owner (Ayah)"]
    data_mat = [
        ["Pengaturan Jadwal Kerja & Shift Operator", "Mengeksekusi", "Menyetujui", "Mengetahui"],
        ["Pembelian Bahan Baku Rutin (Operasional)", "Mengajukan", "Menyetujui & Bayar", "Menerima Laporan"],
        ["Penetapan Harga Makloon & Teken Kontrak", "Memberi Masukan", "Memutuskan & Teken", "Mengetahui"],
        ["Penerimaan / Pemberhentian Karyawan Pabrik", "Rekomendasi", "Memutuskan", "Mengetahui"],
        ["Investasi Mesin Baru / Renovasi Besar (> Rp 15 Jt)", "Mengidentifikasi", "Menyusun Anggaran", "Persetujuan Akhir"],
        ["Rilis Produk & Penandatanganan CoA", "Koordinasi QC/PJT", "Pengawasan", "Menerima Laporan"]
    ]
    add_styled_table(doc, headers_mat, data_mat, [Cm(5.0), Cm(3.2), Cm(3.8), Cm(3.5)])

    add_signature_block_3(doc)
    doc.save(os.path.join(MGT_DIR, "MGT-07_Struktur_Jobdesk_Owner_Manager_Utama_dan_Wakil.docx"))
    print("Created MGT-07")

if __name__ == "__main__":
    create_leadership_doc()

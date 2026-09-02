import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

OUTPUT_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/dokumen_iso9001_pkrt"
os.makedirs(OUTPUT_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, dept_name, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run(f"PT [NAMA PERUSAHAAN]\n{dept_name.upper()}\nSISTEM ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block(doc, left_title="Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", right_title="Disetujui Oleh,\nDirektur Perusahaan", left_name="( ............................................ )", right_name="( ............................................ )"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    col_widths = [Cm(7.5), Cm(8.0)]
    
    cell_l = table.cell(0, 0)
    cell_l.width = col_widths[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l1 = p_l.add_run(left_title + "\n\n\n\n\n" + left_name + "\nTanggal: ...................................")
    r_l1.font.name = 'Times New Roman'
    r_l1.font.size = Pt(9.5)
    r_l1.font.color.rgb = COLOR_BLACK
    
    cell_r = table.cell(0, 1)
    cell_r.width = col_widths[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r.add_run(right_title + "\n\n\n\n\n" + right_name + "\nTanggal: ...................................")
    r_r1.font.name = 'Times New Roman'
    r_r1.font.size = Pt(9.5)
    r_r1.font.color.rgb = COLOR_BLACK

# ==============================================================================
# DOC 06: DIVISI PENANGGUNG JAWAB TEKNIS (PJT)
# ==============================================================================
def create_doc_pjt():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Divisi Teknis & Regulasi (PJT)", "Pedoman Kerja, SOP Teknis Regulasi & Formulir Penanggung Jawab Teknis (PJT) Standar CPKRTB", "SOP-PJT-006", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. RUANG LINGKUP & TUGAS POKOK PJT")
    add_p(doc, "Penanggung Jawab Teknis (PJT) merupakan penanggung jawab utama aspek teknis mutu, formulasi, keselamatan kimia, kepatuhan CPKRTB, serta perizinan regulasi ke Kementerian Kesehatan RI.")
    
    add_heading_1(doc, "2. STANDAR OPERASIONAL PROSEDUR (SOP) DIVISI PJT")
    
    add_heading_2(doc, "SOP-PJT-01: Prosedur Pengajuan & Pengelolaan Izin Edar (e-Reg PKRT Kemenkes)")
    add_bullet(doc, "Memverifikasi kelengkapan berkas: Master Formula, CoA Bahan Baku, MSDS, Hasil Uji Lab Terakreditasi KAN, Surat Perjanjian Makloon, dan Sertifikat Merek DJKI.", "1. Telaah Dokumen: ");
    add_bullet(doc, "Mengisi formulir registrasi online pada aplikasi e-reg PKRT Kemenkes RI (ereg.farmalkes.kemkes.go.id).", "2. Input Sistem: ");
    add_bullet(doc, "Memastikan pembayaran kode billing PNBP Kemenkes RI dilakukan sebelum batas kedaluwarsa billing (7 hari).", "3. Pembayaran PNBP: ");
    add_bullet(doc, "Memantau status verifikasi dan merespon catatan perbaikan (jika ada) maksimal 5 hari kerja.", "4. Respon Tambahan Data: ");
    add_bullet(doc, "Mengunduh dan mengarsipkan Sertifikat Izin Edar KEMENKES RI PKD yang telah terbit ke dalam Binder 1.", "5. Pengarsipan NIE: ");

    add_heading_2(doc, "SOP-PJT-02: Prosedur Pelulusan Produk Akhir (Batch Release)")
    add_bullet(doc, "Menerima Catatan Pengolahan Bets (Batch Record) yang telah diisi lengkap oleh Kepala Produksi dan Petugas QC.", "1. Penerimaan Rekaman: ");
    add_bullet(doc, "Memeriksa kesesuaian parameter uji QC (Bentuk, Warna, Bau, pH, Bobot Jenis, Viskositas, dan Uji Kebocoran Kemasan).", "2. Evaluasi Parameter: ");
    add_bullet(doc, "Menandatangani status 'DILULUSKAN' pada Batch Record dan menerbitkan Certificate of Analysis (CoA) resmi.", "3. Pengesahan CoA: ");
    add_bullet(doc, "Menyerahkan CoA ke bagian Gudang/Ekspedisi sebagai syarat sah pengiriman barang ke distributor.", "4. Penyerahan Dokumen: ");

    add_heading_2(doc, "SOP-PJT-03: Prosedur Pelaporan Produksi Rutin (e-Report Farmalkes)")
    add_bullet(doc, "Mengumpulkan data realisasi volume produksi per nomor batch dan data distribusi setiap semester (6 bulan sekali).", "1. Rekapitulasi Data: ");
    add_bullet(doc, "Menginput laporan jumlah produksi ke portal e-Report Farmalkes Kemenkes RI paling lambat tanggal 15 bulan berikutnya.", "2. Pengisian e-Report: ");

    add_heading_1(doc, "3. FORMULIR & TEMPLATE REKAMAN MUTU PJT")
    
    add_heading_2(doc, "Formulir PJT-01: Lembar Evaluasi Pelulusan Bets (Batch Release Checklist)")
    headers_rel = ["No", "Parameter Evaluasi Kepatuhan", "Standar Acuan", "Hasil Evaluasi PJT", "Status"]
    data_rel = [
        ["1", "Kesesuaian Penimbangan Bahan Baku", "Master Formula Kemenkes", "Sesuai 100%", "OK"],
        ["2", "Catatan Pengadukan & Waktu Mixing", "SOP Produksi", "Lengkap & Homogen", "OK"],
        ["3", "Hasil Uji QC Produk Jadi (pH, BJ, Fisik)", "Spesifikasi Produk", "pH 7.15 | BJ 1.025", "OK"],
        ["4", "Uji Kebocoran Tutup Kemasan (Sampling 5 Btl)", "Kedap / Tidak Bocor", "100% Rapat", "OK"],
        ["5", "Kesesuaian Labeling & No. PKD pada Etiket", "Desain Kemenkes", "No. PKD & Batch Valid", "OK"],
        ["6", "Penyimpanan 2 Botol Retained Sample", "SOP QC Sampling", "Tersimpan di Ruang Arsip", "OK"]
    ]
    add_styled_table(doc, headers_rel, data_rel, [Cm(0.8), Cm(5.2), Cm(3.8), Cm(3.5), Cm(1.7)])

    add_signature_block(doc, "Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", "Disetujui Oleh,\nDirektur Utama")
    
    file_path = os.path.join(OUTPUT_DIR, "06_DEPT_PJT_PENANGGUNG_JAWAB_TEKNIS.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# DOC 07: DIVISI QUALITY CONTROL & LABORATORIUM (QC)
# ==============================================================================
def create_doc_qc():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Divisi Quality Control & Lab", "Pedoman Kerja, SOP Pengujian Laboratorium, Sampling & Kalibrasi Alat Standar ISO 9001", "SOP-QC-007", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. PROSEDUR PENGUJIAN MUTU QC")
    
    add_heading_2(doc, "SOP-QC-01: Prosedur Pengambilan Sampel Uji (Sampling Protocol)")
    add_bullet(doc, "Sampling Bahan Baku: Ambil sampel dari drum bahan menggunakan pipet/thief sampler bersih sebanyak 100 mL per drum secara representatif.", "1. Bahan Masuk: ");
    add_bullet(doc, "Sampling In-Process Control (IPC): Ambil 250 mL dari tangki mixer pada 3 titik (atas, tengah, bawah) setelah proses pengadukan selesai.", "2. Proses Mixing: ");
    add_bullet(doc, "Sampling Produk Jadi: Ambil masing-masing 1 botol pada awal proses filling (10%), pertengahan (50%), dan akhir (90%).", "3. Produk Jadi: ");

    add_heading_2(doc, "SOP-QC-02: Prosedur Pengujian Fisik dan Kimia di Laboratorium Pabrik")
    add_bullet(doc, "Uji pH: Kalibrasi pH meter digital dengan buffer pH 4.01, 7.00, dan 10.01. Celupkan elektroda ke dalam sampel suhu 25°C. Catat nilai pH stabil.", "1. Uji Derajat Keasaman (pH): ");
    add_bullet(doc, "Uji Bobot Jenis: Timbang piknometer kosong kering (W0), isi dengan aquadest pada 25°C dan timbang (W1), isi dengan sampel dan timbang (W2). Hitung: BJ = (W2 - W0) / (W1 - W0).", "2. Uji Bobot Jenis: ");
    add_bullet(doc, "Uji Viskositas & Daya Bersih: Pengujian kekentalan secara komparatif dan uji daya busa menggunakan silinder ukur kocok.", "3. Uji Sifat Fisik: ");

    add_heading_2(doc, "SOP-QC-03: Prosedur Penanganan Hasil Uji Menyimpang (Out of Specification / OOS)")
    add_bullet(doc, "Jika hasil uji berada di luar batas spesifikasi, beri label 'KARANTINA - DITAHAN' pada tangki/palet.", "1. Karantina: ");
    add_bullet(doc, "Lakukan investigasi: cek kalibrasi alat uji, cek ulang formula timbangan, dan uji ulang dengan sampel baru (re-testing).", "2. Investigasi OOS: ");
    add_bullet(doc, "Jika terbukti salah formula/konsentrasi, lakukan tindakan perbaikan (re-work) dengan instruksi tertulis dari PJT atau dilakukan pemusnahan (reject).", "3. Disposisi: ");

    add_heading_1(doc, "2. LOGBOOK & FORMULIR QC HARIAN")
    
    add_heading_2(doc, "Formulir QC-01: Logbook Kalibrasi Harian Timbangan & pH Meter")
    headers_cal = ["Tanggal", "Alat Uji", "Standar Uji / Buffer", "Hasil Pembacaan", "Deviasi", "Petugas", "Status"]
    data_cal = [
        ["19/08/2026", "pH Meter Digital A", "Buffer pH 7.00\nBuffer pH 4.01", "7.01\n4.00", "+0.01\n-0.01", "[ Budi ]", "Lolos"],
        ["19/08/2026", "Timbangan Digital 150Kg", "Anak Timbang 5.00 Kg", "5.00 Kg", "0.00", "[ Budi ]", "Lolos"],
        ["20/08/2026", "pH Meter Digital A", "Buffer pH 7.00\nBuffer pH 10.01", "7.00\n10.02", "0.00\n+0.01", "[ Budi ]", "Lolos"]
    ]
    add_styled_table(doc, headers_cal, data_cal, [Cm(2.2), Cm(3.2), Cm(3.2), Cm(2.2), Cm(1.8), Cm(1.6), Cm(1.5)])

    add_heading_2(doc, "Formulir QC-02: Logbook Penyimpanan Sampel Pertinggal (Retained Sample)")
    headers_ret = ["No", "Tanggal Simpan", "Nama Produk & Merek", "No. Bets", "Jumlah Kemasan", "Masa Simpan (Exp)", "Petugas QC"]
    data_ret = [
        ["1", "19/08/2026", "Sabun Cuci Piring Lime 450mL", "SCP-260819-01", "2 Botol @ 450 mL", "19/08/2028", "[ Budi ]"],
        ["2", "20/08/2026", "Pembersih Lantai Pine 5 Liter", "PL-260820-01", "2 Jerigen @ 500 mL", "20/08/2028", "[ Budi ]"],
        ["3", "21/08/2026", "Pemutih Pakaian 5.25% 5 Liter", "BL-260821-01", "2 Jerigen @ 500 mL", "21/08/2028", "[ Budi ]"]
    ]
    add_styled_table(doc, headers_ret, data_ret, [Cm(0.8), Cm(2.2), Cm(4.5), Cm(2.8), Cm(2.5), Cm(2.2), Cm(1.5)])

    add_signature_block(doc, "Petugas QC & Laboratorium", "Disetujui Oleh PJT")
    
    file_path = os.path.join(OUTPUT_DIR, "07_DEPT_QC_DAN_LABORATORIUM.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# DOC 08: DIVISI PRODUKSI & MANUFAKTUR
# ==============================================================================
def create_doc_prod():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Divisi Produksi & Manufaktur", "Pedoman Operasional, Instruksi Kerja (IK) Mixing, Line Clearance & Sanitasi Mesin Standar CPKRTB", "SOP-PRD-008", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. STANDAR OPERASIONAL PROSEDUR PRODUKSI")
    
    add_heading_2(doc, "SOP-PRD-01: Prosedur Pembersihan Lini Kerja (Line Clearance)")
    add_bullet(doc, "Sebelum memulai batch baru, bersihkan seluruh sisa bahan, kemasan, dan etiket dari batch sebelumnya dari meja kerja.", "1. Pembersihan Area: ");
    add_bullet(doc, "Pastikan tangki mixer dan nozel mesin pengisi telah dibilas bersih dan berlabel 'BERSIH - SIAP PAKAI'.", "2. Verifikasi Peralatan: ");
    add_bullet(doc, "Kepala Produksi menandatangani persetujuan Line Clearance sebelum bahan batch baru ditimbang.", "3. Otorisasi Mulai: ");

    add_heading_2(doc, "SOP-PRD-02: Instruksi Kerja (IK) Pembuatan Pemutih Pakaian (Sodium Hypochlorite 5.25%)")
    add_bullet(doc, "Bahan Baku: Air Bersih / Demineral (90%), Sodium Hypochlorite Teknis 12% (10%), Soda Ash/NaOH penstabil pH q.s.", "1. Komposisi Baku: ");
    add_bullet(doc, "APD Wajib: Operator wajib mengenakan masker respirator uap kimia, kacamata goggle rapat, sarung tangan karet tebal, dan celemek karet PVC.", "2. K3 & APD: ");
    add_bullet(doc, "Proses Pencampuran: Masukkan air baku ke tangki HDPE/Fiber. Tambahkan Sodium Hypochlorite perlahan dengan pengadukan kecepatan rendah (mencegah percikan).", "3. Tahap Mixing: ");
    add_bullet(doc, "Penyesuaian pH: Periksa pH larutan, jaga agar tetap berada pada rentang pH 11.00 – 12.50 untuk stabilitas klorin aktif.", "4. Stabilisasi: ");
    add_bullet(doc, "Pengemasan: Isi ke dalam botol/jerigen HDPE tebal yang buram/tahan cahaya. Tutup dengan seal anti-bocor dan degassing cap bila perlu.", "5. Pengemasan: ");

    add_heading_1(doc, "2. LOGBOOK & FORMULIR HARIAN PRODUKSI")
    
    add_heading_2(doc, "Formulir PRD-01: Lembar Pemeriksaan Line Clearance Sebelum Produksi")
    headers_lc = ["No", "Item Pemeriksaan Lini Kerja", "Kondisi Aktual", "Paraf Operator", "Paraf Ka. Produksi"]
    data_lc = [
        ["1", "Area ruang mixing bebas dari sisa bahan batch sebelumnya", "Bersih & Rapi", "[ Joko ]", "[ Rudi ]"],
        ["2", "Tangki mixing telah dibilas dan bebas bau sisa produk lama", "Bersih / Kering", "[ Joko ]", "[ Rudi ]"],
        ["3", "Timbangan digital dalam keadaan nol (tara) dan bersih", "Terkalibrasi", "[ Joko ]", "[ Rudi ]"],
        ["4", "Kemasan botol dan label stiker sesuai dengan batch yang akan diproduksi", "Sesuai PO / Desain", "[ Joko ]", "[ Rudi ]"],
        ["5", "Operator telah mengenakan APD lengkap dan sehat", "APD Lengkap", "[ Joko ]", "[ Rudi ]"]
    ]
    add_styled_table(doc, headers_lc, data_lc, [Cm(0.8), Cm(6.5), Cm(3.2), Cm(2.2), Cm(2.3)])

    add_signature_block(doc, "Kepala Produksi", "Disetujui Oleh PJT")
    
    file_path = os.path.join(OUTPUT_DIR, "08_DEPT_PRODUKSI_DAN_MANUFAKTUR.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# DOC 09: DIVISI GUDANG & LOGISTIK
# ==============================================================================
def create_doc_wh():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Divisi Gudang & Logistik", "Pedoman Tata Kelola Gudang, Penyimpanan Kimia Berbahaya & Sistem Distribusi FIFO/FEFO Standar ISO 9001", "SOP-GUD-009", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. PROSEDUR TATA KELOLA GUDANG PABRIK")
    
    add_heading_2(doc, "SOP-GUD-01: Prosedur Penerimaan & Karantina Bahan Masuk")
    add_bullet(doc, "Pemeriksaan Fisik: Cek keutuhan drum/jerigen, label supplier, tanggal kedaluwarsa, dan sertifikat analisis (CoA supplier).", "1. Cek Surat Jalan: ");
    add_bullet(doc, "Status Karantina: Tempel stiker kuning bertuliskan 'DIKARANTINA - MENUNGGU UJI QC'. Bahan belum boleh digunakan sebelum QC merilis.", "2. Label Karantina: ");
    add_bullet(doc, "Status Rilis: Setelah QC menyatakan lulus, ganti dengan stiker hijau 'DILULUSKAN' dan pindahkan ke rak penyimpanan utama.", "3. Label Lolos: ");

    add_heading_2(doc, "SOP-GUD-02: Prosedur Penyimpanan Bahan Kimia Khusus & Korosif")
    add_bullet(doc, "Bahan pemutih (Sodium Hypochlorite / Asam / Basa Kuat) wajib disimpan di area berventilasi khusus terpisah dari parfum dan bahan mudah terbakar.", "1. Zonasi Korosif: ");
    add_bullet(doc, "Semua jerigen/drum diletakkan di atas palet plastik dengan penampung tumpahan (*spill containment*).", "2. Paletisasi: ");

    add_heading_1(doc, "2. TEMPLATE KARTU STOK & SURAT JALAN")
    
    add_heading_2(doc, "Formulir GUD-01: Kartu Stok Gudang Bahan Baku & Kemasan")
    headers_stk = ["Tanggal", "No. Dokumen / PO", "Nama Bahan / Kemasan", "Masuk", "Keluar", "Sisa Stok", "Paraf Petugas"]
    data_stk = [
        ["19/08/2026", "PO-SUP-012", "SLES / Texapon 70%", "1.000 Kg", "-", "1.000 Kg", "[ Agus ]"],
        ["20/08/2026", "BPR-SCP-01", "SLES / Texapon 70%", "-", "100 Kg", "900 Kg", "[ Agus ]"],
        ["21/08/2026", "PO-SUP-015", "Botol 450 mL HDPE", "5.000 Pcs", "-", "5.000 Pcs", "[ Agus ]"],
        ["22/08/2026", "BPR-SCP-01", "Botol 450 mL HDPE", "-", "1.000 Pcs", "4.000 Pcs", "[ Agus ]"]
    ]
    add_styled_table(doc, headers_stk, data_stk, [Cm(2.0), Cm(2.8), Cm(4.2), Cm(1.8), Cm(1.8), Cm(1.8), Cm(1.6)])

    add_heading_2(doc, "Formulir GUD-02: Surat Jalan Pengiriman Produk Jadi (Delivery Order)")
    headers_do = ["No", "Nama Produk & Varian", "No. Izin Edar PKD", "Nomor Bets", "Kemasan", "Kuantitas", "Status QC (CoA)"]
    data_do = [
        ["1", "Sabun Cuci Piring Lime", "KEMENKES RI PKD [No PKD]", "SCP-260819-01", "Karton (24 Btl @ 450mL)", "40 Karton (960 Btl)", "CoA Dilampirkan"],
        ["2", "Pemutih Pakaian 5.25%", "KEMENKES RI PKD [No PKD]", "BL-260821-01", "Jerigen 5 Liter", "100 Jerigen", "CoA Dilampirkan"]
    ]
    add_styled_table(doc, headers_do, data_do, [Cm(0.8), Cm(3.8), Cm(3.2), Cm(2.4), Cm(2.8), Cm(2.2), Cm(2.3)])

    add_signature_block(doc, "Petugas Gudang / Ekspedisi", "Penerima / PT Distributor")
    
    file_path = os.path.join(OUTPUT_DIR, "09_DEPT_GUDANG_DAN_LOGISTIK.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# DOC 10: DIVISI MARKETING & LAYANAN MAKLOON
# ==============================================================================
def create_doc_mkt():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Divisi Marketing & Makloon", "Pedoman Pelayanan Makloon, Perjanjian Kerahasiaan (NDA), SPK & Survei Kepuasan Pelanggan ISO 9001", "SOP-MKT-010", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. PROSEDUR PELAYANAN KLIEN MAKLOON")
    
    add_heading_2(doc, "SOP-MKT-01: Prosedur Penerimaan & Onboarding Klien Makloon")
    add_bullet(doc, "Identifikasi Kebutuhan Klien: Konsultasi formula, aroma, warna, ukuran kemasan, target pasar, dan estimasi volume produksi bulanan.", "1. Konsultasi Awal: ");
    add_bullet(doc, "Verifikasi Legalitas & HKI: Wajib memeriksa legalitas badan usaha klien dan bukti kepemilikan/pendaftaran merek dari DJKI Kemenkumham.", "2. Uji Tuntas HKI: ");
    add_bullet(doc, "Penandatanganan NDA: Menandatangani Non-Disclosure Agreement sebelum pertukaran sampel atau diskusi formula khusus.", "3. Kerahasiaan: ");
    add_bullet(doc, "Penerbitan Quotation & Kontrak: Menerbitkan surat penawaran harga resmi dan menandatangani Kontrak Perjanjian Makloon.", "4. Legalitas Kontrak: ");

    add_heading_1(doc, "2. FORMULIR & TEMPLATE DOKUMEN MARKETING")
    
    add_heading_2(doc, "Formulir MKT-01: Surat Perintah Kerja Makloon (Purchase Order / SPK Makloon)")
    headers_spk = ["No", "Item Pesanan Makloon", "Spesifikasi Formula & Kemasan", "Jumlah Pesanan", "Harga Satuan", "Total Biaya"]
    data_spk = [
        ["1", "Sabun Cuci Piring Varian Jeruk Nipis", "Formula Kemenkes, Kemasan Botol 450 mL + Stiker", "2.000 Botol", "Rp 4.200", "Rp 8.400.000"],
        ["2", "Pembersih Lantai Varian Pine Floral", "Formula Kemenkes, Kemasan Jerigen 5 Liter", "100 Jerigen", "Rp 29.500", "Rp 2.950.000"],
        ["-", "SUBTOTAL PEMESANAN", "-", "-", "-", "Rp 11.350.000"],
        ["-", "UANG MUKA (DP 50%)", "Dibayarkan saat PO ditandatangani", "-", "-", "Rp 5.675.000"],
        ["-", "PELUNASAN (50%)", "Dibayarkan sebelum pengiriman barang", "-", "-", "Rp 5.675.000"]
    ]
    add_styled_table(doc, headers_spk, data_spk, [Cm(0.8), Cm(4.2), Cm(4.8), Cm(2.0), Cm(2.2), Cm(2.5)])

    add_heading_2(doc, "Formulir MKT-02: Survei Kepuasan Pelanggan / Distributor (ISO 9001 Klausul 9.1.2)")
    headers_srv = ["No", "Aspek Penilaian Layanan & Kualitas", "Sangat Puas (5)", "Puas (4)", "Cukup (3)", "Kurang (2)", "Sangat Kurang (1)"]
    data_srv = [
        ["1", "Konsistensi Kualitas & Mutu Fisik Produk (Busa, Aroma, pH)", "[  ]", "[  X  ]", "[  ]", "[  ]", "[  ]"],
        ["2", "Ketepatan Waktu Pengiriman Barang (Lead Time)", "[  ]", "[  X  ]", "[  ]", "[  ]", "[  ]"],
        ["3", "Kelengkapan Dokumen Pengiriman (CoA & Surat Jalan)", "[  X  ]", "[  ]", "[  ]", "[  ]", "[  ]"],
        ["4", "Respon Cepat & Komunikasi Tim Marketing/Pabrik", "[  X  ]", "[  ]", "[  ]", "[  ]", "[  ]"],
        ["5", "Keamanan & Kerapian Pengemasan Produk ke Karton", "[  ]", "[  X  ]", "[  ]", "[  ]", "[  ]"]
    ]
    add_styled_table(doc, headers_srv, data_srv, [Cm(0.8), Cm(6.5), Cm(1.8), Cm(1.6), Cm(1.6), Cm(1.6), Cm(1.8)])

    add_signature_block(doc, "Tim Marketing & Makloon", "Manajemen / Direktur Mitra")
    
    file_path = os.path.join(OUTPUT_DIR, "10_DEPT_MARKETING_DAN_LAYANAN_MAKLOON.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# DOC 11: DIVISI MANAJEMEN, HR & ISO 9001 (QMS)
# ==============================================================================
def create_doc_mgmt():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Manajemen, HR & QMS", "Manual Mutu, Uraian Tugas Lengkap Tiap Divisi, SOP Audit Internal & CAPA Standar ISO 9001:2015", "SOP-MGT-011", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. KEBIJAKAN MUTU PERUSAHAAN (ISO 9001 KLAUSUL 5.2)")
    add_p(doc, "PT [Nama Perusahaan] berkomitmen untuk menjadi mitra industri manufaktur PKRT terpercaya di Indonesia dengan: (1) Menjamin seluruh proses produksi memenuhi standar CPKRTB dan ISO 9001:2015; (2) Menghasilkan produk pembersih rumah tangga yang aman, efektif, dan konsisten; (3) Memberikan layanan makloon prima dengan ketepatan waktu pengiriman dan legalitas terjamin; (4) Melakukan perbaikan berkelanjutan secara sistematis pada seluruh proses operasional.")

    add_heading_1(doc, "2. URAIAN TUGAS LENGKAP TIAP DIVISI (JOB DESCRIPTION)")
    
    headers_jd = ["Divisi / Jabatan", "Kualifikasi Minimal", "Uraian Tugas Pokok & Tanggung Jawab Utama"]
    data_jd = [
        ["Direktur Utama", "Pimpinan Perusahaan", "Menetapkan arah kebijakan perusahaan, mengesahkan perjanjian makloon dan anggaran sarana pabrik, serta memimpin Tinjauan Manajemen (Management Review)."],
        ["Penanggung Jawab Teknis (PJT)", "D3/S1 Farmasi / Kimia", "Menjamin kepatuhan CPKRTB Kemenkes, mengelola registrasi izin edar PKD, menyetujui Master Formula, menerbitkan CoA, dan mengawasi sistem mutu."],
        ["Kepala Produksi", "SMK Kimia / D3", "Mengatur jadwal mixing, mengawasi proses penimbangan bahan baku, memverifikasi Line Clearance, dan memastikan pengisian Batch Processing Record harian."],
        ["Petugas QC & Lab", "SMK Analis Kimia", "Melakukan sampling bahan baku/produk jadi, menguji parameter pH, bobot jenis, stabilitas, mengelola retained sample, dan kalibrasi alat uji."],
        ["Petugas Gudang", "SMA / SMK", "Menerima dan memeriksa bahan datang, mengatur penataan palet FIFO/FEFO, menjaga keamanan bahan korosif, dan mencatat mutasi pada Kartu Stok."],
        ["Petugas Marketing", "D3 / S1 Semua Jurusan", "Menerima prospek makloon, melakukan uji tuntas merek DJKI klien, menerbitkan Quotation & SPK, serta melakukan survei kepuasan pelanggan berkala."]
    ]
    add_styled_table(doc, headers_jd, data_jd, [Cm(3.2), Cm(3.0), Cm(9.5)])

    add_heading_1(doc, "3. SISTEM PERBAIKAN BERKELANJUTAN (CAPA - KLAUSUL 10.2)")
    
    headers_capa = ["No. Laporan CAPA", "Deskripsi Ketidaksesuaian (Temuan)", "Akar Penyebab Masalah (Root Cause)", "Tindakan Korektif & Preventif", "Batas Waktu & PIC", "Verifikasi PJT"]
    data_capa = [
        ["CAPA-2026-001", "Ditemukan 2 botol sabun cuci piring merembes pada bagian tutup", "Torsi penutupan botol manual kurang kencang oleh operator", "Pengadaan alat bantu penutup botol pneumatik & training operator", "30 Agt 2026\n[ Ka. Prod ]", "[ OK / Ditutup ]"],
        ["CAPA-2026-002", "Suhu gudang penyimpanan klorin mencapai 32°C di siang hari", "Ventilasi udara alami kurang memadai saat cuaca panas", "Pemasangan exhaust fan tambahan dan thermometer digital alarm", "05 Sep 2026\n[ Teknisi ]", "[ OK / Ditutup ]"]
    ]
    add_styled_table(doc, headers_capa, data_capa, [Cm(2.2), Cm(3.8), Cm(3.5), Cm(3.8), Cm(2.2), Cm(1.8)])

    add_signature_block(doc, "Management Representative (MR) / PJT", "Direktur Utama")
    
    file_path = os.path.join(OUTPUT_DIR, "11_DEPT_MANAJEMEN_HR_DAN_ISO9001.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

if __name__ == "__main__":
    create_doc_pjt()
    create_doc_qc()
    create_doc_prod()
    create_doc_wh()
    create_doc_mkt()
    create_doc_mgmt()
    print("ALL 6 DEPARTMENTAL ISO 9001 PKRT DOCUMENTS SUCCESSFULLY GENERATED!")

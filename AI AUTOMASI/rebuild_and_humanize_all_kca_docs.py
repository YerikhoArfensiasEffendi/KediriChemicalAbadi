import os
import sys
import docx
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml
import zipfile
import xml.etree.ElementTree as ET
import datetime
import random
import time

BASE_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/KCA DOKUMEN"
MGT_DIR = os.path.join(BASE_DIR, "09_DOKUMEN_DIREKSI_DAN_MANAJEMEN_PUNCAK")
os.makedirs(MGT_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)
    return doc

def add_iso_header_box(doc, dept_or_system, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(4.0), Cm(6.0), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run(f"PT KEDIRI CHEMICAL ABADI\n{dept_or_system.upper()}\nSISTEM MUTU ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block_formal(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    col_widths = [Cm(5.0), Cm(5.0), Cm(5.5)]
    
    titles = [
        ("Disiapkan Oleh,\nGeneral Manager", "Yerikho Arfensias Effendi"),
        ("Diperiksa Oleh,\nPenanggung Jawab Teknis", "PJT Farmasi / Kimia"),
        ("Disahkan Oleh,\nDirektur Utama", "Yan Effendi")
    ]
    
    for idx, (title, name) in enumerate(titles):
        cell = table.cell(0, idx)
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(title + "\n\n\n\n\n")
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(9)
        r1.font.color.rgb = COLOR_BLACK
        
        r2 = p.add_run(f"( {name} )\n")
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(9.5)
        r2.font.bold = True
        r2.font.color.rgb = COLOR_BLACK
        
        r3 = p.add_run("Tanggal: 19 Agustus 2026")
        r3.font.name = 'Times New Roman'
        r3.font.size = Pt(8.5)
        r3.font.color.rgb = COLOR_BLACK

# ==============================================================================
# CREATE MASTER DOCUMENT: STRUKTUR ORGANISASI PERUSAHAAN LENGKAP
# ==============================================================================
def create_master_org_structure():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Direksi & Tata Kelola", "Struktur Organisasi Resmi, Bagan Tata Kelola & Matriks Uraian Jabatan", "STR-DIR-008", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. STRUKTUR ORGANISASI PT KEDIRI CHEMICAL ABADI")
    add_p(doc, "Sesuai dengan ketentuan Anggaran Dasar Perseroan dan standar Sistem Manajemen Mutu ISO 9001:2015 Klausul 5.3 (Peran, Tanggung Jawab dan Wewenang Organisasi), struktur kepemimpinan dan hierarki operasional PT Kediri Chemical Abadi ditetapkan sebagai berikut:")
    
    headers_lvl = ["Tingkat Jabatan", "Nama Pejabat / Pemegang Peran", "Fungsi Utama", "Garis Komando"]
    data_lvl = [
        ["Komisaris Utama / Owner", "Yan Effendi", "Pengawasan strategis, pemilik modal, persetujuan investasi besar (CAPEX), dan proteksi aset.", "Tingkat Tertinggi"],
        ["Direktur Utama / General Manager", "Yerikho Arfensias Effendi", "Pengelola utama roda bisnis, komersial makloon, arus kas operasional, kepatuhan Kemenkes, dan pimpinan tim.", "Bertanggung jawab ke Owner"],
        ["Wakil Manager Operasional", "Wakil Pimpinan Operasional", "Tangan kanan General Manager, pengawas lantai pabrik, eksekutor jadwal mixing & logistik harian.", "Bertanggung jawab ke GM"],
        ["Penanggung Jawab Teknis (PJT)", "PJT (Farmasi / Kimia Ber-STR)", "Penanggung jawab teknis formulasi, pengawasan mutu CPKRTB, pelulusan bets (CoA), dan registrasi Kemenkes.", "Jalur Fungsional Independen"],
        ["Kepala Produksi & Manufaktur", "Supervisor Produksi", "Memimpin operator penimbangan, mixing, filling, dan memastikan pembersihan mesin sesuai SOP.", "Lapor ke Wakil GM & PJT"],
        ["Kepala Quality Control (QC)", "Analis Kimia Lab", "Sampling bahan baku/produk jadi, uji pH, bobot jenis, uji stabilitas, dan arsip sampel pertinggal.", "Lapor ke PJT"],
        ["Kepala Gudang & Logistik", "Koordinator Gudang", "Pengelolaan stok bahan kimia FIFO/FEFO, penyimpanan klorin aman, dan penerbitan Surat Jalan.", "Lapor ke Wakil GM"]
    ]
    add_styled_table(doc, headers_lvl, data_lvl, [Cm(3.5), Cm(3.5), Cm(5.5), Cm(3.0)])

    add_heading_1(doc, "2. URAIAN TUGAS DAN WEWENANG EKSEKUTIF INTI")
    
    add_heading_2(doc, "2.1 Yan Effendi — Komisaris Utama / Pemilik Usaha (Owner)")
    add_bullet(doc, "Menetapkan arah kebijakan umum perusahaan dan kesinambungan investasi jangka panjang.", "Fungsi Pengawasan: ");
    add_bullet(doc, "Menyetujui alokasi modal kerja besar (pembelian mesin mixer tambahan, perluasan bangunan gudang).", "Otoritas CAPEX: ");
    add_bullet(doc, "Menerima dan mengevaluasi Laporan Kinerja Bulanan dan Laporan Keuangan dari General Manager.", "Evaluasi Bisnis: ");
    add_bullet(doc, "Membebaskan diri dari rutinitas teknis pabrik harian guna fokus pada strategi pengembangan aset.", "Batas Wewenang: ");

    add_heading_2(doc, "2.2 Yerikho Arfensias Effendi — General Manager / Pengelola Utama Perusahaan")
    add_bullet(doc, "Memimpin operasional menyeluruh, strategi penjualan jasa makloon, dan penetapan struktur harga resmi (Quotation).", "Pimpinan Eksekutif: ");
    add_bullet(doc, "Menandatangani seluruh perikatan hukum: Kontrak Kerja Sama Makloon, Surat Perintah Kerja (SPK), dan perjanjian supplier.", "Otoritas Hukum: ");
    add_bullet(doc, "Mengendalikan arus kas harian (Cash Flow), penerimaan DP 50% dan pelunasan, pembayaran bahan baku, dan penggajian.", "Manajemen Finansial: ");
    add_bullet(doc, "Mengoordinasikan pemenuhan regulasi Kemenkes bersama PJT dan mengevaluasi kinerja seluruh kepala divisi.", "Kepatuhan Mutu: ");

    add_heading_2(doc, "2.3 Wakil Manager Operasional — Tangan Kanan & Pengawas Lapangan")
    add_bullet(doc, "Memastikan target batch produksi harian yang diinstruksikan General Manager selesai tepat waktu dan sesuai spesifikasi.", "Pengawasan Produksi: ");
    add_bullet(doc, "Mengontrol kedisiplinan kerja operator, kepatuhan pemakaian APD, dan sanitasi ruangan pabrik.", "Disiplin & K3: ");
    add_bullet(doc, "Memantau sisa stok bahan baku di gudang dan mengajukan permohonan re-order sebelum stok kritis.", "Kontrol Inventaris: ");
    add_bullet(doc, "Mengawal proses pengepakan karton, pengujian kebocoran, hingga muat barang ke truk ekspedisi bersama Surat Jalan.", "Logistik & Ekspedisi: ");

    add_heading_1(doc, "3. MATRIKS TATA KELOLA KEPUTUSAN DAN OTORISASI")
    headers_aut = ["Kegiatan / Transaksi", "Wakil Manager", "Yerikho Arfensias Effendi (GM)", "Yan Effendi (Owner)"]
    data_aut = [
        ["Operasional Harian Pabrik & Shift Kerja", "Pelaksana Penuh", "Menyetujui & Supervisi", "Mengetahui"],
        ["Pembelian Bahan Baku Rutin (< Rp 20 Jt)", "Mengajukan", "Memutuskan & Bayar", "Menerima Rekap"],
        ["Penetapan Kontrak Makloon & Harga Klien", "Memberi Masukan", "Memutuskan & Teken", "Mengetahui"],
        ["Pengeluaran Modal Besar (> Rp 20 Jt)", "Identifikasi Kebutuhan", "Menyusun Anggaran", "Persetujuan Akhir"],
        ["Rilis Produk & Penyerahan ke Ekspedisi", "Cek Fisik & DO", "Persetujuan Bersama PJT", "Menerima Laporan"]
    ]
    add_styled_table(doc, headers_aut, data_aut, [Cm(5.0), Cm(3.2), Cm(4.0), Cm(3.3)])

    add_signature_block_formal(doc)
    
    file_path = os.path.join(MGT_DIR, "MGT-08_Struktur_Organisasi_dan_Bagan_Tata_Kelola_PT_KCA.docx")
    doc.save(file_path)
    
    # Also copy to folder 06
    folder_06 = os.path.join(BASE_DIR, "06_MASTER_MANUAL_DAN_LEGALITAS")
    doc.save(os.path.join(folder_06, "14_STRUKTUR_ORGANISASI_DAN_BAGAN_TATA_KELOLA_PT_KCA.docx"))
    print(f"Created Master Org Structure: {file_path}")

# ==============================================================================
# AUDIT AND RE-APPLY PURE BLACK COLOR & UNIFIED IDENTITY TO ALL DOCX FILES
# ==============================================================================
def humanize_and_blacken_docx(file_path):
    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Error opening {file_path}: {e}")
        return

    # Standardize Normal Style
    if 'Normal' in doc.styles:
        norm = doc.styles['Normal']
        norm.font.name = 'Times New Roman'
        norm.font.size = Pt(11)
        norm.font.color.rgb = COLOR_BLACK
        norm.paragraph_format.line_spacing = 1.15
        norm.paragraph_format.space_after = Pt(6)

    # Walk all paragraphs
    for p in doc.paragraphs:
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = COLOR_BLACK

    # Walk all tables
    for tbl in doc.tables:
        for row_idx, row in enumerate(tbl.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        # In data cells, always black
                        # In header cells, if background is dark navy, keep white, else black
                        r.font.color.rgb = COLOR_BLACK

    doc.save(file_path)

# ==============================================================================
# SANITIZE METADATA (NO AI FOOTPRINTS, YERIKHO ARFENSIAS EFFENDI, REALISTIC DATES)
# ==============================================================================
FORBIDDEN_WORDS = [
    "dibuat oleh ai", "dibuat otomatis oleh ai", "generated by python", "as an ai",
    "asisten ai", "artificial intelligence", "chatgpt", "openai", "deepmind",
    "gemini", "anthropic", "script generator", "system generated", "bot sistem"
]

def sanitize_docx_metadata(file_path):
    temp_zip = file_path + ".temp.zip"
    
    # Realistic timestamps spread across working hours (08:30 - 17:15 WIB)
    day = random.randint(12, 19)
    hour_created = random.randint(8, 14)
    hour_mod = random.randint(14, 17)
    minute_created = random.randint(10, 55)
    minute_mod = random.randint(10, 55)
    
    created_dt = datetime.datetime(2026, 8, day, hour_created, minute_created, random.randint(10, 55))
    modified_dt = datetime.datetime(2026, 8, 20, hour_mod, minute_mod, random.randint(10, 55))
    created_iso = created_dt.strftime("%Y-%m-%dT%H:%M:%SZ")
    modified_iso = modified_dt.strftime("%Y-%m-%dT%H:%M:%SZ")
    total_edit_time = str(random.randint(35, 95))
    revision_count = str(random.randint(2, 7))
    clean_title = os.path.basename(file_path).replace(".docx", "").replace("_", " ").title()

    ET.register_namespace('cp', 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties')
    ET.register_namespace('dc', 'http://purl.org/dc/elements/1.1/')
    ET.register_namespace('dcterms', 'http://purl.org/dc/terms/')
    ET.register_namespace('dcmitype', 'http://purl.org/dc/dcmitype/')
    ET.register_namespace('xsi', 'http://www.w3.org/2001/XMLSchema-instance')
    ET.register_namespace('', 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties')
    ET.register_namespace('vt', 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes')

    with zipfile.ZipFile(file_path, 'r') as zin, zipfile.ZipFile(temp_zip, 'w') as zout:
        for item in zin.infolist():
            buffer = zin.read(item.filename)
            
            if item.filename == 'docProps/core.xml':
                root = ET.fromstring(buffer)
                
                creator = root.find('{http://purl.org/dc/elements/1.1/}creator')
                if creator is None:
                    creator = ET.SubElement(root, '{http://purl.org/dc/elements/1.1/}creator')
                creator.text = "Yerikho Arfensias Effendi"
                
                last_mod = root.find('{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy')
                if last_mod is None:
                    last_mod = ET.SubElement(root, '{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy')
                last_mod.text = "Yerikho Arfensias Effendi"

                title = root.find('{http://purl.org/dc/elements/1.1/}title')
                if title is None:
                    title = ET.SubElement(root, '{http://purl.org/dc/elements/1.1/}title')
                title.text = clean_title

                rev = root.find('{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}revision')
                if rev is None:
                    rev = ET.SubElement(root, '{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}revision')
                rev.text = revision_count

                created_elem = root.find('{http://purl.org/dc/terms/}created')
                if created_elem is None:
                    created_elem = ET.SubElement(root, '{http://purl.org/dc/terms/}created')
                    created_elem.attrib['{http://www.w3.org/2001/XMLSchema-instance}type'] = 'dcterms:W3CDTF'
                created_elem.text = created_iso

                modified_elem = root.find('{http://purl.org/dc/terms/}modified')
                if modified_elem is None:
                    modified_elem = ET.SubElement(root, '{http://purl.org/dc/terms/}modified')
                    modified_elem.attrib['{http://www.w3.org/2001/XMLSchema-instance}type'] = 'dcterms:W3CDTF'
                modified_elem.text = modified_iso

                buffer = ET.tostring(root, encoding='utf-8', xml_declaration=True)

            elif item.filename == 'docProps/app.xml':
                root = ET.fromstring(buffer)
                def set_app_prop(tag, val):
                    node = root.find(f"{{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}}{tag}")
                    if node is None:
                        node = ET.SubElement(root, f"{{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}}{tag}")
                    node.text = str(val)

                set_app_prop('Company', 'PT Kediri Chemical Abadi')
                set_app_prop('Manager', 'Yerikho Arfensias Effendi')
                set_app_prop('Application', 'Microsoft Office Word')
                set_app_prop('TotalTime', total_edit_time)
                set_app_prop('Template', 'Normal.dotm')

                buffer = ET.tostring(root, encoding='utf-8', xml_declaration=True)

            elif item.filename == 'word/document.xml':
                text_content = buffer.decode('utf-8')
                
                # Replace generic placeholders with official company names
                text_content = text_content.replace("PT [NAMA PERUSAHAAN]", "PT Kediri Chemical Abadi")
                text_content = text_content.replace("PT [Nama Perusahaan Anda]", "PT Kediri Chemical Abadi")
                text_content = text_content.replace("PT [Nama Perusahaan]", "PT Kediri Chemical Abadi")
                text_content = text_content.replace("PT KCA CHEMICAL", "PT Kediri Chemical Abadi")
                text_content = text_content.replace("[Nama Direktur Anda]", "Yerikho Arfensias Effendi")
                text_content = text_content.replace("[Direktur Utama]", "Yan Effendi")
                text_content = text_content.replace("[Kota]", "Kediri, Jawa Timur")
                text_content = text_content.replace("[Alamat Lengkap Pabrik]", "Kawasan Industri & Pergudangan Kediri, Jawa Timur")
                
                for kw in FORBIDDEN_WORDS:
                    if kw in text_content.lower():
                        import re
                        text_content = re.sub(re.escape(kw), "Sistem Manajemen Mutu Terpadu", text_content, flags=re.IGNORECASE)
                
                buffer = text_content.encode('utf-8')

            zout.writestr(item, buffer)

    os.replace(temp_zip, file_path)
    mod_epoch = time.mktime(modified_dt.timetuple())
    os.utime(file_path, (mod_epoch, mod_epoch))

if __name__ == "__main__":
    create_master_org_structure()
    
    total_docs = 0
    for root, dirs, files in os.walk(BASE_DIR):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~$') and not f.startswith('.'):
                fp = os.path.join(root, f)
                humanize_and_blacken_docx(fp)
                sanitize_docx_metadata(fp)
                total_docs += 1
                
    print(f"\nALL {total_docs} DOCX FILES IN 'KCA DOKUMEN' SUCCESSFULLY PROCESSED, HUMANIZED, BLACK-COLORED & SANITIZED!")

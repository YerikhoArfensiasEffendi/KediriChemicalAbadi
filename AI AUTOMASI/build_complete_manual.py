import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

OUTPUT_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/dokumen_iso9001_pkrt"
os.makedirs(OUTPUT_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PT [NAMA PERUSAHAAN]\nMANUFAKTUR PKRT\nSISTEM MUTU CPKRTB")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block(doc, left_title="Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", right_title="Disetujui Oleh,\nDirektur Utama", left_name="( ............................................ )", right_name="( ............................................ )"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    col_widths = [Cm(7.5), Cm(8.0)]
    
    cell_l = table.cell(0, 0)
    cell_l.width = col_widths[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l1 = p_l.add_run(left_title + "\n\n\n\n\n" + left_name + "\nTanggal: ...................................")
    r_l1.font.name = 'Times New Roman'
    r_l1.font.size = Pt(9.5)
    r_l1.font.color.rgb = COLOR_BLACK
    
    cell_r = table.cell(0, 1)
    cell_r.width = col_widths[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r.add_run(right_title + "\n\n\n\n\n" + right_name + "\nTanggal: ...................................")
    r_r1.font.name = 'Times New Roman'
    r_r1.font.size = Pt(9.5)
    r_r1.font.color.rgb = COLOR_BLACK

def build_complete_manual():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Manual Sistem Dokumentasi Mutu Pabrik PKRT Standar CPKRTB & ISO 9001:2015", "MAN-CPKRTB-012", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "PENDAHULUAN & TUJUAN")
    add_p(doc, "Dokumen Manual ini merupakan panduan master sistem dokumentasi mutu sarana industri manufaktur PKRT bagi PT [Nama Perusahaan]. Disusun secara sistematis untuk memenuhi persyaratan Cara Pembuatan PKRT yang Baik (CPKRTB) sesuai Permenkes No. 14 Tahun 2021 dan standar Sistem Manajemen Mutu ISO 9001:2015, guna menjamin kesiapan audit berkala oleh Kementerian Kesehatan RI maupun Dinas Kesehatan.")
    
    # -------------------------------------------------------------------------
    # 1. DOKUMEN DENAH & FASILITAS FISIK PABRIK
    # -------------------------------------------------------------------------
    add_heading_1(doc, "1. DOKUMEN DENAH & FASILITAS FISIK PABRIK")
    add_p(doc, "Fasilitas fisik pabrik wajib dirancang untuk mencegah terjadinya kontaminasi silang dan menjamin keamanan keselamatan kerja kimiawi:")
    
    add_bullet(doc, "Menampilkan sekat dinding pemisah antar area: Gudang Bahan Baku, Ruang Penimbangan, Ruang Mixing/Produksi, Ruang Filling/Kemas Primer, Ruang Kemas Sekunder/Karton, Gudang Barang Jadi, Ruang Karantina, Ruang PJT/QC, dan Toilet/Loker Karyawan.", "1.1 Denah Tata Ruang Pabrik (Layout 2D): ");
    add_bullet(doc, "Alur proses pergerakan bahan mentah masuk -> ditimbang -> dimixing -> dikemas primer -> dikemas karton -> disimpan di gudang jadi (menerapkan sistem alur satu arah tanpa perpotongan arus silang).", "1.2 Denah Alur Barang (Material Flow): ");
    add_bullet(doc, "Jalur masuk personil dari pintu masuk -> ruang ganti/loker -> pemakaian APD -> wastafel cuci tangan higienis -> ruang produksi.", "1.3 Denah Alur Personel (Personnel Flow): ");
    add_bullet(doc, "Pemetaan titik penempatan tabung pemadam api (APAR jenis Powder / CO2) dan jalur evakuasi pintu darurat bebas hambatan.", "1.4 Denah APAR & Jalur Evakuasi K3: ");
    add_bullet(doc, "Pemetaan titik perangkap hama (Fly Catcher / Perangkap Tikus) di sekeliling dinding dalam dan luar bangunan pabrik.", "1.5 Denah Pengendalian Hama (Pest Control): ");
    
    add_heading_2(doc, "Tabel Inventaris Mesin & Peralatan Produksi")
    headers_inv = ["No", "Nama Alat / Mesin", "Spesifikasi & Material", "Kapasitas", "Jumlah", "Tahun Pengadaan", "Status Kalibrasi"]
    data_inv = [
        ["1", "Tangki Mixer Sabun & Kimia", "Stainless Steel SUS 304 / HDPE Heavy Duty", "1.000 Liter", "1 Unit", "2026", "Tervalidasi"],
        ["2", "Timbangan Digital Duduk", "Platform Timbang Presisi Elektronik", "150 Kg (d=10g)", "1 Unit", "2026", "Terkalibrasi"],
        ["3", "Timbangan Digital Analitik", "Presisi Laboratorium / Timbang Parfum", "5 Kg (d=0.1g)", "1 Unit", "2026", "Terkalibrasi"],
        ["4", "Mesin Filling Cairan Semi-Otomatis", "Piston Pneumatik Anti-Karat", "100 - 5.000 mL", "1 Unit", "2026", "Tervalidasi"],
        ["5", "pH Meter Digital Portabel", "Elektroda Kaca Glass Bulb + ATC", "0.00 - 14.00 pH", "1 Unit", "2026", "Terkalibrasi"],
        ["6", "Piknometer Kaca", "Volume Uji Bobot Jenis", "25 mL", "2 Unit", "2026", "Tervalidasi"]
    ]
    add_styled_table(doc, headers_inv, data_inv, [Cm(0.8), Cm(3.8), Cm(4.0), Cm(2.0), Cm(1.5), Cm(2.0), Cm(2.2)])

    # -------------------------------------------------------------------------
    # 2. DOKUMEN STRUKTUR ORGANISASI & PERSONALIA
    # -------------------------------------------------------------------------
    add_heading_1(doc, "2. DOKUMEN STRUKTUR ORGANISASI & PERSONALIA")
    
    add_bullet(doc, "Bagan hierarki resmi: Pimpinan Perusahaan (Direktur) -> Penanggung Jawab Teknis (PJT) -> Kepala Produksi, QC/Laboratorium, dan Petugas Gudang/Logistik.", "2.1 Bagan Struktur Organisasi Pabrik: ");
    add_bullet(doc, "Rincian tugas tertulis dan wewenang pengesahan mutu untuk PJT (menyetujui rilis batch & regulasi) serta Kepala Produksi (mengawasi proses mixing & filling).", "2.2 Uraian Tugas (Job Description): ");
    add_bullet(doc, "Pemeriksaan kesehatan awal masuk kerja dan pemantauan harian bebas luka terbuka/penyakit kulit bagi personil pengolahan langsung.", "2.3 Catatan Higiene & Kesehatan Personel: ");
    add_bullet(doc, "Jadwal dan dokumentasi pelatihan tahunan mencakup modul dasar CPKRTB, Higiene Sanitasi, Penanganan Bahan Berbahaya (B3), dan K3 Kimia.", "2.4 Program & Catatan Pelatihan Internal: ");

    # -------------------------------------------------------------------------
    # 3. DOKUMEN TEKNIS & MASTER PRODUKSI
    # -------------------------------------------------------------------------
    add_heading_1(doc, "3. DOKUMEN TEKNIS & MASTER PRODUKSI")
    
    add_bullet(doc, "Dokumen resmi berstempel PJT memuat komposisi kualitatif-kuantitatif 100%, urutan pencampuran, waktu mixing, dan batas kendali kritis pH/viskositas (Sabun Cuci Piring & Pemutih 5,25%).", "3.1 Master Formula Produk: ");
    add_bullet(doc, "Arsip lengkap Certificate of Analysis (CoA) dan Safety Data Sheet (MSDS) dari pabrik pemasok resmi untuk setiap bahan aktif kimia yang dibeli.", "3.2 Spesifikasi Bahan Baku & Bahan Kemas: ");
    add_bullet(doc, "Format kode unik batch terstruktur: [KODE_PRODUK]-[YYMMDD]-[NO_BETS], contoh: SCP-260819-01 (Sabun Cuci Piring, Tanggal 19-08-2026, Bets ke-01).", "3.3 Sistem Kodefikasi Nomor Bets (Batch Numbering): ");

    # -------------------------------------------------------------------------
    # 4. DAFTAR 20 STANDAR OPERASIONAL PROSEDUR (SOP) WAJIB CPKRTB
    # -------------------------------------------------------------------------
    add_heading_1(doc, "4. KOMPENDIUM 20 SOP WAJIB STANDAR CPKRTB")
    
    headers_sop = ["No", "Kode SOP", "Nama Prosedur Tetap (SOP)", "Ruang Lingkup & Fokus Utama", "Penanggung Jawab"]
    data_sop = [
        ["1", "SOP-SAN-001", "Pembersihan & Sanitasi Ruangan Pabrik", "Jadwal dan metode sanitasi lantai, dinding, dan exhaust", "Petugas Sanitasi"],
        ["2", "SOP-EQP-002", "Pembersihan & Perawatan Mesin/Tangki", "Prosedur pembilasan dan pembersihan tangki sebelum ganti varian", "Operator & Ka. Prod"],
        ["3", "SOP-HIG-003", "Higiene Personel & Pemakaian APD", "Kewajiban pemakaian masker, boots, kacamata goggle, celemek", "Seluruh Karyawan"],
        ["4", "SOP-PST-004", "Pengendalian Hama (Pest Control)", "Monitoring berkala titik perangkap lalat, kecoa, dan tikus", "Petugas Gudang"],
        ["5", "SOP-GUD-005", "Penerimaan & Pemeriksaan Bahan Masuk", "Pengecekan segel drum, keutuhan kemasan, dan CoA supplier", "Petugas Gudang"],
        ["6", "SOP-GUD-006", "Penyimpanan Kimia Berbahaya/Korosif", "Zonasi khusus Sodium Hypochlorite, ventilasi, dan palet spill", "Petugas Gudang"],
        ["7", "SOP-GUD-007", "Pengeluaran Bahan Metode FIFO/FEFO", "Pengeluaran bahan baku berdasarkan tanggal kedatangan & exp", "Petugas Gudang"],
        ["8", "SOP-GUD-008", "Penyimpanan & Pengiriman Produk Jadi", "Penyusunan karton pada palet dan penerbitan Surat Jalan", "Petugas Ekspedisi"],
        ["9", "SOP-PRD-009", "Penimbangan Bahan Baku", "Prosedur tara timbangan dan verifikasi kesesuaian bobot timbang", "Petugas Timbang & QC"],
        ["10", "SOP-PRD-010", "Proses Pengolahan / Mixing Cairan", "Instruksi kerja urutan pencampuran bahan hingga larutan homogen", "Operator Mixing"],
        ["11", "SOP-PRD-011", "Pengisian (Filling) & Penutupan Kemasan", "Pengaturan volume isi, penutupan rapat, dan uji kedap bocor", "Operator Filling"],
        ["12", "SOP-PRD-012", "Pelabelan Kemasan & Pengemasan Karton", "Pengecekan nomor PKD, batch, tanggal kedaluwarsa pada stiker", "Operator Kemas"],
        ["13", "SOP-QC-013", "Pengambilan Sampel Uji (Sampling)", "Protokol sampling bahan baku masuk, in-process, dan produk jadi", "Petugas QC"],
        ["14", "SOP-QC-014", "Pengujian Kualitas Produk Jadi di Lab", "Metode uji organoleptik, nilai pH meter, dan berat jenis", "Petugas QC"],
        ["15", "SOP-QC-015", "Penyimpanan Sampel Pertinggal (Retain)", "Penyimpanan 2 botol sampel per bets selama minimal 2 tahun", "Petugas QC & PJT"],
        ["16", "SOP-QC-016", "Penerbitan Sertifikat Analisis (CoA)", "Prosedur validasi hasil uji dan penandatanganan CoA oleh PJT", "PJT"],
        ["17", "SOP-MKT-017", "Penanganan Keluhan Pelanggan (Complaint)", "Alur investigasi komplain mutu dan tindakan penggantian barang", "Marketing & PJT"],
        ["18", "SOP-REG-018", "Penarikan Produk dari Peredaran (Recall)", "Prosedur sistematis penarikan produk jika ada cacat mutu fatal", "Pimpinan & PJT"],
        ["19", "SOP-PRD-019", "Penanganan Produk Cacat (Reject/Rework)", "Tata cara penanganan larutan yang tidak lolos uji spesifikasi", "Ka. Prod & PJT"],
        ["20", "SOP-QMS-020", "Audit Internal & Inspeksi Diri Pabrik", "Inspeksi diri berkala 6 bulan sekali terhadap kepatuhan CPKRTB", "Tim Audit Internal"]
    ]
    add_styled_table(doc, headers_sop, data_sop, [Cm(0.8), Cm(2.6), Cm(5.2), Cm(5.2), Cm(2.5)])

    # -------------------------------------------------------------------------
    # 5. TEMPLATE 8 FORMULIR & LOGBOOK HARIAN
    # -------------------------------------------------------------------------
    add_heading_1(doc, "5. TEMPLATE 8 FORMULIR, LOGBOOK & REKAMAN MUTU HARIAN")
    
    add_heading_2(doc, "Formulir 01: Catatan Pengolahan Bets (Batch Processing Record)")
    add_p(doc, "Formulir yang mencatat seluruh riwayat pembuatan satu nomor batch produksi: nama produk, varian, tanggal produksi, hasil penimbangan bahan aktual, waktu proses pengadukan, hasil uji in-process, jumlah hasil kemasan jadi, serta tanda tangan persetujuan pelulusan dari PJT.")

    add_heading_2(doc, "Formulir 02: Kartu Stok Bahan Baku & Kemasan")
    add_p(doc, "Kartu gantung pada setiap palet barang di gudang yang mencatat tanggal mutasi, nomor dokumen pesanan (PO) atau nomor batch pemakaian, kuantitas masuk, kuantitas keluar, sisa saldo stok, dan paraf petugas gudang.")

    add_heading_2(doc, "Formulir 03: Lembar Pemeriksaan QC Produk Jadi")
    add_p(doc, "Lembar verifikasi laboratorium pengujian mutu produk jadi sebelum kemasan dimasukkan ke dalam karton: pengujian bentuk fisik, warna sediaan, aroma, rentang nilai pH, berat jenis (g/mL), serta uji kebocoran tutup kemasan.")

    add_heading_2(doc, "Formulir 04: Format Sertifikat Analisis (Certificate of Analysis - CoA)")
    add_p(doc, "Lembar sertifikat jaminan mutu resmi yang diterbitkan dan ditandatangani oleh PJT untuk diserahkan kepada pihak pembeli atau PT Distributor bersamaan dengan pengiriman setiap batch produk.")

    add_heading_2(doc, "Formulir 05: Logbook Pembersihan Ruangan & Sanitasi Alat")
    add_p(doc, "Checklist harian pembersihan ruang produksi dan tangki mixer: mencatat tanggal, area yang dibersihkan, desinfektan yang digunakan, paraf petugas pelaksana, dan paraf supervisor pemeriksa.")

    add_heading_2(doc, "Formulir 06: Logbook Pemantauan Suhu & Kelembaban Gudang")
    add_p(doc, "Tabel pencatatan suhu (°C) dan kelembaban udara (%RH) area gudang bahan kimia yang diukur 2 kali sehari (pukul 08:00 dan 16:00 WIB) untuk memastikan bahan aktif tetap stabil.")

    add_heading_2(doc, "Formulir 07: Surat Jalan Pengiriman Produk Jadi (Delivery Order)")
    add_p(doc, "Dokumen resmi pengantar barang yang mencantumkan nama PT Distributor penerima, rincian nama produk, nomor izin edar PKD, nomor batch yang dikirim, jumlah kemasan, dan tanda terima ekspedisi.")

    add_heading_2(doc, "Formulir 08: Formulir Penanganan Keluhan Pelanggan")
    add_p(doc, "Formulir pencatatan laporan ketidaksesuaian produk dari pasar: identitas pelapor, nomor batch produk yang dikeluhkan, investigasi sampel arsip oleh PJT, kesimpulan akar masalah, dan tindakan korektif.")

    # -------------------------------------------------------------------------
    # 6. PANDUAN SISTEM PENGARSIPAN 3 BINDER BESAR
    # -------------------------------------------------------------------------
    add_heading_1(doc, "6. PANDUAN PENGARSIPAN FISIK (SISTEM 3 BINDER MASTER)")
    add_p(doc, "Seluruh dokumen fisik di pabrik wajib diklasifikasikan dan disimpan rapi ke dalam 3 Ordner / Ring Binder besar:")
    
    add_bullet(doc, "Berisi Akta Pendirian PT Perorangan, SK Kemenkumham, NPWP, NIB KBLI 20231, SPPL Lingkungan, Berkas PJT (Ijazah/STRTTK), Sertifikat Standar Produksi PKRT, Sertifikat Izin Edar PKD, Denah Tata Ruang, dan Bagan Struktur Organisasi.", "Binder 1 (Legalitas & Manual Mutu): ");
    add_bullet(doc, "Berisi Master Dokumen ke-20 SOP CPKRTB yang telah ditandatangani pengesahannya oleh Direktur dan PJT, Master Formula Baku, serta Master Kontrak Makloon.", "Binder 2 (Kumpulan Master SOP & Kontrak): ");
    add_bullet(doc, "Berisi seluruh lembar kerja harian yang sudah terisi dan berjalan: Catatan Pengolahan Bets (Batch Record), Rekaman Hasil Uji QC, Salinan CoA yang diterbitkan, Logbook Sanitasi, dan Salinan Surat Jalan Pengiriman.", "Binder 3 (Rekaman Operasional & Batch Record): ");

    add_signature_block(doc, "Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", "Disahkan Oleh,\nDirektur Utama PT [Nama Perusahaan]")
    
    file_path = os.path.join(OUTPUT_DIR, "12_MANUAL_SISTEM_DOKUMENTASI_MUTU_PABRIK_CPKRTB_LENGKAP.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

if __name__ == "__main__":
    build_complete_manual()

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from kca_doc_humanizer import sanitize_docx_metadata

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_presentation_document():
    base_dir = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical"
    output_dir = os.path.join(base_dir, "Keuangan/Proposal")
    charts_dir = os.path.join(output_dir, "charts")
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, "PROPOSAL_EKSEKUTIF_INVESTASI_DISTRIBUTOR_ROI_500PCT_KCA.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # ---------------- KOP SURAT RESMI ----------------
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k1 = kop.add_run("PT KEDIRI CHEMICAL ABADI\n")
    k1.font.name = "Arial"
    k1.font.size = Pt(16)
    k1.font.bold = True
    k1.font.color.rgb = RGBColor(0, 0, 0)

    k2 = kop.add_run("MANUFAKTUR & DISTRIBUSI BAHAN KIMIA PEMBERSIH, LAUNDRY & MAKLOON INDUSTRI\n")
    k2.font.name = "Arial"
    k2.font.size = Pt(9.5)
    k2.font.bold = True
    k2.font.color.rgb = RGBColor(0, 0, 0)

    k3 = kop.add_run("Pabrik: RT.1/RW.6, Pagung, Kec. Semen, Kabupaten Kediri, Jawa Timur 64161\n"
                    "Hotline/WhatsApp: 0822-4400-6699 | Email: kdrchemicals@gmail.com\n")
    k3.font.name = "Arial"
    k3.font.size = Pt(9)
    k3.font.color.rgb = RGBColor(0, 0, 0)

    p_line = doc.add_paragraph("═" * 58)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_line.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- JUDUL PRESENTASI ----------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title_p.add_run("EXECUTIVE INVESTMENT PITCH & FINANCIAL SIMULATION\n")
    t1.font.name = "Arial"
    t1.font.size = Pt(13.5)
    t1.font.bold = True
    t1.font.color.rgb = RGBColor(0, 0, 0)

    t2 = title_p.add_run("SKEMA PEMBIAYAAN MESIN PRODUKSI DEDIKASI & PENGEMBALIAN INVESTASI 500% ROI\n"
                        "UNTUK MITRA DISTRIBUTOR UTAMA & PEMODAL STRATEGIS\n")
    t2.font.name = "Arial"
    t2.font.size = Pt(11)
    t2.font.bold = True
    t2.font.color.rgb = RGBColor(0, 0, 0)

    t3 = title_p.add_run("Nomor Dokumen: 002/PCH-KCA/ROI-500/IX/2026 | Alokasi Modal: Rp 200.000.000,-\n")
    t3.font.name = "Arial"
    t3.font.size = Pt(9.5)
    t3.font.italic = True
    t3.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- RINGKASAN EKSEKUTIF ----------------
    h_exec = doc.add_heading("RINGKASAN EKSEKUTIF (EXECUTIVE SUMMARY)", level=2)
    for r in h_exec.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_exec = doc.add_paragraph(
        "Proposal ini memaparkan model kemitraan strategis dari sudut pandang (Point of View / POV) Distributor dan Pemodal. "
        "Melalui skema Dedicated CapEx Offset Financing senilai Rp 200.000.000,- (Dua Ratus Juta Rupiah), Distributor tidak hanya "
        "memperoleh pengembalian modal pokok 100% secara otomatis, namun juga melipatgandakan keuntungan dagang dengan margin "
        "hingga 66,67% per unit, menghasilkan total akumulasi profit tunai bersih sebesar Rp 1.000.000.000,- (Satu Miliar Rupiah) "
        "atau setara Return on Investment (ROI) sebesar 500%."
    )
    p_exec.paragraph_format.line_spacing = 1.15
    for r in p_exec.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 1 ----------------
    h1 = doc.add_heading("1. ANATOMI STRUKTUR MARGIN & KEUNTUNGAN PER UNIT", level=2)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p1 = doc.add_paragraph(
        "Pada skema kemitraan mesin dedikasi KCA, setiap botol/jerigen yang dipesan mendapatkan potongan langsung 10% pada invoice, "
        "yang secara instan memangkas Harga Pokok Pembelian (COGS) dan melipatgandakan margin riil distributor:"
    )
    p1.paragraph_format.line_spacing = 1.15
    for r in p1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # VISUAL GRAFIK 2: PERBANDINGAN MARGIN
    chart2_path = os.path.join(charts_dir, "chart2_margin_comparison.png")
    if os.path.exists(chart2_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(6)
        doc.add_picture(chart2_path, width=Inches(6.4))
        p_cap = doc.add_paragraph("Gambar 1: Perbandingan Harga Beli, Laba Kotor per Unit, dan Persentase Margin Riil")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        for r in p_cap.runs:
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    t_margin = doc.add_table(rows=1, cols=4)
    t_margin.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_margin.rows[0].cells
    m_hdr[0].text = "Komponen Perhitungan"
    m_hdr[1].text = "Distributor Konvensional"
    m_hdr[2].text = "Mitra Mesin Dedikasi KCA"
    m_hdr[3].text = "Selisih Keuntungan Tambahan"
    for c in m_hdr:
        set_cell_shading(c, "1E293B")
        set_cell_margins(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    margin_items = [
        ("Harga Beli dari Pabrik", "Rp 10.000 / unit", "Rp 9.000 / unit (Diskon 10%)", "Hemat Rp 1.000 / unit (10%)"),
        ("Harga Jual ke Pasar (Laundry/Hotel)", "Rp 15.000 / unit", "Rp 15.000 / unit", "Standar Pasar Kompetitif"),
        ("Laba Kotor per Unit", "Rp 5.000 / unit", "Rp 6.000 / unit", "Ekstra Profit +Rp 1.000 / unit"),
        ("Persentase Margin Keuntungan", "50,00%", "66,67%", "Lonjakan Margin +16,67% Murni"),
        ("Pengembalian Modal Pokok", "Tidak Ada", "Rp 1.000 / unit masuk kas modal", "Modal Rp 200 Jt Balik Otomatis")
    ]
    for komp, konv, kca, sel in margin_items:
        row = t_margin.add_row().cells
        row[0].text = komp
        row[1].text = konv
        row[2].text = kca
        row[3].text = sel
        for c in row:
            set_cell_margins(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)
        if "Lonjakan" in sel:
            set_cell_shading(row[0], "F1F5F9")
            set_cell_shading(row[1], "F1F5F9")
            set_cell_shading(row[2], "F1F5F9")
            set_cell_shading(row[3], "F1F5F9")
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 2 ----------------
    doc.add_page_break()
    h2 = doc.add_heading("2. DIAGRAM VISUAL SIKLUS PERPUTARAN KAS (CASH FLOW CYCLE)", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_diag = doc.add_paragraph(
        "Siklus perputaran dana dan arus kas distributor berlangsung dalam 4 tahapan tertutup dan terukur seperti diilustrasikan pada diagram berikut:"
    )
    p_diag.paragraph_format.line_spacing = 1.15
    for r in p_diag.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # VISUAL GRAFIK 1: FLOWCHART
    chart1_path = os.path.join(charts_dir, "chart1_flowchart_cashflow.png")
    if os.path.exists(chart1_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(6)
        doc.add_picture(chart1_path, width=Inches(6.4))
        p_cap = doc.add_paragraph("Gambar 2: Diagram Alur Perputaran Kas & Akumulasi ROI 500% Kemitraan KCA")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        for r in p_cap.runs:
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 3 ----------------
    h3 = doc.add_heading("3. SIMULASI PERTUMBUHAN LABA & SKENARIO PENJUALAN", level=2)
    for r in h3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p3 = doc.add_paragraph(
        "Grafik di bawah ini menggambarkan penurunan sisa saldo modal mesin dari Rp 200 Juta menuju Rp 0 (lunas), "
        "seiring dengan peningkatan akumulasi laba bersih tunai distributor hingga mencapai Rp 1.000.000.000,- (Satu Miliar Rupiah):"
    )
    p3.paragraph_format.line_spacing = 1.15
    for r in p3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # VISUAL GRAFIK 3: PENGURANGAN SALDO & AKUMULASI LABA
    chart3_path = os.path.join(charts_dir, "chart3_capital_recovery_growth.png")
    if os.path.exists(chart3_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(6)
        doc.add_picture(chart3_path, width=Inches(6.4))
        p_cap = doc.add_paragraph("Gambar 3: Grafik Amortisasi Modal Mesin vs Akumulasi Laba Bersih Tunai Distributor")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        for r in p_cap.runs:
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    # VISUAL GRAFIK 4: 3 SKENARIO TIMELINE
    chart4_path = os.path.join(charts_dir, "chart4_scenario_timeline.png")
    if os.path.exists(chart4_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(6)
        doc.add_picture(chart4_path, width=Inches(6.4))
        p_cap = doc.add_paragraph("Gambar 4: Proyeksi Arus Kas Masuk & Laba Bersih Bulanan pada 3 Skenario Distribusi")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        for r in p_cap.runs:
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    t_scen = doc.add_table(rows=1, cols=5)
    t_scen.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_hdr = t_scen.rows[0].cells
    s_hdr[0].text = "Parameter Kinerja"
    s_hdr[1].text = "Skenario Konservatif\n(12 Bulan)"
    s_hdr[2].text = "Skenario Moderat\n(8 Bulan)"
    s_hdr[3].text = "Skenario Agresif\n(5 Bulan)"
    s_hdr[4].text = "Keterangan Finansial"
    for c in s_hdr:
        set_cell_shading(c, "1E293B")
        set_cell_margins(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    scen_data = [
        ("Volume Penjualan / Bulan", "16.667 unit / bln", "25.000 unit / bln", "40.000 unit / bln", "Rata-rata penjualan harian"),
        ("Penjualan Harian (25 Hari)", "667 unit / hari", "1.000 unit / hari", "1.600 unit / hari", "Area Jawa Timur & sekitarnya"),
        ("Kas Masuk Pasar / Bulan", "Rp 250.000.000", "Rp 375.000.000", "Rp 600.000.000", "Omzet bruto distributor"),
        ("Belanja ke Pabrik / Bulan", "Rp 150.000.000", "Rp 225.000.000", "Rp 360.000.000", "90% kas dibayar ke KCA"),
        ("Pengembalian Modal Mesin/Bln", "Rp 16.667.000", "Rp 25.000.000", "Rp 40.000.000", "10% amortisasi modal"),
        ("Laba Bersih Tunai / Bulan", "Rp 83.333.000", "Rp 125.000.000", "Rp 200.000.000", "Profit bersih bulanan"),
        ("WAKTU LUNAS MODAL 200 JT", "Bulan ke-12", "Bulan ke-8", "Bulan ke-5", "100% Modal Pokok Kembali"),
        ("TOTAL PROFIT AKUMULASI", "Rp 1.000.000.000", "Rp 1.000.000.000", "Rp 1.000.000.000", "500% ROI Bersih Murni")
    ]
    for p_name, sk1, sk2, sk3, ket in scen_data:
        row = t_scen.add_row().cells
        row[0].text = p_name
        row[1].text = sk1
        row[2].text = sk2
        row[3].text = sk3
        row[4].text = ket
        for c in row:
            set_cell_margins(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)
        if "TOTAL PROFIT" in p_name or "WAKTU LUNAS" in p_name:
            set_cell_shading(row[0], "F1F5F9")
            set_cell_shading(row[1], "F1F5F9")
            set_cell_shading(row[2], "F1F5F9")
            set_cell_shading(row[3], "F1F5F9")
            set_cell_shading(row[4], "F1F5F9")
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 4: SPESIFIKASI ALOKASI MESIN & PRODUK ----------------
    h4 = doc.add_heading("4. SPESIFIKASI ALOKASI MESIN & PORTOFOLIO PRODUK UNGGULAN", level=2)
    for r in h4.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p4 = doc.add_paragraph(
        "Alokasi dana investasi sebesar Rp 200.000.000,- dialokasikan 100% secara transparan untuk pengadaan unit mesin dedikasi:\n"
        "1. Mesin Mixer Homogenizer SUS316 Double Jacket (1.000 - 2.000 L/Batch) : Rp 115.000.000,-\n"
        "2. Semi-Automatic Liquid Filling Machine 4-Nozzle (600 - 800 Botol/Jam) : Rp 45.000.000,-\n"
        "3. Mesin Continuous Induction Sealer & Capping Otomatis : Rp 25.000.000,-\n"
        "4. Sistem Pompa Kimia, Filter Presisi & Sanitary Piping : Rp 15.000.000,-\n\n"
        "Portofolio produk prioritas meliputi: Liquid Detergent Matic, Liquid DishWasher, Pencerah Warna Oxygen Bleach, "
        "Emulsifier Pengangkat Lemak, Bleach Klorin, Rust Tex (Karat), Blood Tex (Darah), dan Peluruh Noda Minyak Spa (10 kg Pail)."
    )
    p4.paragraph_format.line_spacing = 1.15
    for r in p4.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- LEMBAR PENGESAHAN ----------------
    doc.add_paragraph("\n")
    h5 = doc.add_heading("5. LEMBAR KESEPAKATAN & KOMITMEN KEMITRAAN", level=2)
    for r in h5.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_penutup = doc.add_paragraph(
        "Melalui penandatanganan lembar komitmen ini, kedua belah pihak sepakat untuk memulai persiapan teknis "
        "pengadaan mesin dedikasi dan alur pemesanan produk sesuai ketentuan yang tercantum dalam proposal ini."
    )
    p_penutup.paragraph_format.line_spacing = 1.15
    for r in p_penutup.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("\n")
    t_sign = doc.add_table(rows=3, cols=2)
    t_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t_sign.rows[0].cells[0].text = "PENGELOLA PABRIK & MANUFAKTUR\nPT KEDIRI CHEMICAL ABADI"
    t_sign.rows[0].cells[1].text = "MITRA PEMODAL & DISTRIBUTOR UTAMA\n"
    
    t_sign.rows[1].cells[0].text = "\n\n\n\n"
    t_sign.rows[1].cells[1].text = "\n\n\n\n"

    t_sign.rows[2].cells[0].text = "YAN EFFENDI / YERIKHO ARFENSIAS E.\nDirektur Utama / General Manager"
    t_sign.rows[2].cells[1].text = "(......................................................)\nNama Lengkap & Jabatan"

    for r in t_sign.rows:
        for c in r.cells:
            set_cell_margins(c)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_run in p.runs:
                    r_run.font.bold = True
                    r_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(doc_path)
    
    # Salin ke folder arsip legalitas
    legal_dir = os.path.join(base_dir, "KCA DOKUMEN/06_MASTER_MANUAL_DAN_LEGALITAS")
    os.makedirs(legal_dir, exist_ok=True)
    legal_path = os.path.join(legal_dir, "PROPOSAL_EKSEKUTIF_INVESTASI_DISTRIBUTOR_ROI_500PCT_KCA.docx")
    doc.save(legal_path)
    
    # Sanitasi Metadata
    sanitize_docx_metadata(doc_path, title="Proposal Eksekutif Investasi Distributor ROI 500% KCA", author="Yerikho Arfensias Effendi")
    sanitize_docx_metadata(legal_path, title="Proposal Eksekutif Investasi Distributor ROI 500% KCA", author="Yerikho Arfensias Effendi")
    
    print(f"SUCCESS: Clean Focused Proposal generated at {doc_path} and {legal_path}")
    return doc_path

if __name__ == "__main__":
    create_presentation_document()

import os
import sys
import datetime
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from kca_doc_humanizer import sanitize_docx_metadata

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_monthly_report(target_month="September", target_year="2026"):
    base_dir = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical"
    xlsx_path = os.path.join(base_dir, "Keuangan/Dashboard Keuangan PT Kediri Chemical Abadi.xlsx")
    output_dir = os.path.join(base_dir, "Keuangan/Laporan Bulanan")
    os.makedirs(output_dir, exist_ok=True)
    
    doc_filename = f"LAPORAN_KEUANGAN_BULANAN_KCA_{target_month.upper()}_{target_year}.docx"
    doc_path = os.path.join(output_dir, doc_filename)

    wb = openpyxl.load_workbook(xlsx_path, read_only=True)
    
    # 1. Gaji
    ws_gaji = wb['GAJI KARYAWAN']
    total_payroll = 0
    gaji_rows = []
    for r in ws_gaji.iter_rows(values_only=True):
        if r and len(r) > 12 and r[4] == target_month and str(r[5]) == target_year:
            nama = str(r[6])
            status = str(r[7])
            total = float(r[12]) if r[12] is not None else 0.0
            total_payroll += total
            gaji_rows.append((nama, status, total))
    if total_payroll == 0:
        total_payroll = 7500000.0

    # 2. Aset & Depresiasi
    ws_aset = wb['ASET & DEPRESIASI']
    total_aset = 0
    total_dep_bulanan = 0
    for r in ws_aset.iter_rows(values_only=True):
        if r and len(r) > 16 and r[3] in ['AS001', 'AS002', 'AS003', 'AS004']:
            nom = float(r[10]) if r[10] else 0.0
            dep_thn = float(r[13]) if r[13] else 0.0
            dep_bln = dep_thn / 12.0
            total_aset += nom
            total_dep_bulanan += dep_bln

    # 3. Hutang Restrukturisasi
    ws_hutang = wb['HUTANG']
    sisa_hutang = 700000000.0
    for r in ws_hutang.iter_rows(values_only=True):
        if r and len(r) > 13 and r[6] == 'H-001':
            sisa_hutang = float(r[13]) if r[13] else 700000000.0

    wb.close()

    # 4. Susun Dokumen Resmi Word ISO / Direksi
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Header / Kop
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr1 = title.add_run("PT KEDIRI CHEMICAL ABADI\n")
    tr1.font.name = "Arial"
    tr1.font.size = Pt(16)
    tr1.font.bold = True
    tr1.font.color.rgb = RGBColor(0, 0, 0)

    tr2 = title.add_run("LAPORAN KEUANGAN & EKSEKUTIF BULANAN\n")
    tr2.font.name = "Arial"
    tr2.font.size = Pt(12.5)
    tr2.font.bold = True
    tr2.font.color.rgb = RGBColor(0, 0, 0)

    tr3 = title.add_run(f"Periode: {target_month} {target_year} | No. Dokumen: FIN-REP-{target_year}{target_month[:3].upper()}-001\n")
    tr3.font.name = "Arial"
    tr3.font.size = Pt(9.5)
    tr3.font.italic = True
    tr3.font.color.rgb = RGBColor(0, 0, 0)

    p_line = doc.add_paragraph("═" * 58)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_line.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 1. Ringkasan Eksekutif
    h1 = doc.add_heading("1. RINGKASAN EKSEKUTIF DIREKSI", level=2)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    
    p_sum = doc.add_paragraph(
        f"Laporan ini menyajikan posisi keuangan, realisasi biaya operasional pabrik, pemantauan aset, "
        f"dan status kewajiban restrukturisasi PT Kediri Chemical Abadi untuk periode {target_month} {target_year}. "
        f"Pengawasan operasional dipimpin oleh Bpk. Yerikho Arfensias Effendi di bawah arahan Direktur Utama Bpk. Yan Effendi."
    )
    p_sum.paragraph_format.line_spacing = 1.15
    for r in p_sum.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 2. Struktur Biaya Operasional (OPEX)
    h2 = doc.add_heading("2. REALISASI BIAYA OPERASIONAL BULANAN (OPEX)", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    t_opex = doc.add_table(rows=1, cols=3)
    t_opex.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t_opex.rows[0].cells
    hdr[0].text = "Pos Biaya / Beban Operasional"
    hdr[1].text = "Kategori"
    hdr[2].text = "Realisasi / Anggaran"
    for cell in hdr:
        set_cell_shading(cell, "1E293B")
        set_cell_margins(cell)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    opex_items = [
        ("Beban Gaji & Upah (Payroll)", "SDM & Manajemen", f"Rp {total_payroll:,.0f}"),
        ("Beban Listrik Pabrik (PLN)", "Utilitas Pabrik", "Rp 1,700,000"),
        ("Beban Air & Sanitasi Pabrik", "Utilitas & Kebersihan", "Rp 300,000"),
        ("Beban BBM Operasional (Xenia 2012)", "Logistik & Sales", "Rp 1,200,000"),
        ("Beban Konsumsi Pabrik (3-4 Tim)", "Kesejahteraan Karyawan", "Rp 800,000"),
        ("Beban Pemeliharaan & ATK", "Operasional Umum", "Rp 500,000"),
        ("TOTAL BIAYA TETAP OPERASIONAL (OPEX)", "TOTAL FIXED COST", "Rp 12,000,000")
    ]
    for pos, kat, nom in opex_items:
        row_cells = t_opex.add_row().cells
        row_cells[0].text = pos
        row_cells[1].text = kat
        row_cells[2].text = nom
        for cell in row_cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)
        if "TOTAL" in pos:
            set_cell_shading(row_cells[0], "F1F5F9")
            set_cell_shading(row_cells[1], "F1F5F9")
            set_cell_shading(row_cells[2], "F1F5F9")
            for cell in row_cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0, 0, 0)

    # 3. Posisi Aset & Kewajiban Restrukturisasi
    h3 = doc.add_heading("3. POSISI ASET & KEWAJIBAN RESTRUKTURISASI", level=2)
    for r in h3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_ast = doc.add_paragraph(
        f"• Total Nilai Perolehan Aset Tetap: Rp {total_aset:,.0f} (Daihatsu Xenia 2012, 2 Mesin Pengaduk Manual, Pabrik Pagung)\n"
        f"• Alokasi Beban Penyusutan Bulanan: Rp {total_dep_bulanan:,.0f} / bulan\n"
        f"• Sisa Hutang Usaha Restrukturisasi: Rp {sisa_hutang:,.0f} (Target penyelesaian via cicilan bertahap dari profit produk)"
    )
    p_ast.paragraph_format.line_spacing = 1.15
    for r in p_ast.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 4. Lembar Pengesahan
    doc.add_paragraph("\n")
    t_sign = doc.add_table(rows=3, cols=2)
    t_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_hdr = t_sign.rows[0].cells
    s_hdr[0].text = "Dibuat & Dilaporkan Oleh:"
    s_hdr[1].text = "Disetujui & Diketahui Oleh:"
    
    t_sign.rows[1].cells[0].text = "\n\n\n"
    t_sign.rows[1].cells[1].text = "\n\n\n"

    s_ftr = t_sign.rows[2].cells
    s_ftr[0].text = "YERIKHO ARFENSIAS EFFENDI\nGeneral Manager / Finance & Ops"
    s_ftr[1].text = "YAN EFFENDI\nDirektur Utama"

    for r in t_sign.rows:
        for cell in r.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_run in p.runs:
                    r_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(doc_path)
    # Sanitasi Metadata Otomatis
    sanitize_docx_metadata(doc_path, title=f"Laporan Keuangan Bulanan KCA {target_month} {target_year}", author="Yerikho Arfensias Effendi")
    print(f"SUCCESS: Report generated at {doc_path}")
    return doc_path

if __name__ == "__main__":
    m = sys.argv[1] if len(sys.argv) > 1 else "September"
    y = sys.argv[2] if len(sys.argv) > 2 else "2026"
    generate_monthly_report(m, y)

import os
import sys
import shutil
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

BASE_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/KCA DOKUMEN"
os.makedirs(BASE_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PT KCA CHEMICAL\nMANUFAKTUR PKRT\nSISTEM ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_signature_block_3(doc, pjt_name="( ............................................ )", dir_name="( ............................................ )", kabag_name="( ............................................ )"):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    col_widths = [Cm(5.0), Cm(5.0), Cm(5.5)]
    
    titles = [
        ("Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", pjt_name),
        ("Diperiksa Oleh,\nKepala Bagian Operasional", kabag_name),
        ("Disahkan Oleh,\nDirektur Utama", dir_name)
    ]
    
    for idx, (title, name) in enumerate(titles):
        cell = table.cell(0, idx)
        cell.width = col_widths[idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + "\n\n\n\n\n" + name + "\nTanggal: ...................................")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)
        r.font.color.rgb = COLOR_BLACK

def create_individual_sop(folder_name, filename, doc_no, title, tujuan, ruang_lingkup, tanggung_jawab, definisi, prosedur_list, formulir_terkait):
    folder_path = os.path.join(BASE_DIR, folder_name)
    os.makedirs(folder_path, exist_ok=True)
    
    doc = setup_iso_document()
    add_iso_header_box(doc, f"Standar Operasional Prosedur:\n{title}", doc_no, "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. TUJUAN")
    add_p(doc, tujuan)
    
    add_heading_1(doc, "2. RUANG LINGKUP")
    add_p(doc, ruang_lingkup)
    
    add_heading_1(doc, "3. TANGGUNG JAWAB DAN WEWENANG")
    for tj in tanggung_jawab:
        add_bullet(doc, tj[1], tj[0])
        
    add_heading_1(doc, "4. DEFINISI DAN DASAR HUKUM ACUAN")
    add_p(doc, definisi)
    
    add_heading_1(doc, "5. PROSEDUR KERJA LANGKAH DEMI LANGKAH")
    for step in prosedur_list:
        add_bullet(doc, step[1], step[0])
        
    add_heading_1(doc, "6. DOKUMEN DAN FORMULIR TERKAIT")
    for form in formulir_terkait:
        add_bullet(doc, form[1], form[0])
        
    add_signature_block_3(doc)
    
    file_path = os.path.join(folder_path, filename)
    doc.save(file_path)
    print(f"Created: {file_path}")

# ==============================================================================
# DATA DEFINITIONS FOR ALL 20 SOPS
# ==============================================================================
sops_data = [
    # -------------------------------------------------------------------------
    # FOLDER 1: 01_BANGUNAN_FASILITAS_DAN_SANITASI
    # -------------------------------------------------------------------------
    (
        "01_BANGUNAN_FASILITAS_DAN_SANITASI",
        "SOP-SAN-001_Pembersihan_dan_Sanitasi_Ruangan_Pabrik.docx",
        "SOP-SAN-001",
        "Pembersihan dan Sanitasi Ruangan Pabrik",
        "Memastikan seluruh area produksi, gudang, dan koridor pabrik berada dalam kondisi bersih, higienis, bebas dari debu, kotoran, dan sisa bahan kimia agar tidak terjadi kontaminasi silang.",
        "Prosedur ini berlaku untuk seluruh ruang pengolahan/mixing, ruang penimbangan, ruang filling, ruang kemas karton, dan gudang di PT KCA Chemical.",
        [
            ("Petugas Sanitasi: ", "Melaksanakan pembersihan rutin harian dan mingguan sesuai jadwal."),
            ("Kepala Produksi: ", "Memeriksa dan memverifikasi kebersihan ruangan sebelum produksi dimulai."),
            ("PJT: ", "Memantau efektivitas disinfektan dan kepatuhan standar CPKRTB.")
        ],
        "Dasar hukum acuan: Permenkes No. 14 Tahun 2021 tentang Standar Usaha dan Produk Sektor Kesehatan (CPKRTB) dan ISO 9001:2015 Klausul 7.1.3 (Infrastruktur) & 7.1.4 (Lingkungan untuk Operasi Proses).",
        [
            ("1. Persiapan: ", "Siapkan peralatan kebersihan khusus area produksi (kain pel mikrofiber, ember ganda, sapu nilon, cairan disinfektan Benzalkonium Klorida 0.1% atau Karbol wangi)."),
            ("2. Pembersihan Harian (Pagi & Sore): ", "Sapu lantai dari sudut terdalam menuju pintu keluar. Lakukan pengepelan lantai dengan larutan disinfektan. Bersihkan meja timbang dan meja filling dengan lap basah berdisinfektan."),
            ("3. Pembersihan Mingguan: ", "Lap seluruh dinding keramik, daun pintu, kaca jendela, dan bersihkan penutup lampu serta kisi-kisi exhaust fan dari debu."),
            ("4. Pembuangan Sampah: ", "Keluarkan kantong sampah padat dari tempat sampah berpedal setiap sore hari dan buang ke TPS luar pabrik."),
            ("5. Pencatatan: ", "Petugas mencatat kegiatan pada Logbook Pembersihan Ruangan dan meminta paraf verifikasi Kepala Produksi.")
        ],
        [
            ("FRM-SAN-01: ", "Logbook Pembersihan dan Sanitasi Ruangan Harian."),
            ("FRM-SAN-02: ", "Checklist Pemeriksaan Kebersihan Mingguan.")
        ]
    ),
    (
        "01_BANGUNAN_FASILITAS_DAN_SANITASI",
        "SOP-EQP-002_Pembersihan_dan_Perawatan_Mesin_Tangki_Mixer.docx",
        "SOP-EQP-002",
        "Pembersihan dan Perawatan Mesin / Tangki Mixer",
        "Mencegah terjadinya percampuran sisa formula produk dari batch sebelumnya dan menjaga keawetan tangki mixer serta mesin filling.",
        "Berlaku untuk tangki mixer sabun/pembersih (Stainless SUS 304 / HDPE), mesin filling pneumatik, dan pompa transfer cairan.",
        [
            ("Operator Produksi: ", "Melakukan pembersihan fisik dan pembilasan tangki mixer segera setelah proses selesai."),
            ("Teknisi Mesin: ", "Melakukan perawatan preventif berkala (pelumasan motor mixer, cek seal pompa)."),
            ("PJT & QC: ", "Memeriksa air bilasan terakhir dan mengesahkan label status tangki.")
        ],
        "CPKRTB Aspek Peralatan: Peralatan yang digunakan dalam proses produksi harus dirancang, ditempatkan, dan dirawat dengan baik agar mudah dibersihkan dan tidak mencemari produk.",
        [
            ("1. Pengurasan: ", "Kuras habis sisa cairan produk pada dasar tangki mixer."),
            ("2. Pembilasan Awal: ", "Semprotkan air bersih bertekanan ke seluruh dinding dalam tangki mixer untuk melarutkan sisa surfaktan/busa."),
            ("3. Penyikatan: ", "Gunakan sikat bertangkai halus dan deterjen netral untuk membersihkan kerak formula yang menempel pada baling-baling pengaduk."),
            ("4. Pembilasan Akhir: ", "Bilas dengan air bersih hingga air buangan jernih, tidak berbusa, berbau netral, dan pH air bilasan sama dengan pH air baku (pH 6.5 - 7.5)."),
            ("5. Penempelan Label: ", "Pasang label gantung bertuliskan 'BERSIH - SIAP PAKAI' lengkap dengan tanggal pembersihan dan paraf operator.")
        ],
        [
            ("FRM-EQP-01: ", "Logbook Pembersihan Tangki Mixer dan Peralatan."),
            ("FRM-EQP-02: ", "Kartu Riwayat Pemeliharaan Mesin (Maintenance Card).")
        ]
    ),
    (
        "01_BANGUNAN_FASILITAS_DAN_SANITASI",
        "SOP-HIG-003_Higiene_Personalia_dan_Pemakaian_APD.docx",
        "SOP-HIG-003",
        "Higiene Personalia dan Penggunaan Alat Pelindung Diri (APD)",
        "Menjamin perlindungan keselamatan dan kesehatan kerja (K3) personil dari uap/kontak kimia serta mencegah kontaminasi dari personil ke produk.",
        "Berlaku untuk seluruh karyawan, teknisi, analis lab, dan tamu yang memasuki area pengolahan dan pengemasan.",
        [
            ("Seluruh Karyawan: ", "Wajib mematuhi protokol higiene dan memakai APD yang disyaratkan."),
            ("Kepala Bagian Operasional: ", "Mengawasi kepatuhan pemakaian APD di area kerja harian."),
            ("PJT: ", "Menetapkan standar APD yang sesuai dengan tingkat risiko bahan kimia.")
        ],
        "Dasar hukum acuan: Permenkes 14/2021 dan UU No. 1 Tahun 1970 tentang Keselamatan Kerja.",
        [
            ("1. Memasuki Ruang Produksi: ", "Tanggalkan sepatu luar di ruang loker, kenakan pakaian kerja pabrik/jas lab dan sepatu boots tertutup."),
            ("2. Cuci Tangan: ", "Cuci tangan dengan sabun antiseptik pada wastafel alir selama minimal 20 detik, keringkan dengan tisu bersih."),
            ("3. Pemakaian APD: ", "Kenakan masker medis/respirator uap kimia, tutup kepala (hairnet), kacamata pelindung (goggle) untuk cairan asam/klorin, dan sarung tangan nitril/karet."),
            ("4. Larangan di Area Produksi: ", "Dilarang memakai cincin, gelang, jam tangan, dilarang makan, minum, meludah, atau merokok."),
            ("5. Kondisi Kesehatan: ", "Personel yang menderita flu, luka terbuka, atau sakit kulit dilarang berada di ruang penimbangan dan mixing langsung.")
        ],
        [
            ("FRM-HIG-01: ", "Checklist Harian Higiene dan Kerapian APD Karyawan."),
            ("FRM-HIG-02: ", "Rekaman Pemeriksaan Kesehatan Personel.")
        ]
    ),
    (
        "01_BANGUNAN_FASILITAS_DAN_SANITASI",
        "SOP-PST-004_Pengendalian_Hama_Pest_Control.docx",
        "SOP-PST-004",
        "Pengendalian Hama dan Binatang Pengganggu (Pest Control)",
        "Mencegah masuk, berkembang biak, dan mencemarnya serangga (lalat, nyamuk, kecoa) serta hewan pengerat (tikus) di lingkungan pabrik.",
        "Meliputi area luar gedung pabrik, gudang bahan baku, ruang produksi, dan gudang barang jadi.",
        [
            ("Petugas Gudang / Sanitasi: ", "Melakukan pengecekan dan penggantian umpan/perangkap hama."),
            ("Kepala Operasional: ", "Mengevaluasi laporan temuan hama mingguan."),
            ("PJT: ", "Memastikan tidak ada racun kimia berbahaya yang berpotensi mencemari formula produk.")
        ],
        "Standar Sanitasi Industri CPKRTB Kemenkes: Sarana produksi harus memiliki sistem pencegahan dan pengendalian hama yang efektif.",
        [
            ("1. Pemasangan Perangkap Serangga: ", "Pasang Insect Fly Catcher dengan lampu UV di dekat pintu masuk gudang dan produksi (posisi tidak di atas tangki terbuka)."),
            ("2. Pemasangan Perangkap Tikus: ", "Letakkan glue trap atau box trap mekanis di sepanjang dinding luar dan sudut-sudut gudang."),
            ("3. Pengecekan Rutin: ", "Lakukan inspeksi seluruh titik perangkap setiap hari Senin pagi."),
            ("4. Tindakan Koreksi: ", "Jika ditemukan bangkai serangga/tikus, segera buang dan kubur di luar area, ganti lem perangkap baru, dan catat titik temuan."),
            ("5. Pengendalian Fisik: ", "Pastikan semua pintu selalu tertutup rapat dan kawat kasa ventilasi tidak ada yang robek.")
        ],
        [
            ("FRM-PST-01: ", "Denah Titik Perangkap Hama (Pest Trap Map)."),
            ("FRM-PST-02: ", "Logbook Pemantauan dan Temuan Hama Mingguan.")
        ]
    ),

    # -------------------------------------------------------------------------
    # FOLDER 2: 02_GUDANG_DAN_PENGADAAN_BAHAN
    # -------------------------------------------------------------------------
    (
        "02_GUDANG_DAN_PENGADAAN_BAHAN",
        "SOP-GUD-005_Penerimaan_dan_Pemeriksaan_Bahan_Masuk.docx",
        "SOP-GUD-005",
        "Penerimaan dan Pemeriksaan Bahan Baku / Kemas Masuk",
        "Menjamin bahwa seluruh bahan kimia, bahan kemasan, dan label yang diterima dari pemasok sesuai dengan pesanan dan spesifikasi mutu resmi.",
        "Berlaku untuk seluruh penerimaan bahan mentah cair, bubuk, botol, jerigen, tutup, dan karton boks di area bongkar muat gudang.",
        [
            ("Petugas Penerimaan Gudang: ", "Memeriksa surat jalan, keutuhan fisik kemasan drum, dan menempelkan label karantina."),
            ("Petugas QC: ", "Mengambil sampel bahan dan melakukan uji verifikasi mutu laboratorium."),
            ("PJT: ", "Menandatangani status pelepasan (release) bahan baku.")
        ],
        "ISO 9001:2015 Klausul 8.4 (Pengendalian Proses, Produk dan Jasa yang Disediakan Eksternal).",
        [
            ("1. Pemeriksaan Dokumen: ", "Cocokkan Surat Jalan ekspedisi dengan Surat Pesanan (Purchase Order) dan minta Certificate of Analysis (CoA) dari supplier."),
            ("2. Pemeriksaan Fisik: ", "Periksa keutuhan wadah (drum tidak bocor, tidak berkarat, segel utuh, label nama bahan jelas)."),
            ("3. Penempatan Karantina: ", "Turunkan barang ke area Karantina Bahan Masuk dan tempel stiker kuning 'DIKARANTINA - MENUNGGU UJI QC'."),
            ("4. Pengambilan Sampel QC: ", "Laporkan kepada bagian QC untuk dilakukan pengambilan sampel uji bahan sesuai SOP-QC-013."),
            ("5. Pelepasan ke Rak: ", "Setelah QC meluluskan, ganti stiker menjadi hijau 'DILULUSKAN' dan pindahkan ke rak/palet bahan baku."),
            ("6. Pencatatan: ", "Catat jumlah bahan masuk pada Kartu Stok Gudang.")
        ],
        [
            ("FRM-GUD-01: ", "Formulir Pemeriksaan Bahan Masuk (Incoming Material Form)."),
            ("FRM-GUD-02: ", "Kartu Stok Bahan Baku Gudang.")
        ]
    ),
    (
        "02_GUDANG_DAN_PENGADAAN_BAHAN",
        "SOP-GUD-006_Penyimpanan_Bahan_Kimia_Berbahaya_dan_Korosif.docx",
        "SOP-GUD-006",
        "Penyimpanan Bahan Kimia Berbahaya dan Korosif",
        "Menjamin keselamatan penyimpanan bahan kimia keras/korosif (Sodium Hypochlorite, LABSA, Asam Sitrat, Caustic Soda) agar tidak mencemari lingkungan dan aman dari risiko kebocoran.",
        "Berlaku untuk area penyimpanan bahan kimia B3 di gudang PT KCA Chemical.",
        [
            ("Petugas Gudang: ", "Menata dan memantau kondisi fisik drum bahan korosif."),
            ("PJT: ", "Menyediakan Safety Data Sheet (MSDS) dan menetapkan rambu bahaya."),
            ("Kepala Operasional: ", "Memastikan ketersediaan fasilitas tanggap darurat tumpahan (Spill Kit).")
        ],
        "Dasar hukum acuan: Permenaker No. 187/MEN/1999 tentang Pengendalian Bahan Kimia Berbahaya di Tempat Kerja.",
        [
            ("1. Penataan Zonasi: ", "Simpan bahan korosif di area berventilasi udara lancar, terpisah dari bahan parfum/pewarna organik."),
            ("2. Paletisasi Khusus: ", "Wadah drum/jerigen wajib diletakkan di atas palet plastik yang dilengkapi bak penampung tumpahan (spill pallet)."),
            ("3. Pemasangan Rambu & MSDS: ", "Tempel simbol bahaya korosif (GHS) dan lembar MSDS bahasa Indonesia pada dinding area penyimpanan."),
            ("4. Alat Tanggap Darurat: ", "Sediakan pasir kering/serbuk gergaji, sekop plastik, dan eye washer darurat dalam jarak jangkau maksimal 10 meter."),
            ("5. Pemantauan Kebocoran: ", "Lakukan inspeksi visual setiap pagi terhadap potensi rembesan atau pembengkakan wadah.")
        ],
        [
            ("FRM-GUD-03: ", "Checklist Inspeksi Area Penyimpanan Kimia Berbahaya."),
            ("FRM-GUD-04: ", "Logbook Penggunaan Bahan Korosif.")
        ]
    ),
    (
        "02_GUDANG_DAN_PENGADAAN_BAHAN",
        "SOP-GUD-007_Pengeluaran_Bahan_Baku_Metode_FIFO_FEFO.docx",
        "SOP-GUD-007",
        "Pengeluaran Bahan Baku Sistem FIFO dan FEFO",
        "Mencegah penggunaan bahan kimia yang kedaluwarsa atau rusak akibat penyimpanan terlalu lama melalui sistem FIFO (First In First Out) dan FEFO (First Expired First Out).",
        "Berlaku untuk seluruh bahan aktif kimia, bahan pembantu, parfum, pewarna, dan kemasan primer/sekunder.",
        [
            ("Petugas Gudang: ", "Mengeluarkan bahan sesuai nomor urut tanggal masuk dan tanggal expired."),
            ("Petugas Penimbangan: ", "Memverifikasi tanggal expired bahan sebelum ditimbang."),
            ("Kepala Gudang: ", "Mengawasi rotasi stok dan mutasi barang.")
        ],
        "ISO 9001:2015 Klausul 8.5.4 (Preservasi / Pemeliharaan Produk dan Bahan).",
        [
            ("1. Pemberian Identitas: ", "Tempel label identitas yang memuat Tanggal Masuk dan Tanggal Expired pada setiap drum saat diterima."),
            ("2. Penataan Rak: ", "Bahan yang masuk lebih awal atau memiliki tanggal kedaluwarsa lebih dekat diletakkan di posisi paling depan/mudah dijangkau."),
            ("3. Permintaan Produksi: ", "Terima Formulir Permintaan Bahan dari bagian produksi yang disetujui Kepala Produksi."),
            ("4. Pengambilan Bahan: ", "Ambil bahan dengan prinsip FEFO terlebih dahulu, kemudian FIFO."),
            ("5. Update Kartu Stok: ", "Kurangi jumlah saldo stok pada Kartu Stok Fisik dan sistem database inventaris."),
            ("6. Penyerahan: ", "Serahkan bahan ke ruang penimbangan disertai paraf serah terima.")
        ],
        [
            ("FRM-GUD-02: ", "Kartu Stok Bahan Baku & Kemasan."),
            ("FRM-GUD-05: ", "Formulir Permintaan dan Pengeluaran Bahan (Material Requisition Form).")
        ]
    ),
    (
        "02_GUDANG_DAN_PENGADAAN_BAHAN",
        "SOP-GUD-008_Penyimpanan_dan_Pengiriman_Produk_Jadi.docx",
        "SOP-GUD-008",
        "Penyimpanan dan Pengiriman Produk Jadi (Delivery)",
        "Menjamin produk jadi tersimpan dalam kondisi aman, tidak rusak, dan dikirimkan kepada distributor secara tepat jumlah, tepat nomor batch, dan lengkap dokumen mutunya.",
        "Berlaku untuk seluruh produk berizin PKD yang telah diluluskan QC di gudang barang jadi.",
        [
            ("Petugas Gudang Produk Jadi: ", "Menata karton produk jadi di atas palet dan menjaga kondisi gudang."),
            ("Petugas Ekspedisi: ", "Membuat Surat Jalan dan memeriksa kesesuaian muatan kendaraan."),
            ("PJT: ", "Menyediakan Certificate of Analysis (CoA) untuk dilampirkan pada pengiriman.")
        ],
        "ISO 9001:2015 Klausul 8.5.4 dan Permenkes 14/2021 tentang Distribusi Produk PKRT.",
        [
            ("1. Penerimaan dari Kemas: ", "Terima karton produk jadi dari ruang pengemasan sekunder lengkap dengan bukti kelulusan QC."),
            ("2. Penyimpanan Palet: ", "Susun karton di atas palet kayu/plastik dengan tumpukan maksimal 5 susun karton (jangan langsung menyentuh lantai)."),
            ("3. Penerimaan Order: ", "Terima Surat Perintah Pengiriman / Delivery Order dari bagian marketing/distributor."),
            ("4. Penyiapan Surat Jalan: ", "Buat Surat Jalan resmi yang mencantumkan nama distributor, nama produk, nomor PKD, nomor batch, dan jumlah koli/botol."),
            ("5. Pengecekan Muatan: ", "Periksa kebersihan bak truk pengangkut (kering, tidak bocor, tidak berbau tajam)."),
            ("6. Pelepasan Barang: ", "Muat barang ke truk, lampirkan lembar CoA asli dari PJT, dan minta tanda tangan penerima pada Surat Jalan.")
        ],
        [
            ("FRM-GUD-06: ", "Surat Jalan Pengiriman Barang Jadi (Delivery Order)."),
            ("FRM-GUD-07: ", "Kartu Stok Barang Jadi Gudang.")
        ]
    ),

    # -------------------------------------------------------------------------
    # FOLDER 3: 03_OPERASIONAL_PRODUKSI_DAN_PENGEMASAN
    # -------------------------------------------------------------------------
    (
        "03_OPERASIONAL_PRODUKSI_DAN_PENGEMASAN",
        "SOP-PRD-009_Penimbangan_Bahan_Baku_dan_Tara_Timbangan.docx",
        "SOP-PRD-009",
        "Penimbangan Bahan Baku dan Kalibrasi Tara Timbangan",
        "Menjamin keakuratan takaran setiap komponen bahan kimia sesuai dengan komposisi Master Formula yang didaftarkan ke Kementerian Kesehatan RI.",
        "Berlaku untuk penimbangan bahan cair dan padat di Ruang Penimbangan PT KCA Chemical.",
        [
            ("Petugas Penimbangan: ", "Melakukan penimbangan bahan secara presisi dan mencatat bobot aktual."),
            ("Petugas QC / PJT: ", "Memverifikasi kebenaran bobot timbang bahan aktif utama."),
            ("Kepala Produksi: ", "Memastikan seluruh bahan telah siap sebelum proses mixing dimulai.")
        ],
        "CPKRTB Aspek Penimbangan: Penimbangan bahan baku harus dilakukan oleh personil yang berwenang dengan menggunakan peralatan yang terkalibrasi.",
        [
            ("1. Cek Kebersihan: ", "Pastikan ruang timbang bersih dan meja timbangan bebas dari tumpahan debu bahan kimia."),
            ("2. Cek Timbangan: ", "Pastikan timbangan digital menunjukkan angka 0.00 (Tara) dan letakkan wadah penimbang bersih di atas platform timbang."),
            ("3. Penimbangan: ", "Timbang bahan baku satu per satu sesuai urutan pada Lembar Catatan Pengolahan Bets (Batch Record)."),
            ("4. Labelisasi Wadah: ", "Tempel label bertuliskan 'SUDAH DITIMBANG' memuat nama bahan, berat aktual, tanggal timbang, dan nomor batch target."),
            ("5. Verifikasi: ", "Petugas QC melakukan verifikasi ganda (double-check) terhadap bahan berisiko tinggi (klorin/surfaktan utama)."),
            ("6. Pencatatan: ", "Tuliskan hasil timbang aktual dan berikan paraf pada dokumen Batch Record.")
        ],
        [
            ("FRM-PRD-01: ", "Catatan Pengolahan Bets (Batch Processing Record) - Lembar Penimbangan."),
            ("FRM-PRD-02: ", "Label Identitas Bahan 'SUDAH DITIMBANG'.")
        ]
    ),
    (
        "03_OPERASIONAL_PRODUKSI_DAN_PENGEMASAN",
        "SOP-PRD-010_Proses_Pencampuran_Mixing_Sabun_dan_Pembersih.docx",
        "SOP-PRD-010",
        "Proses Pencampuran (Mixing) Sabun dan Kimia Pembersih",
        "Menghasilkan larutan sabun cuci piring, pembersih lantai, deterjen cair, dan pemutih yang homogen, stabil, tidak memisah, dan berpenampilan jernih.",
        "Berlaku untuk proses mixing di tangki Stainless Steel / HDPE di Ruang Produksi.",
        [
            ("Operator Mixing: ", "Mengoperasikan mesin pengaduk sesuai instruksi kerja kecepatan dan urutan bahan."),
            ("Kepala Produksi: ", "Mengawasi jalannya proses mixing dan durasi pengadukan."),
            ("Petugas QC: ", "Mengambil sampel in-process untuk verifikasi viskositas dan homogenitas.")
        ],
        "CPKRTB Aspek Pengolahan: Proses pengolahan harus mengikuti prosedur tervalidasi untuk mencegah kesalahan formulasi.",
        [
            ("1. Pemeriksaan Awal (Line Clearance): ", "Pastikan tangki mixer berlabel 'BERSIH - SIAP PAKAI' dan kran pembuangan tertutup rapat."),
            ("2. Pemasukan Air Baku: ", "Masukkan air demineral/air bersih terukur ke dalam tangki mixer (sekitar 80% dari total volume)."),
            ("3. Pelarutan Surfaktan: ", "Masukkan SLES / LABSA perlahan sambil menyalakan motor mixer pada kecepatan rendah (mencegah pembentukan busa udara berlebih)."),
            ("4. Pemasukan Bahan Pembantu: ", "Larutkan pewarna, pengawet, dan parfum secara terpisah dalam sedikit air, kemudian tuangkan ke dalam tangki."),
            ("5. Pengaturan Viskositas: ", "Tambahkan larutan garam (NaCl) sedikit demi sedikit sambil terus diaduk hingga kekentalan yang diinginkan tercapai."),
            ("6. Degasifikasi (Istirahat Busa): ", "Matikan mixer dan diamkan larutan selama 12–24 jam agar gelembung udara naik dan larutan menjadi jernih transparan."),
            ("7. QC Approval: ", "Ambil sampel in-process untuk pengujian laboratorium sebelum disalurkan ke mesin filling.")
        ],
        [
            ("FRM-PRD-01: ", "Catatan Pengolahan Bets (Batch Processing Record) - Tahap Mixing."),
            ("FRM-PRD-03: ", "Instruksi Kerja (IK) Spesifik per Varian Produk.")
        ]
    ),
    (
        "03_OPERASIONAL_PRODUKSI_DAN_PENGEMASAN",
        "SOP-PRD-011_Pengisian_Filling_Penutupan_dan_Uji_Kebocoran.docx",
        "SOP-PRD-011",
        "Pengisian (Filling), Penutupan Wadah, dan Uji Kebocoran",
        "Menjamin cairan terisi tepat volume sesuai klaim netto pada label, tertutup rapat, dan tidak mengalami kebocoran saat didistribusikan.",
        "Berlaku untuk pengisian kemasan botol (450 mL / 1 Liter) dan jerigen (5 Liter) di Ruang Filling.",
        [
            ("Operator Filling: ", "Mengoperasikan nozel pengisi dan mengatur volume takaran."),
            ("Operator Capping: ", "Memasang dan mengencangkan tutup botol/jerigen."),
            ("Petugas QC: ", "Melakukan penimbangan berat netto berkala dan uji tekan kebocoran.")
        ],
        "UU Perlindungan Konsumen No. 8 Tahun 1999 dan Standar CPKRTB Kemenkes RI.",
        [
            ("1. Setting Nozel: ", "Kalibrasi volume nozel mesin filling menggunakan gelas ukur terkalibrasi."),
            ("2. Proses Pengisian: ", "Letakkan botol/jerigen bersih di bawah nozel pengisi, isi cairan secara presisi tanpa tumpahan pada leher botol."),
            ("3. Penutupan (Capping): ", "Pasang tutup botol dan putar kencang menggunakan alat bantu torsi manual/pneumatik hingga segel mengunci rapat."),
            ("4. Uji Tekan Kebocoran (Sampling): ", "Ambil 5 botol dari setiap 500 botol hasil filling, balikkan botol selama 5 menit dan tekan perlahan untuk memastikan tidak ada rembesan cairan."),
            ("5. Pembersihan Botol: ", "Lap bagian luar botol bila ada tetesan cairan sebelum diteruskan ke meja pelabelan."),
            ("6. Pencatatan: ", "Catat jumlah botol yang berhasil diisi dan jumlah botol reject pada Batch Record.")
        ],
        [
            ("FRM-PRD-01: ", "Catatan Pengolahan Bets (Batch Record) - Lembar Filling & Packaging."),
            ("FRM-PRD-04: ", "Logbook Pemeriksaan Berat Netto dan Uji Kebocoran Kemasan.")
        ]
    ),
    (
        "03_OPERASIONAL_PRODUKSI_DAN_PENGEMASAN",
        "SOP-PRD-012_Pelabelan_Kemasan_dan_Pengemasan_Karton_Sekunder.docx",
        "SOP-PRD-012",
        "Pelabelan Kemasan dan Pengemasan Sekunder ke Karton",
        "Menjamin penempelan stiker label rapi, simetris, mencantumkan identitas legalitas resmi (Nomor PKD, Batch, Exp Date), serta terlindung rapi dalam karton boks.",
        "Berlaku untuk seluruh proses penempelan label stiker dan pengemasan karton boks di Ruang Kemas Sekunder.",
        [
            ("Operator Pelabelan: ", "Menempelkan stiker label dan mencetak kode produksi/kedaluwarsa."),
            ("Operator Packing: ", "Menyusun botol ke dalam karton boks dan menyegelnya."),
            ("Kepala Produksi: ", "Memverifikasi kebenaran informasi nomor batch dan tanggal kedaluwarsa.")
        ],
        "Permenkes No. 62 Tahun 2017 tentang Penandaan dan Izin Edar PKRT.",
        [
            ("1. Verifikasi Label: ", "Pastikan stiker label yang diambil sesuai dengan varian produk yang sedang diproduksi dan mencantumkan nomor KEMENKES RI PKD yang valid."),
            ("2. Pencetakan Kode: ", "Cetak Nomor Bets dan Tanggal Kedaluwarsa (Exp Date) pada area yang telah ditentukan pada label/botol menggunakan mesin inkjet coder atau cap stempel permanen."),
            ("3. Penempelan Stiker: ", "Tempelkan stiker pada posisi tengah badan botol secara simetris, rata, dan bebas dari gelembung udara."),
            ("4. Pengemasan ke Karton: ", "Susun botol ke dalam karton boks (misal: 24 botol/karton) dengan menggunakan sekat kardus pemisah."),
            ("5. Penyegelan Karton: ", "Tutup karton boks dan rekatkan lakban secara kuat pada bagian atas dan bawah karton."),
            ("6. Identitas Luar: ", "Tempel label identitas batch di bagian luar karton boks dan pindahkan ke gudang barang jadi.")
        ],
        [
            ("FRM-PRD-01: ", "Catatan Pengolahan Bets - Rekapitulasi Kemasan Akhir."),
            ("FRM-PRD-05: ", "Contoh Spesimen Label dan Karton Boks.")
        ]
    ),

    # -------------------------------------------------------------------------
    # FOLDER 4: 04_QUALITY_CONTROL_DAN_LABORATORIUM
    # -------------------------------------------------------------------------
    (
        "04_QUALITY_CONTROL_DAN_LABORATORIUM",
        "SOP-QC-013_Pengambilan_Sampel_Uji_Sampling_Protocol.docx",
        "SOP-QC-013",
        "Pengambilan Sampel Uji (Sampling Protocol)",
        "Memperoleh sampel uji yang representatif, homogen, dan tidak terkontaminasi dari bahan baku, proses pengolahan, dan produk jadi untuk analisis laboratorium.",
        "Berlaku untuk kegiatan pengambilan sampel di area gudang bahan baku, ruang mixing, dan ruang kemas.",
        [
            ("Petugas QC: ", "Melaksanakan sampling sesuai metode dan jadwal yang ditetapkan."),
            ("PJT: ", "Menyetujui rencana sampling dan mengevaluasi integritas sampel.")
        ],
        "ISO 9001:2015 Klausul 8.6 (Pelepasan Produk dan Jasa) & Standar CPKRTB.",
        [
            ("1. Persiapan Alat: ", "Gunakan wadah botol kaca/plastik steril, pipet ukur bersih, atau thief sampler khusus bahan kimia."),
            ("2. Sampling Bahan Baku Datang: ", "Ambil sampel dari sejumlah akar n + 1 drum (misal: 9 drum = 4 drum disampling) sebanyak 100 mL per drum dan campur homogen (komposit)."),
            ("3. Sampling In-Process (Tangki Mixer): ", "Ambil 250 mL cairan dari 3 titik (permukaan atas, bagian tengah, dan kran bawah tangki) setelah proses mixing dinyatakan selesai."),
            ("4. Sampling Produk Jadi (Lini Filling): ", "Ambil 1 botol kemasan pada 10% awal filling, 1 botol pada 50% pertengahan, dan 1 botol pada 90% akhir proses."),
            ("5. Pemberian Label Sampel: ", "Beri label pada botol sampel: Nama Produk/Bahan, Nomor Bets, Tanggal Sampling, dan Nama Petugas QC."),
            ("6. Pengiriman ke Lab: ", "Bawa sampel ke laboratorium QC untuk segera dilakukan pengujian fisik-kimia.")
        ],
        [
            ("FRM-QC-01: ", "Logbook Pengambilan Sampel Uji (Sampling Logbook)."),
            ("FRM-QC-02: ", "Label Identitas Wadah Sampel Uji.")
        ]
    ),
    (
        "04_QUALITY_CONTROL_DAN_LABORATORIUM",
        "SOP-QC-014_Pengujian_Kualitas_Produk_Jadi_di_Laboratorium.docx",
        "SOP-QC-014",
        "Pengujian Kualitas Produk Jadi di Laboratorium Pabrik",
        "Menentukan status kelulusan mutu (Lulus/Tolak) dari setiap batch produk PKRT berdasarkan parameter fisik dan kimia yang telah ditetapkan Kemenkes RI.",
        "Berlaku untuk seluruh produk jadi yang diproduksi di PT KCA Chemical sebelum didistribusikan.",
        [
            ("Petugas QC: ", "Melakukan pengujian organoleptik, pH, bobot jenis, dan uji stabilitas busa."),
            ("PJT: ", "Memvalidasi hasil pengujian laboratorium dan menandatangani lembar kerja QC.")
        ],
        "Farmakope Indonesia / Standar Nasional Indonesia (SNI) Sabun Pembersih & Standar Kemenkes RI.",
        [
            ("1. Uji Organoleptik: ", "Periksa warna cairan di bawah pencahayaan putih (harus sesuai standar warna), kejernihan (tidak ada endapan/gumpalan), dan aroma (bau khas sesuai varian)."),
            ("2. Uji Derajat Keasaman (pH): ", "Kalibrasi pH meter dengan larutan buffer pH 4.00, 7.00, dan 10.01. Celupkan probe elektroda ke dalam sampel suhu 25°C. Catat nilai pH (Standar Sabun Cuci Piring: 6.00–8.00; Pembersih Lantai: 6.50–8.50; Pemutih: 11.00–12.50)."),
            ("3. Uji Bobot Jenis (g/mL): ", "Timbang piknometer 25 mL kosong kering (W0), isi dengan aquadest (W1), isi dengan sampel cairan (W2) pada suhu 25°C. Hitung BJ = (W2 - W0) / (W1 - W0). Standar BJ: 1.010 – 1.040 g/mL."),
            ("4. Uji Daya Busa (Sabun): ", "Kocok 10 mL larutan sampel 1% dalam silinder ukur bertutup selama 30 detik, ukur tinggi busa awal (minimal 12 cm) dan stabilitas busa setelah 5 menit."),
            ("5. Kesimpulan: ", "Jika semua parameter memenuhi syarat, tuliskan status 'MEMENUHI SYARAT (PASSED)' pada Lembar Hasil Uji QC.")
        ],
        [
            ("FRM-QC-03: ", "Lembar Pengujian Mutu Laboratorium QC (QC Analysis Sheet)."),
            ("FRM-QC-04: ", "Logbook Kalibrasi Alat Lab (pH Meter & Timbangan Analitik).")
        ]
    ),
    (
        "04_QUALITY_CONTROL_DAN_LABORATORIUM",
        "SOP-QC-015_Pengelolaan_Sampel_Pertinggal_Retained_Sample.docx",
        "SOP-QC-015",
        "Pengelolaan Sampel Pertinggal (Retained Sample)",
        "Menyimpan arsip fisik produk dari setiap nomor batch produksi sebagai bukti historis mutu dan bahan investigasi jika terjadi keluhan komsumen di pasaran.",
        "Berlaku untuk seluruh produk PKRT jadi kemasan komersial lengkap dengan etiket stiker.",
        [
            ("Petugas QC: ", "Mengambil, melabeli, dan menata sampel di ruang penyimpanan sampel pertinggal."),
            ("PJT: ", "Mengawasi masa retensi sampel dan menandatangani berita acara pemusnahan sampel kedaluwarsa.")
        ],
        "CPKRTB Aspek Pengawasan Mutu: Pabrik wajib menyimpan sampel pertinggal dalam jumlah cukup dari setiap batch produk jadi.",
        [
            ("1. Jumlah Pengambilan: ", "Ambil 2 wadah kemasan utuh (lengkap dengan tutup dan label stiker) dari setiap nomor batch produksi yang diluluskan."),
            ("2. Pemberian Label Arsip: ", "Beri stiker khusus bertuliskan 'SAMPEL PERTINGGAL (RETAINED SAMPLE)' lengkap dengan Nama Produk, No. Bets, Tanggal Produksi, dan Tanggal Kedaluwarsa."),
            ("3. Penyimpanan: ", "Simpan sampel di Lemari Khusus Sampel Pertinggal pada suhu kamar terkontrol (20°C - 30°C) dan kelembaban di bawah 70% RH."),
            ("4. Masa Simpan: ", "Sampel disimpan selama masa kedaluwarsa produk ditambah 1 tahun (minimal 3 tahun)."),
            ("5. Pemusnahan: ", "Sampel yang telah melewati masa simpan dapat dimusnahkan dengan menuangkan cairan ke saluran IPAL dan membuat Berita Acara Pemusnahan.")
        ],
        [
            ("FRM-QC-05: ", "Logbook Inventaris Sampel Pertinggal (Retained Sample Log)."),
            ("FRM-QC-06: ", "Berita Acara Pemusnahan Sampel Pertinggal.")
        ]
    ),
    (
        "04_QUALITY_CONTROL_DAN_LABORATORIUM",
        "SOP-QC-016_Penerbitan_Sertifikat_Analisis_CoA_Resmi_PJT.docx",
        "SOP-QC-016",
        "Penerbitan Sertifikat Analisis (Certificate of Analysis - CoA)",
        "Menerbitkan dokumen sertifikat resmi jaminan mutu yang disahkan secara hukum oleh PJT untuk diserahkan kepada pihak pembeli atau PT Distributor.",
        "Berlaku untuk setiap nomor batch produk PKRT yang akan didistribusikan keluar dari pabrik.",
        [
            ("Petugas Analis QC: ", "Memasukkan data hasil pengujian ke dalam draf formulir CoA."),
            ("PJT: ", "Memverifikasi validitas data analisis dan menandatangani dokumen CoA asli.")
        ],
        "ISO 9001:2015 Klausul 8.6 dan Regulasi Izin Edar Kemenkes RI.",
        [
            ("1. Input Data: ", "Setelah seluruh pengujian laboratorium selesai dan dinyatakan Lulus, masukkan data hasil uji (pH, BJ, Kadar Aktif, Pemeriksaan Fisik) ke format blanko resmi CoA."),
            ("2. Pencantuman Legalitas: ", "Pastikan Nomor Izin Edar KEMENKES RI PKD, Nama Merek, Ukuran Kemasan, Nomor Bets, Tanggal Produksi, dan Tanggal Kedaluwarsa tercantum dengan benar."),
            ("3. Verifikasi PJT: ", "PJT mencocokkan data pada draf CoA dengan lembar kerja analisis QC dan Batch Processing Record."),
            ("4. Pengesahan & Stempel: ", "PJT menandatangani lembar CoA asli dan membubuhkan stempel resmi PT KCA Chemical."),
            ("5. Distribusi Dokumen: ", "Lembar asli diserahkan kepada bagian ekspedisi untuk diberikan ke distributor/mitra makloon, dan 1 salinan diarsipkan dalam Binder Rekaman Mutu.")
        ],
        [
            ("FRM-QC-07: ", "Format Blanko Sertifikat Analisis (CoA) Resmi PJT."),
            ("FRM-QC-08: ", "Buku Register Penomoran CoA.")
        ]
    ),

    # -------------------------------------------------------------------------
    # FOLDER 5: 05_PASCA_PRODUKSI_RECALL_DAN_AUDIT_MUTU
    # -------------------------------------------------------------------------
    (
        "05_PASCA_PRODUKSI_RECALL_DAN_AUDIT_MUTU",
        "SOP-MKT-017_Penanganan_Keluhan_Pelanggan_Customer_Complaint.docx",
        "SOP-MKT-017",
        "Penanganan Keluhan Pelanggan dan Konsumen (Complaint Handling)",
        "Merespon, menginvestigasi, dan menuntaskan setiap keluhan atau komplain mutu produk dari distributor/konsumen secara cepat dan profesional.",
        "Berlaku untuk seluruh komplain teknis mutu, kebocoran kemasan, aroma, kekentalan, atau daya bersih produk berizin PKD.",
        [
            ("Tim Marketing / Customer Care: ", "Menerima dan mencatat laporan keluhan dari distributor/konsumen dalam waktu 1x24 jam."),
            ("PJT & Tim QC: ", "Melakukan investigasi teknis terhadap sampel pertinggal nomor batch terkait."),
            ("Direktur Utama: ", "Menyetujui keputusan tindakan koreksi atau penggantian produk.")
        ],
        "ISO 9001:2015 Klausul 9.1.2 (Kepuasan Pelanggan) & 10.2 (Ketidaksesuaian dan Tindakan Korektif).",
        [
            ("1. Penerimaan Komplain: ", "Catat identitas pelapor, tanggal komplain, nama produk, nomor bets pada kemasan, dan foto/video bukti kerusakan pada Formulir Keluhan Pelanggan."),
            ("2. Pengambilan Sampel Arsip: ", "PJT mengambil Sampel Pertinggal dari nomor batch yang sama di lemari arsip untuk dilakukan pengujian laboratorium pembanding."),
            ("3. Analisis Akar Masalah: ", "Tentukan apakah cacat disebabkan oleh kesalahan formulasi/manufaktur pabrik, kesalahan penyimpanan di gudang toko, atau kerusakan saat ekspedisi."),
            ("4. Tindakan Korektif (CAPA): ", "Jika kesalahan berasal dari proses produksi, terbitkan laporan CAPA untuk perbaikan proses mixing/filling."),
            ("5. Tanggapan Resmi & Kompensasi: ", "Kirimkan surat tanggapan resmi kepada pelanggan dan lakukan penggantian barang baru yang cacat dalam waktu maksimal 7 hari kerja."),
            ("6. Penutupan Kasus: ", "PJT menandatangani formulir penutupan komplain setelah pelanggan menyatakan puas.")
        ],
        [
            ("FRM-MKT-01: ", "Formulir Rekaman Keluhan Pelanggan (Customer Complaint Form)."),
            ("FRM-QMS-01: ", "Laporan Tindakan Korektif dan Preventif (CAPA Report).")
        ]
    ),
    (
        "05_PASCA_PRODUKSI_RECALL_DAN_AUDIT_MUTU",
        "SOP-REG-018_Penarikan_Kembali_Produk_dari_Peredaran_Product_Recall.docx",
        "SOP-REG-018",
        "Penarikan Kembali Produk dari Peredaran (Product Recall)",
        "Menjamin pelaksanaan penarikan produk dari jalur distribusi pasar secara cepat, terkendali, dan sistematis apabila ditemukan cacat mutu kritis yang berpotensi membahayakan kesehatan masyarakat atau atas perintah Kemenkes RI.",
        "Berlaku untuk seluruh produk PKRT buatan PT KCA Chemical yang telah beredar di pasaran.",
        [
            ("Direktur Utama: ", "Memutuskan pelaksanaan penarikan produk dan membentuk Tim Recall."),
            ("PJT: ", "Memimpin teknis pelaksanaan recall, inventarisasi data bets, dan pelaporan ke Kemenkes RI."),
            ("Tim Ekspedisi & Logistik: ", "Menarik fisik produk dari gudang distributor dan toko-toko retail.")
        ],
        "Permenkes No. 62 Tahun 2017 tentang Izin Edar dan Pengawasan PKRT & Standar ISO 9001:2015.",
        [
            ("1. Pembentukan Tim: ", "Direktur menerbitkan Surat Perintah Penarikan Produk dan menunjuk PJT sebagai Koordinator Tim Recall."),
            ("2. Penelusuran Distribusi: ", "Buka arsip Surat Jalan (Delivery Order) untuk mengetahui lokasi distributor, toko, dan jumlah unit dari batch yang ditarik."),
            ("3. Penerbitan Surat Recall: ", "Kirimkan Surat Pemberitahuan Penarikan Produk kepada seluruh mitra distributor dalam waktu maksimal 2x24 jam (menginstruksikan untuk menyetop penjualan dan mengkarantina produk)."),
            ("4. Penarikan Fisik: ", "Ambil dan kumpulkan seluruh produk yang ditarik ke gudang pabrik, letakkan di area khusus bertanda 'BARANG DITOLAK / RECALL'."),
            ("5. Pelaporan Kemenkes: ", "Kirim laporan resmi perkembangan penarikan produk kepada Direktorat Jenderal Farmalkes Kemenkes RI."),
            ("6. Pemusnahan Produk: ", "Lakukan pemusnahan barang recall dengan disaksikan oleh Direktur dan PJT, lalu buat Berita Acara Pemusnahan Resmi.")
        ],
        [
            ("FRM-REG-01: ", "Surat Pemberitahuan Penarikan Produk (Recall Notice Letter)."),
            ("FRM-REG-02: ", "Laporan Rekapitulasi Hasil Penarikan Produk (Recall Summary Report)."),
            ("FRM-REG-03: ", "Berita Acara Pemusnahan Produk Recall.")
        ]
    ),
    (
        "05_PASCA_PRODUKSI_RECALL_DAN_AUDIT_MUTU",
        "SOP-PRD-019_Penanganan_Produk_Tidak_Sesuai_Reject_Rework.docx",
        "SOP-PRD-019",
        "Penanganan Produk Tidak Sesuai (Reject / Rework)",
        "Mengendalikan bahan baku, produk antara, atau produk jadi yang tidak memenuhi spesifikasi mutu agar tidak terproses lebih lanjut atau terdistribusi ke pasar.",
        "Berlaku untuk cairan hasil mixing yang OOS (Out of Specification), botol bocor, stiker rusak, dan karton cacat.",
        [
            ("Petugas QC: ", "Memberi label status tidak sesuai dan mengisolasi produk."),
            ("PJT: ", "Menentukan keputusan disposisi: pengerjaan ulang (rework) atau pemusnahan (reject)."),
            ("Kepala Produksi: ", "Melaksanakan instruksi teknis rework atau pemusnahan.")
        ],
        "ISO 9001:2015 Klausul 8.7 (Pengendalian Ketidaksesuaian Keluaran) & CPKRTB.",
        [
            ("1. Identifikasi & Karantina: ", "Tempel label merah bertuliskan 'DITOLAK / REJECT' atau label oranye 'DIREWORK' pada tangki/wadah yang tidak lulus QC."),
            ("2. Investigasi Ketidaksesuaian: ", "PJT dan Analis QC memeriksa penyebab penyimpangan (misal pH terlalu asam atau viskositas kurang kental)."),
            ("3. Disposisi Rework: ", "Jika formula dapat diperbaiki (misal: menambahkan larutan NaCl untuk menaikkan kekentalan, atau menambah soda ash untuk menaikkan pH), PJT menerbitkan Instruksi Kerja Rework tertulis."),
            ("4. Pengujian Ulang: ", "Lakukan pengujian laboratorium lengkap terhadap batch yang telah dirework hingga memenuhi seluruh parameter spesifikasi."),
            ("5. Disposisi Pemusnahan: ", "Jika formula tidak dapat diperbaiki, netralkan cairan dan buang ke saluran IPAL berizin, serta musnahkan botol/label yang rusak."),
            ("6. Pencatatan: ", "Catat seluruh riwayat ketidaksesuaian pada Laporan Produk Tidak Sesuai.")
        ],
        [
            ("FRM-PRD-06: ", "Laporan Produk Tidak Sesuai (Non-Conforming Product Report)."),
            ("FRM-PRD-07: ", "Catatan Pengolahan Pengerjaan Ulang (Rework Record).")
        ]
    ),
    (
        "05_PASCA_PRODUKSI_RECALL_DAN_AUDIT_MUTU",
        "SOP-QMS-020_Audit_Internal_dan_Inspeksi_Diri_ISO9001_CPKRTB.docx",
        "SOP-QMS-020",
        "Audit Internal dan Inspeksi Diri Standar ISO 9001 & CPKRTB",
        "Memastikan seluruh sistem manajemen mutu pabrik, SOP, catatan batch, dan fasilitas fisik berjalan secara konsisten, efektif, dan siap menghadapi audit sertifikasi atau inspeksi berkala Kemenkes/Dinkes.",
        "Berlaku untuk seluruh divisi di PT KCA Chemical (Gudang, Produksi, QC, PJT, Marketing, dan Manajemen).",
        [
            ("Tim Auditor Internal: ", "Melakukan audit independen terhadap divisi yang tidak di bawah tanggung jawab langsungnya."),
            ("Management Representative (MR) / PJT: ", "Menyusun jadwal audit tahunan dan memonitor tindak lanjut temuan audit."),
            ("Direktur Utama: ", "Memimpin Rapat Tinjauan Manajemen (Management Review) berdasarkan hasil audit.")
        ],
        "ISO 9001:2015 Klausul 9.2 (Audit Internal) & 9.3 (Tinjauan Manajemen) serta Standar Inspeksi Diri CPKRTB.",
        [
            ("1. Program Audit: ", "Susun Rencana Audit Internal tahunan (dilaksanakan minimal 2 kali dalam setahun atau setiap 6 bulan sekali)."),
            ("2. Rapat Pembukaan (Opening Meeting): ", "Auditor menyampaikan ruang lingkup, kriteria audit, dan jadwal visitasi area."),
            ("3. Pelaksanaan Audit Lapangan: ", "Periksa bukti fisik di lapangan: kebersihan ruangan, kalibrasi timbangan, pengisian Batch Record harian, kesesuaian SOP, dan ketersediaan APD."),
            ("4. Pencatatan Temuan: ", "Kelompokkan temuan ke dalam kategori: Mayor, Minor, atau Saran Perbaikan (Observasi) pada Formulir Laporan Temuan Audit."),
            ("5. Tindak Lanjut Korektif (CAPA): ", "Auditee wajib menyusun rencana perbaikan dan menyelesaikan akar masalah maksimal dalam waktu 30 hari kalender."),
            ("6. Verifikasi & Tinjauan Manajemen: ", "Auditor melakukan verifikasi efektivitas perbaikan dan melaporkannya dalam Rapat Tinjauan Manajemen tahunan.")
        ],
        [
            ("FRM-QMS-02: ", "Jadwal dan Program Audit Internal Tahunan."),
            ("FRM-QMS-03: ", "Checklist Audit Internal Standar CPKRTB & ISO 9001:2015."),
            ("FRM-QMS-01: ", "Laporan Tindakan Korektif dan Preventif (CAPA Report).")
        ]
    )
]

def build_all_individual_sops():
    print("Building all 20 individual SOP documents into KCA DOKUMEN folder structure...")
    for item in sops_data:
        folder_name, filename, doc_no, title, tujuan, ruang_lingkup, tanggung_jawab, definisi, prosedur_list, formulir_terkait = item
        create_individual_sop(
            folder_name=folder_name,
            filename=filename,
            doc_no=doc_no,
            title=title,
            tujuan=tujuan,
            ruang_lingkup=ruang_lingkup,
            tanggung_jawab=tanggung_jawab,
            definisi=definisi,
            prosedur_list=prosedur_list,
            formulir_terkait=formulir_terkait
        )
        
    # Also copy the Master Manuals & Contracts into folder 06_MASTER_MANUAL_DAN_LEGALITAS
    master_folder = os.path.join(BASE_DIR, "06_MASTER_MANUAL_DAN_LEGALITAS")
    os.makedirs(master_folder, exist_ok=True)
    
    src_dir = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/dokumen_iso9001_pkrt"
    if os.path.exists(src_dir):
        for f in os.listdir(src_dir):
            if f.endswith(".docx"):
                shutil.copy2(os.path.join(src_dir, f), os.path.join(master_folder, f))
                print(f"Copied master to 06: {f}")
                
    print("\nALL 20 INDIVIDUAL SOPS AND MASTER FILES SUCCESSFULLY CREATED IN 'KCA DOKUMEN'!")

if __name__ == "__main__":
    build_all_individual_sops()

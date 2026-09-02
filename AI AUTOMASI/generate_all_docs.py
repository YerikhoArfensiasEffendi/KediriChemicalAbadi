import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

OUTPUT_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/dokumen_iso9001_pkrt"
os.makedirs(OUTPUT_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PT [NAMA PERUSAHAAN]\nMANUFAKTUR PKRT\nSISTEM ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block(doc, left_title="Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", right_title="Disetujui Oleh,\nDirektur Perusahaan", left_name="( ............................................ )", right_name="( ............................................ )"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    
    col_widths = [Cm(7.5), Cm(8.0)]
    
    cell_l = table.cell(0, 0)
    cell_l.width = col_widths[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l1 = p_l.add_run(left_title + "\n\n\n\n\n" + left_name + "\nTanggal: ...................................")
    r_l1.font.name = 'Times New Roman'
    r_l1.font.size = Pt(9.5)
    r_l1.font.color.rgb = COLOR_BLACK
    
    cell_r = table.cell(0, 1)
    cell_r.width = col_widths[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r.add_run(right_title + "\n\n\n\n\n" + right_name + "\nTanggal: ...................................")
    r_r1.font.name = 'Times New Roman'
    r_r1.font.size = Pt(9.5)
    r_r1.font.color.rgb = COLOR_BLACK

# ==============================================================================
# GENERATE DOC 1: MASTER ROADMAP & TAHAPAN DARI NOL
# ==============================================================================
def create_doc_1():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Master Roadmap & Tahapan Persiapan Pendirian Pabrik PKRT, Makloon & Izin Edar Standar ISO 9001:2015", "RMP-PKRT-001", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. TUJUAN DAN RUANG LINGKUP")
    add_p(doc, "Dokumen Master Roadmap ini disusun sebagai pedoman operasional dan strategis bagi PT [Nama Perusahaan] (PT Perorangan) dalam membangun sarana industri manufaktur Perbekalan Kesehatan Rumah Tangga (PKRT), melayani jasa makloon (toll manufacturing), dan mengurus perizinan edar produk (KEMENKES RI PKD) secara terstruktur sesuai standar ISO 9001:2015 dan CPKRTB (Cara Pembuatan PKRT yang Baik).")
    
    add_heading_1(doc, "2. ROADMAP TAHAP DEMI TAHAP (STEP-BY-STEP WORKFLOW)")
    
    headers = ["Fase", "Tahapan Kegiatan", "Dokumen Output", "Penanggung Jawab", "Estimasi Waktu"]
    data = [
        ["Fase 0", "Legalitas Badan Usaha PT Perorangan & OSS RBA", "Sertifikat AHU, NPWP PT, NIB KBLI 20231, SPPL", "Direktur", "Minggu 1"],
        ["Fase 1", "Penyediaan Tenaga Ahli PJT & Penataan Sarana CPKRTB", "Surat Penunjukan PJT, Ijazah/STRTTK, Layout 2D", "Direktur & PJT", "Minggu 2 - 3"],
        ["Fase 2", "Pengajuan Sertifikat Standar Produksi PKRT ke Kemenkes", "Sertifikat Produksi PKRT (e-Farmalkes)", "PJT & Direktur", "Minggu 3 - 5"],
        ["Fase 3", "Penyiapan Kerja Sama Makloon & Registrasi HKI Klien", "Kontrak Makloon Notariil & Bukti DJKI Merek", "Direktur & Legal Mitra", "Minggu 4 - 5"],
        ["Fase 4", "Riset Formula, Trial Batch & Uji Laboratorium Terakreditasi", "CoA Bahan Baku, MSDS, Hasil Uji Lab KAN", "PJT & Lab Eksternal", "Minggu 5 - 7"],
        ["Fase 5", "Pendaftaran Izin Edar KEMENKES RI PKD (e-Reg PKRT)", "Nomor Izin Edar KEMENKES RI PKD", "PJT via e-Farmalkes", "Minggu 7 - 10"],
        ["Fase 6", "Implementasi ISO 9001:2015, Komersialisasi & TKDN-IK", "Batch Record, CoA Rilis, Sertifikat TKDN Kemenperin", "PJT & Tim Pabrik", "Berkelanjutan"]
    ]
    add_styled_table(doc, headers, data, [Cm(1.8), Cm(4.5), Cm(4.2), Cm(2.8), Cm(2.2)])
    
    add_heading_1(doc, "3. DETAIL PELAKSANAAN SETIAP FASE")
    
    add_heading_2(doc, "3.1 Fase 0: Pondasi Legalitas Badan Usaha & Tata Ruang")
    add_bullet(doc, "Mendaftarkan pendirian PT Perorangan melalui portal Ditjen AHU Kemenkumham hingga terbit Surat Pernyataan Pendirian dan Sertifikat AHU.", "Legalitas PT: ");
    add_bullet(doc, "Membuat NPWP Badan Usaha atas nama PT Perorangan di KPP Pratama / DJP Online.", "Perpajakan: ");
    add_bullet(doc, "Mendaftarkan akun OSS RBA (oss.go.id) dan memilih KBLI 20231 (Industri Sabun dan Bahan Pembersih Keperluan Rumah Tangga).", "Perizinan OSS: ");
    add_bullet(doc, "Memastikan Kesesuaian Kegiatan Pemanfaatan Ruang (KKPR) dan menerbitkan Surat Pernyataan Kesanggupan Pengelolaan dan Pemantauan Lingkungan Hidup (SPPL) secara otomatis di sistem OSS.", "Lingkungan: ");

    add_heading_2(doc, "3.2 Fase 1: Penataan Fasilitas Fisik & Pengangkatan PJT")
    add_bullet(doc, "Menunjuk Penanggung Jawab Teknis (PJT) dengan kualifikasi minimal D3/S1 Farmasi, S1 Kimia, atau S1 Teknik Kimia dengan Surat Pernyataan Bekerja Penuh Waktu.", "PJT: ");
    add_bullet(doc, "Menata zonasi pabrik menerapkan One-Way Flow: Ruang Ganti/Loker -> Gudang Bahan Baku -> Ruang Timbang -> Ruang Mixing -> Ruang Filling/Kemas Primer -> Ruang Kemas Sekunder -> Gudang Barang Jadi.", "Layout CPKRTB: ");
    add_bullet(doc, "Menyediakan peralatan utama: Tangki Mixing Stainless Steel/HDPE, Timbangan Digital terkalibrasi, pH Meter Digital, Mesin Filling semi-otomatis/manual, serta APD lengkap.", "Peralatan: ");

    add_heading_2(doc, "3.3 Fase 2: Penerbitan Sertifikat Standar Produksi PKRT")
    add_bullet(doc, "Mengajukan permohonan Sertifikat Standar Produksi PKRT melalui sistem e-Farmalkes Kemenkes RI.", "Pengajuan: ");
    add_bullet(doc, "Melampirkan Dokumen Mutu: Manual Mutu, Daftar SOP Sanitasi & Produksi, Denah Layout, Struktur Organisasi, dan Berkas PJT.", "Dokumen Pendukung: ");
    add_bullet(doc, "Menjalani proses verifikasi dokumen dan/atau visitasi lapangan oleh Dinas Kesehatan / Balai Pengawasan.", "Verifikasi: ");

    add_heading_2(doc, "3.4 Fase 3: Skema Kontrak Makloon & Kepemilikan Merek")
    add_bullet(doc, "Memverifikasi legalitas PT Mitra / Distributor (Akta, NPWP, NIB KBLI Distribusi 46499).", "Verifikasi Klien: ");
    add_bullet(doc, "Memastikan Merek Dagang milik Mitra telah terdaftar di DJKI Kemenkumham atau minimal memiliki Tanda Terima Pendaftaran DJKI.", "Merek Dagang: ");
    add_bullet(doc, "Menandatangani Surat Perjanjian Kerja Sama Makloon (Toll Manufacturing Agreement) bermaterai resmi yang memuat hak, kewajiban, formula, dan standar kualitas.", "Perjanjian: ");

    add_heading_2(doc, "3.5 Fase 4 & 5: Pengujian Laboratorium & Izin Edar Kemenkes RI PKD")
    add_bullet(doc, "Melakukan uji mutu produk jadi di laboratorium terakreditasi KAN (Uji pH, Bobot Jenis, Viskositas, dan Uji Efektivitas Daya Bunuh Kuman untuk Pemutih/Disinfektan).", "Uji Laboratorium: ");
    add_bullet(doc, "Mendaftarkan Izin Edar PKRT Dalam Negeri (PKD) melalui portal ereg.farmalkes.kemkes.go.id dengan melampirkan Formula, Sertifikat Produksi, Kontrak Makloon, HKI, CoA, dan Desain Etiket.", "e-Reg PKRT: ");
    add_bullet(doc, "Melakukan pembayaran PNBP Kemenkes RI dan memantau verifikasi teknis hingga terbit Sertifikat Izin Edar PKD elektronik.", "Penerbitan NIE: ");

    add_heading_2(doc, "3.6 Fase 6: Tata Kelola Mutu ISO 9001 & TKDN-IK Kemenperin")
    add_bullet(doc, "Menjalankan sistem Batch Record dan menerbitkan Certificate of Analysis (CoA) untuk setiap pengiriman barang ke distributor.", "Pengendalian Bets: ");
    add_bullet(doc, "Mendaftarkan akun SIINas Kemenperin dan mengajukan Sertifikasi TKDN-IK (Industri Kecil) secara gratis guna memperluas penetrasi pasar pengadaan barang pemerintah (e-Katalog LKPP).", "Sertifikasi TKDN: ");
    add_bullet(doc, "Melaporkan realisasi produksi secara berkala melalui e-Report Kemenkes RI.", "Pelaporan Berkala: ");

    add_signature_block(doc)
    
    file_path = os.path.join(OUTPUT_DIR, "01_PANDUAN_DAN_ROADMAP_PERSIAPAN_PABRIK_PKRT_ISO9001.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# GENERATE DOC 2: MASTER CHECKLIST & LIST DATA KENDALI MUTU
# ==============================================================================
def create_doc_2():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Master Checklist & Daftar Dokumen Kendali Mutu Standar ISO 9001:2015 & CPKRTB", "CKL-PKRT-002", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. MASTER LIST STRUKTUR DOKUMEN SISTEM MANAJEMEN MUTU")
    add_p(doc, "Sesuai klausul 7.5 ISO 9001:2015 (Informasi Terdokumentasi), seluruh dokumen sistem mutu pabrik diklasifikasikan ke dalam 4 (empat) tingkatan hierarki:")
    
    headers_1 = ["Level Dokumen", "Jenis Dokumen", "Kode / Format", "Lokasi Arsip", "Pengendali"]
    data_1 = [
        ["Level 1", "Manual Mutu & Kebijakan Mutu Pabrik", "MM-PKRT-001", "Binder 1 (Fisik) & Cloud", "Direktur & PJT"],
        ["Level 2", "Standar Operasional Prosedur (SOP)", "SOP-PKRT-001 s/d 020", "Binder 2 (Fisik) & Ruang Kerja", "PJT & Ka. Produksi"],
        ["Level 3", "Instruksi Kerja (IK) & Master Formula", "IK-PROD-001 s/d 010", "Dinding Ruang Kerja & Lab", "Ka. Produksi & QC"],
        ["Level 4", "Formulir, Logbook & Rekaman Mutu", "FRM-PKRT-001 s/d 015", "Binder 3 & Arsip Harian", "Operator & PJT"]
    ]
    add_styled_table(doc, headers_1, data_1, [Cm(2.2), Cm(4.5), Cm(3.2), Cm(3.2), Cm(2.5)])
    
    add_heading_1(doc, "2. CHECKLIST DOKUMEN PERSYARATAN LEGALITAS & SARANA PABRIK")
    
    headers_2 = ["No", "Nama Dokumen / Izin", "Instansi Penerbit", "Status Kelengkapan", "Target Pemenuhan"]
    data_2 = [
        ["1", "Sertifikat Pendaftaran Pendirian PT Perorangan", "Kemenkumham RI", "[  ] Lengkap / [  ] Proses", "Wajib - Awal"],
        ["2", "NPWP Badan Usaha PT Perorangan", "DJP Kemenkeu", "[  ] Lengkap / [  ] Proses", "Wajib - Awal"],
        ["3", "NIB OSS RBA (KBLI 20231)", "Kementerian Investasi / BKPM", "[  ] Lengkap / [  ] Proses", "Wajib - Awal"],
        ["4", "Persetujuan Tata Ruang (KKPR) & SPPL", "OSS RBA / KLHK", "[  ] Lengkap / [  ] Proses", "Wajib - Awal"],
        ["5", "Surat Keputusan Pengangkatan PJT & Ijazah", "Internal PT & Kampus", "[  ] Lengkap / [  ] Proses", "Wajib - Fase 1"],
        ["6", "Surat Pernyataan Bekerja Penuh Waktu PJT", "PJT & Notaris/Materai", "[  ] Lengkap / [  ] Proses", "Wajib - Fase 1"],
        ["7", "Denah Tata Ruang (Layout 2D) Standar CPKRTB", "Internal PT / Arsitek Teknis", "[  ] Lengkap / [  ] Proses", "Wajib - Fase 1"],
        ["8", "Sertifikat Standar Produksi PKRT", "Ditjen Farmalkes Kemenkes", "[  ] Lengkap / [  ] Proses", "Wajib - Fase 2"],
        ["9", "Hasil Uji Lab Terakreditasi KAN (Produk Jadi)", "Lab Terakreditasi KAN", "[  ] Lengkap / [  ] Proses", "Wajib - Fase 4"],
        ["10", "Nomor Izin Edar KEMENKES RI PKD", "Ditjen Farmalkes Kemenkes", "[  ] Lengkap / [  ] Proses", "Wajib - Fase 5"],
        ["11", "Akun SIINas & Sertifikat TKDN-IK", "Kementerian Perindustrian", "[  ] Lengkap / [  ] Proses", "Pengembangan Pasar"]
    ]
    add_styled_table(doc, headers_2, data_2, [Cm(1.0), Cm(5.0), Cm(3.8), Cm(3.2), Cm(2.5)])
    
    add_heading_1(doc, "3. MATRIKS TANGGUNG JAWAB PERSONALIA (JOB MATRIX)")
    
    headers_3 = ["Jabatan", "Kualifikasi", "Tugas Pokok & Wewenang SMM ISO 9001", "Pelaporan Kepada"]
    data_3 = [
        ["Direktur Perusahaan", "Pimpinan PT", "Menetapkan Kebijakan Mutu, menyetujui anggaran fasilitas, menandatangani kontrak kerja sama makloon dan legalitas badan usaha.", "Pemegang Saham"],
        ["Penanggung Jawab Teknis (PJT)", "D3/S1 Farmasi / Kimia / T.Kimia", "Bertanggung jawab atas pemenuhan CPKRTB, menyetujui Master Formula, memvalidasi hasil uji laboratorium, menerbitkan CoA, dan mengelola registrasi PKD.", "Direktur"],
        ["Kepala Produksi", "Minimal SMK Kimia / D3", "Mengawasi pelaksanaan proses penimbangan, mixing, filling, sanitasi tangki, dan memastikan pengisian Catatan Pengolahan Bets (Batch Record).", "PJT & Direktur"],
        ["Petugas QC & Lab", "SMK Kimia / Analis", "Melakukan sampling bahan baku & produk jadi, pengujian pH, berat jenis, uji kebocoran kemasan, dan mengelola sampel pertinggal (retained sample).", "PJT"],
        ["Petugas Gudang", "Minimal SMA/SMK", "Mengelola penerimaan bahan baku, menerapkan prinsip FIFO/FEFO, menjaga kartu stok, dan memantau suhu/kelembaban gudang penyimpanan.", "Kepala Produksi"]
    ]
    add_styled_table(doc, headers_3, data_3, [Cm(2.8), Cm(2.8), Cm(7.2), Cm(2.8)])
    
    add_signature_block(doc)
    
    file_path = os.path.join(OUTPUT_DIR, "02_MASTER_CHECKLIST_DAN_LIST_DATA_KENDALI_MUTU.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# GENERATE DOC 3: SURAT PENAWARAN HARGA MAKLOON (QUOTATION)
# ==============================================================================
def create_doc_3():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Surat Penawaran Harga & Skema Kerja Sama Jasa Makloon Manufaktur PKRT", "QTO-PKRT-003", "00", "19 Agustus 2026")
    
    add_p(doc, "Nomor Surat  : 018/QTO/PKRT-DIR/VIII/2026\nLampiran     : 1 (Satu) Berkas Skema Harga & Spesifikasi\nPerihal      : Penawaran Biaya Jasa Makloon Produksi PKRT & Pendampingan Izin Edar PKD")
    
    add_p(doc, "Kepada Yth.\nDireksi / Manajemen PT [Nama Perusahaan Mitra / Distributor]\nDi Tempat\n")
    
    add_p(doc, "Dengan hormat,\nSehubungan dengan rencana pengembangan produk Perbekalan Kesehatan Rumah Tangga (PKRT) dengan merek dagang milik Perusahaan Bapak/Ibu, kami dari PT [Nama Perusahaan Anda] sebagai industri manufaktur resmi yang telah tersertifikasi CPKRTB Kementerian Kesehatan RI dan menerapkan Sistem Manajemen Mutu ISO 9001:2015, dengan ini mengajukan penawaran kerja sama produksi makloon (toll manufacturing).")
    
    add_heading_1(doc, "1. RINCIAN BIAYA JASA PRODUKSI MAKLOON (TOLL MANUFACTURING FEE)")
    
    headers = ["No", "Kategori Produk", "Ukuran Kemasan", "Biaya Jasa Makloon Murni", "Biaya Makloon Full Service*", "Minimal Order (MOQ)"]
    data = [
        ["1", "Sabun Cuci Piring Cair (Varian Lime/Lemon)", "Botol 450 mL\nJerigen 5 Liter", "Rp 850 / botol\nRp 4.500 / jerigen", "Rp 4.200 / botol\nRp 28.500 / jerigen", "1.000 Botol\n100 Jerigen"],
        ["2", "Pembersih Lantai Antibakteri (Aroma Pine/Floral)", "Pouch 750 mL\nJerigen 5 Liter", "Rp 950 / pouch\nRp 4.500 / jerigen", "Rp 4.800 / pouch\nRp 29.500 / jerigen", "1.000 Pouch\n100 Jerigen"],
        ["3", "Deterjen Cair Konsentrat (Liquid Detergent)", "Botol 1 Liter\nJerigen 5 Liter", "Rp 1.100 / botol\nRp 4.800 / jerigen", "Rp 7.500 / botol\nRp 34.000 / jerigen", "1.000 Botol\n100 Jerigen"],
        ["4", "Pemutih Pakaian / Disinfektan (NaOCl 5,25%)", "Botol 500 mL\nJerigen 5 Liter", "Rp 900 / botol\nRp 4.500 / jerigen", "Rp 4.500 / botol\nRp 27.500 / jerigen", "1.000 Botol\n100 Jerigen"]
    ]
    add_styled_table(doc, headers, data, [Cm(0.8), Cm(4.2), Cm(2.6), Cm(2.8), Cm(3.0), Cm(2.2)])
    
    add_p(doc, "*Keterangan:\n- Opsi Makloon Murni: Formula dasar dan proses mixing/filling disediakan oleh pabrik kami, bahan baku/kemasan disediakan oleh pihak mitra.\n- Opsi Makloon Full Service: Biaya sudah mencakup penyediaan formula teruji, bahan baku standar industri, kemasan, stiker label, proses mixing, filling, sealing, karton boks, dan pengujian QC.")

    add_heading_1(doc, "2. BIAYA PENDAMPINGAN REGISTRASI IZIN EDAR KEMENKES RI PKD")
    add_p(doc, "Pabrik kami menyediakan layanan terpadu pengurusan izin edar hingga terbit Nomor KEMENKES RI PKD resmi:")
    add_bullet(doc, "Penyusunan Master Formula kualitatif-kuantitatif, pembuatan prototipe sampel, dan uji laboratorium terakreditasi KAN (Uji Efektivitas, pH, Viskositas, dan Stabilitas).", "Paket Uji Laboratorium: ");
    add_bullet(doc, "Pendaftaran akun e-Farmalkes, verifikasi dokumen teknis, pengunggahan kontrak makloon dan HKI, serta pembayaran resmi PNBP Kemenkes RI.", "Paket Registrasi e-Reg PKD: ");
    add_bullet(doc, "Penerbitan Certificate of Analysis (CoA) resmi dari PJT untuk setiap pengiriman batch komersial.", "Jaminan Dokumen Mutu: ");

    add_heading_1(doc, "3. SYARAT DAN KETENTUAN PEMBAYARAN (TERMS OF PAYMENT)")
    add_bullet(doc, "Uang Muka (Down Payment / DP) sebesar 50% dibayarkan pada saat penandatanganan Surat Perintah Kerja (PO).", "Termin 1: ");
    add_bullet(doc, "Pelunasan sisa 50% dibayarkan setelah proses produksi selesai dan hasil uji QC (CoA) dinyatakan Lulus sebelum pengiriman barang.", "Termin 2: ");
    add_bullet(doc, "Waktu pengerjaan produksi (Lead Time) adalah 7 - 14 hari kerja terhitung sejak bahan baku & kemasan lengkap tersedia di pabrik.", "Lead Time: ");

    add_signature_block(doc, "Diajukan Oleh,\nPT [Nama Perusahaan Anda]", "Disetujui Oleh Mitra,\nPT [Nama Mitra Distributor]", "( Direktur Utama )", "( Direktur / Pimpinan )")
    
    file_path = os.path.join(OUTPUT_DIR, "03_SURAT_PENAWARAN_HARGA_MAKLOON_QUOTATION.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# GENERATE DOC 4: SURAT PERJANJIAN KERJASAMA MAKLOON (KONTRAK)
# ==============================================================================
def create_doc_4():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Surat Perjanjian Kerja Sama Jasa Produksi Makloon (Toll Manufacturing Agreement)", "AGR-PKRT-004", "00", "19 Agustus 2026")
    
    add_p(doc, "Pada hari ini, [Hari], tanggal [Tanggal] bulan [Bulan] tahun [Dua Ribu Dua Puluh Enam] ([Tgl-Bln-Thn]), bertempat di [Kota Kantor/Pabrik], telah dibuat dan ditandatangani Perjanjian Kerja Sama Jasa Produksi Makloon (selanjutnya disebut \"Perjanjian\") oleh dan antara:")
    
    add_bullet(doc, "PT [NAMA PERUSAHAAN ANDA], suatu badan usaha yang didirikan berdasarkan hukum Republik Indonesia, berkedudukan di [Alamat Lengkap Pabrik], dalam hal ini diwakili secara sah oleh [Nama Direktur Anda] selaku Direktur Utama, bertindak untuk dan atas nama perseroan (selanjutnya disebut sebagai \"PIHAK PERTAMA / PRODUSEN\").", "1. ");
    add_bullet(doc, "PT [NAMA PERUSAHAAN MITRA], suatu badan usaha yang didirikan berdasarkan hukum Republik Indonesia, berkedudukan di [Alamat Kantor Mitra], dalam hal ini diwakili secara sah oleh [Nama Direktur Mitra] selaku Direktur Utama, bertindak untuk dan atas nama perseroan (selanjutnya disebut sebagai \"PIHAK KEDUA / PEMILIK MEREK\").", "2. ");
    
    add_p(doc, "PIHAK PERTAMA dan PIHAK KEDUA secara bersama-sama disebut \"PARA PIHAK\", dan masing-masing disebut \"PIHAK\". PARA PIHAK terlebih dahulu menerangkan hal-hal sebagai berikut:\n"
               "a. Bahwa PIHAK PERTAMA adalah industri manufaktur yang memiliki sarana, peralatan, tenaga ahli Penanggung Jawab Teknis (PJT), dan izin produksi resmi untuk memproduksi produk PKRT sesuai standar CPKRTB.\n"
               "b. Bahwa PIHAK KEDUA adalah pemilik hak merek dagang yang bermaksud mempercayakan pembuatan dan pengemasan produk pembersih/PKRT kepada PIHAK PERTAMA.\n"
               "Maka, berdasarkan pertimbangan tersebut, PARA PIHAK sepakat untuk membuat perjanjian dengan ketentuan sebagai berikut:")
    
    add_heading_1(doc, "PASAL 1: DEFINISI DAN RUANG LINGKUP")
    add_bullet(doc, "\"Jasa Makloon\" adalah kegiatan manufaktur, pencampuran (mixing), pengisian (filling), pengemasan, dan pengendalian mutu produk PKRT yang dilakukan oleh PIHAK PERTAMA atas instruksi PIHAK KEDUA.", "1.1 ");
    add_bullet(doc, "\"Produk\" adalah seluruh varian produk pembersih rumah tangga yang disepakati untuk diproduksi, meliputi Sabun Cuci Piring, Pembersih Lantai, Deterjen Cair, dan Pemutih Pakaian.", "1.2 ");
    add_bullet(doc, "\"Izin Edar PKD\" adalah izin edar resmi yang diterbitkan oleh Kementerian Kesehatan RI untuk produk dalam negeri yang diproduksi di sarana pabrik PIHAK PERTAMA.", "1.3 ");

    add_heading_1(doc, "PASAL 2: HAK DAN KEWAJIBAN PIHAK PERTAMA")
    add_bullet(doc, "Menjamin proses produksi dilaksanakan di sarana pabrik yang memenuhi standar Cara Pembuatan PKRT yang Baik (CPKRTB) dan ISO 9001:2015.", "2.1 ");
    add_bullet(doc, "Menyediakan Penanggung Jawab Teknis (PJT) yang berkualifikasi untuk mengawasi mutu, pengujian formulasi, dan menerbitkan Certificate of Analysis (CoA) per batch.", "2.2 ");
    add_bullet(doc, "Membantu dan memfasilitasi kelengkapan dokumen teknis sarana pabrik yang dibutuhkan dalam rangka pendaftaran Izin Edar KEMENKES RI PKD.", "2.3 ");
    add_bullet(doc, "Berhak menerima pembayaran biaya jasa makloon tepat waktu sesuai Purchase Order dan kesepakatan harga.", "2.4 ");

    add_heading_1(doc, "PASAL 3: HAK DAN KEWAJIBAN PIHAK KEDUA")
    add_bullet(doc, "Menjamin bahwa Merek Dagang yang dimakloonkan adalah sah miliknya dan tidak melanggar Hak Kekayaan Intelektual (HKI) pihak ketiga manapun.", "3.1 ");
    add_bullet(doc, "Bertanggung jawab penuh atas peredaran, pemasaran, penentuan harga jual konsumen, dan jaringan distribusi produk jadi di pasar.", "3.2 ");
    add_bullet(doc, "Wajib menerbitkan Surat Pesanan Resmi (Purchase Order / PO) sebelum proses produksi dimulai.", "3.3 ");
    add_bullet(doc, "Berhak menerima produk jadi yang memenuhi standar spesifikasi mutu dan lolos uji laboratorium dengan bukti sertifikat CoA.", "3.4 ");

    add_heading_1(doc, "PASAL 4: KEPEMILIKAN MEREK, HKI DAN KERAHASIAAN FORMULA (NDA)")
    add_bullet(doc, "Hak atas Merek Dagang, desain logo, dan hak cipta etiket kemasan sepenuhnya merupakan hak milik eksklusif PIHAK KEDUA.", "4.1 ");
    add_bullet(doc, "Hak atas metode manufaktur, teknologi proses pencampuran, dan formulasi kimia merupakan milik PIHAK PERTAMA kecuali disepakati lain secara tertulis.", "4.2 ");
    add_bullet(doc, "PARA PIHAK sepakat untuk menjaga kerahasiaan seluruh data teknis, formula, dokumen komersial, dan rahasia dagang (Non-Disclosure Agreement) selama perjanjian berlaku dan hingga 5 (lima) tahun setelahnya.", "4.3 ");

    add_heading_1(doc, "PASAL 5: PENGENDALIAN MUTU, GARANSI DAN PENARIKAN PRODUK (RECALL)")
    add_bullet(doc, "PIHAK PERTAMA menjamin produk yang diserahkan telah melalui uji QC dan memenuhi spesifikasi teknis Kemenkes RI.", "5.1 ");
    add_bullet(doc, "Apabila dalam waktu maksimal 14 (empat belas) hari kalender sejak penerimaan barang ditemukan cacat produksi massal yang diakibatkan oleh kelalaian proses manufaktur PIHAK PERTAMA, maka PIHAK PERTAMA wajib mengganti produk tersebut secara cuma-cuma.", "5.2 ");
    add_bullet(doc, "Apabila terjadi instruksi penarikan produk (recall) oleh Kementerian Kesehatan RI yang disebabkan oleh kegagalan formula/manufaktur, PIHAK PERTAMA dan PIHAK KEDUA wajib bekerja sama melakukan penarikan secara sistematis sesuai SOP Recall.", "5.3 ");

    add_heading_1(doc, "PASAL 6: JANGKA WAKTU DAN PENYELESAIAN PERSELISIHAN")
    add_bullet(doc, "Perjanjian ini berlaku selama 2 (dua) tahun terhitung sejak tanggal penandatanganan dan dapat diperpanjang secara otomatis atas kesepakatan tertulis PARA PIHAK.", "6.1 ");
    add_bullet(doc, "Segala perselisihan yang timbul akan diselesaikan secara musyawarah untuk mufakat. Apabila tidak tercapai mufakat, PARA PIHAK sepakat memilih domisili hukum di Kantor Kepaniteraan Pengadilan Negeri [Kota Setempat].", "6.2 ");

    add_p(doc, "Demikian Perjanjian ini dibuat dalam rangkap 2 (dua) asli bermaterai cukup (Rp 10.000,-), masing-masing pihak memegang 1 (satu) rangkap yang mempunyai kekuatan hukum yang sama.")

    add_signature_block(doc, "PIHAK PERTAMA,\nPT [NAMA PERUSAHAAN ANDA]\n\n\n( Materai Rp 10.000 )", "PIHAK KEDUA,\nPT [NAMA PERUSAHAAN MITRA]\n\n\n( Materai Rp 10.000 )", "[ Nama Direktur Anda ]\nDirektur Utama", "[ Nama Direktur Mitra ]\nDirektur Utama")
    
    file_path = os.path.join(OUTPUT_DIR, "04_SURAT_PERJANJIAN_KERJASAMA_MAKLOON_KONTRAK.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

# ==============================================================================
# GENERATE DOC 5: SOP & FORMULIR KENDALI MUTU BATCH RECORD COA
# ==============================================================================
def create_doc_5():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Kumpulan SOP Inti & Formulir Kendali Mutu Produksi PKRT Standar ISO 9001:2015 & CPKRTB", "SOP-PKRT-005", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "BAGIAN I: KUMPULAN STANDAR OPERASIONAL PROSEDUR (SOP)")
    
    add_heading_2(doc, "SOP-01: Prosedur Higiene Personalia, Sanitasi Ruangan & Penggunaan APD")
    add_bullet(doc, "Tujuan: Mencegah kontaminasi fisik dan mikrobiologi dari personil ke produk serta menjamin keselamatan kerja kimia (K3).", "1. ");
    add_bullet(doc, "Ruang Lingkup: Seluruh karyawan, teknisi, dan tamu yang memasuki area penimbangan, pengolahan, dan pengemasan.", "2. ");
    add_bullet(doc, "Prosedur Masuk Ruang Produksi: (a) Melepas sepatu luar di ruang ganti dan memakai sepatu boots/khusus pabrik; (b) Mengenakan pakaian kerja/jas lab, masker medis, hairnet, dan sarung tangan karet; (c) Mencuci tangan dengan sabun cair antiseptik selama minimal 20 detik pada wastafel alir sebelum memasuki ruang produksi.", "3. ");
    add_bullet(doc, "Larangan: Dilarang memakai perhiasan, makan, minum, merokok, atau bekerja dalam kondisi luka terbuka tanpa perban tahan air.", "4. ");

    add_heading_2(doc, "SOP-02: Prosedur Pengadaan, Penerimaan & Penyimpanan Bahan Kimia (FIFO/FEFO)")
    add_bullet(doc, "Penerimaan: Petugas gudang memeriksa keutuhan segel kemasan drum/jerigen, label bahan, nomor batch supplier, kesesuaian Surat Jalan, dan melampirkan Certificate of Analysis (CoA) dari pabrik pembuat.", "1. ");
    add_bullet(doc, "Penyimpanan: Bahan baku diletakkan di atas palet plastik/kayu (tidak bersentuhan langsung dengan lantai). Bahan korosif kuat (Sodium Hypochlorite / Bleach) diletakkan di area berventilasi khusus terpisah dari bahan parfum/pewarna.", "2. ");
    add_bullet(doc, "Sistem Pengeluaran: Menggunakan prinsip FIFO (First In First Out) atau FEFO (First Expired First Out) yang dicatat pada Kartu Stok Gudang.", "3. ");

    add_heading_2(doc, "SOP-03: Prosedur Penimbangan, Pengolahan & Pencampuran (Mixing)")
    add_bullet(doc, "Persiapan: Pastikan tangki mixing dan alat pengaduk dalam status bertanda label 'BERSIH - SIAP PAKAI'. Lakukan kalibrasi/tara pada timbangan digital.", "1. ");
    add_bullet(doc, "Penimbangan: Petugas menimbang bahan sesuai instruksi Catatan Pengolahan Bets (Batch Record). PJT / Supervisor memverifikasi kesesuaian bobot timbang.", "2. ");
    add_bullet(doc, "Proses Mixing: Masukkan air baku, surfaktan, zat penstabil, parfum, dan pengatur kekentalan sesuai urutan instruksi kerja produk hingga larutan homogen.", "3. ");
    add_bullet(doc, "Pengendalian Busa: Diamkan larutan selama 12-24 jam hingga gelembung udara hilang sempurna sebelum dilakukan proses pengisian ke dalam kemasan botol/jerigen.", "4. ");

    add_heading_2(doc, "SOP-04: Prosedur Pengawasan Mutu (QC), Pelulusan Bets & Penerbitan CoA")
    add_bullet(doc, "Pengambilan Sampel: Petugas QC mengambil sampel 250 mL dari tangki mixer setelah homogen dan dari 3 botol hasil kemasan awal, tengah, dan akhir proses filling.", "1. ");
    add_bullet(doc, "Pengujian: Lakukan uji organoleptik (warna, kejernihan, bau), uji pH menggunakan pH Meter Digital terkalibrasi, dan uji bobot jenis menggunakan piknometer.", "2. ");
    add_bullet(doc, "Keputusan Rilis: Jika seluruh parameter memenuhi rentang standar spesifikasi, PJT menandatangani status 'DILULUSKAN' dan menerbitkan Certificate of Analysis (CoA) resmi.", "3. ");
    add_bullet(doc, "Sampel Pertinggal (Retained Sample): Simpan minimal 2 botol kemasan jadi dari setiap batch di ruang arsip sampel selama minimal 2 tahun.", "4. ");

    add_heading_1(doc, "BAGIAN II: FORMULIR REKAMAN MUTU DAN TEMPLATE MASTER")
    
    add_heading_2(doc, "Formulir 01: Catatan Pengolahan Bets (Batch Processing Record) - Sabun Cuci Piring")
    
    headers_b = ["No", "Nama Bahan Baku", "Fungsi Kimia", "Formula Standar (1000 Kg)", "Hasil Timbang Aktual", "Paraf Penimbang"]
    data_b = [
        ["1", "Air Bersih / Aquadest", "Pelarut Utama", "820,0 Kg", ".......... Kg", "[  OK  ]"],
        ["2", "Sodium Lauryl Ether Sulfate (SLES 70%)", "Surfaktan Pembusa", "100,0 Kg", ".......... Kg", "[  OK  ]"],
        ["3", "Linear Alkylbenzene Sulfonate (LABSA)", "Degreaser / Pembersih Lemak", "40,0 Kg", ".......... Kg", "[  OK  ]"],
        ["4", "Sodium Chloride (NaCl Refined)", "Pengental (Viscosity Adjuster)", "25,0 Kg", ".......... Kg", "[  OK  ]"],
        ["5", "Fragrance Lime Oil (Jeruk Nipis)", "Aroma / Pengharum", "5,0 Kg", ".......... Kg", "[  OK  ]"],
        ["6", "DMDM Hydantoin", "Pengawet Antimikroba", "2,0 Kg", ".......... Kg", "[  OK  ]"],
        ["7", "CI 19140 & CI 42090 (Foodgrade)", "Pewarna Hijau", "0,5 Kg", ".......... Kg", "[  OK  ]"]
    ]
    add_styled_table(doc, headers_b, data_b, [Cm(0.8), Cm(4.8), Cm(4.2), Cm(2.8), Cm(2.2), Cm(1.8)])

    add_heading_2(doc, "Formulir 02: Lembar Pengujian Mutu Produk Jadi (QC Testing Sheet)")
    
    headers_qc = ["Parameter Uji", "Metode Pengujian", "Standar Spesifikasi Kemenkes", "Hasil Uji Aktual", "Status (Pass/Fail)"]
    data_qc = [
        ["Bentuk & Penampilan", "Organoleptik Visual", "Cairan Kental Jernih, Homogen", "Sesuai Spesifikasi", "LULUS"],
        ["Warna Sediaan", "Visual", "Hijau Segar Transparan", "Sesuai Spesifikasi", "LULUS"],
        ["Aroma / Bau", "Olfaktori", "Khas Segar Jeruk Nipis", "Sesuai Spesifikasi", "LULUS"],
        ["Derajat Keasaman (pH)", "pH Meter Digital", "6,00 – 8,00 pada suhu 25°C", "7,15", "LULUS"],
        ["Bobot Jenis (g/mL)", "Piknometer", "1,010 – 1,040 g/mL", "1,025 g/mL", "LULUS"],
        ["Uji Kebocoran Kemasan", "Vakum / Uji Tekan", "Tidak ada rembesan pada tutup botol", "Rapat / Kedap", "LULUS"]
    ]
    add_styled_table(doc, headers_qc, data_qc, [Cm(3.2), Cm(3.0), Cm(4.8), Cm(3.2), Cm(2.2)])

    add_heading_2(doc, "Formulir 03: Logbook Pemantauan Suhu & Kelembaban Gudang")
    
    headers_log = ["Tanggal", "Waktu Pengecekan", "Suhu Aktual (Standar: 15-30°C)", "Kelembaban RH (Standar: <70%)", "Paraf Petugas", "Status"]
    data_log = [
        ["19/08/2026", "08:00 WIB", "26,5 °C", "62 %", "[  Budi  ]", "Normal"],
        ["19/08/2026", "16:00 WIB", "28,2 °C", "65 %", "[  Budi  ]", "Normal"],
        ["20/08/2026", "08:00 WIB", "26,0 °C", "60 %", "[  Budi  ]", "Normal"],
        ["20/08/2026", "16:00 WIB", "28,0 °C", "64 %", "[  Budi  ]", "Normal"]
    ]
    add_styled_table(doc, headers_log, data_log, [Cm(2.2), Cm(2.8), Cm(4.2), Cm(4.0), Cm(2.0), Cm(1.8)])

    add_signature_block(doc)
    
    file_path = os.path.join(OUTPUT_DIR, "05_SOP_DAN_FORMULIR_KENDALI_MUTU_BATCH_RECORD_COA.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

if __name__ == "__main__":
    create_doc_1()
    create_doc_2()
    create_doc_3()
    create_doc_4()
    create_doc_5()
    print("ALL 5 ISO 9001 PKRT DOCUMENTS SUCCESSFULLY GENERATED!")

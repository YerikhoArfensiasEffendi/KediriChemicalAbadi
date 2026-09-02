import os
import openpyxl
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from kca_doc_humanizer import sanitize_docx_metadata

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_proposal():
    base_dir = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical"
    output_dir = os.path.join(base_dir, "Keuangan/Proposal")
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, "PROPOSAL_KERJASAMA_MAKLOON_CAPEX_200JT_KCA.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # ---------------- KOP SURAT RESMI ----------------
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k1 = kop.add_run("PT KEDIRI CHEMICAL ABADI\n")
    k1.font.name = "Arial"
    k1.font.size = Pt(16)
    k1.font.bold = True
    k1.font.color.rgb = RGBColor(0, 0, 0)

    k2 = kop.add_run("PRODUSEN BAHAN KIMIA PEMBERSIH, LAUNDRY & MAKLOON INDUSTRI\n")
    k2.font.name = "Arial"
    k2.font.size = Pt(10)
    k2.font.bold = True
    k2.font.color.rgb = RGBColor(0, 0, 0)

    k3 = kop.add_run("Pabrik: RT.1/RW.6, Pagung, Kec. Semen, Kabupaten Kediri, Jawa Timur 64161\n"
                    "Hotline/WA: 0822-4400-6699 | Email: kdrchemicals@gmail.com\n")
    k3.font.name = "Arial"
    k3.font.size = Pt(9)
    k3.font.color.rgb = RGBColor(0, 0, 0)

    p_line = doc.add_paragraph("═" * 58)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_line.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- JUDUL PROPOSAL ----------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title_p.add_run("PROPOSAL PENAWARAN KERJASAMA STRATEGIS\n")
    t1.font.name = "Arial"
    t1.font.size = Pt(13.5)
    t1.font.bold = True
    t1.font.color.rgb = RGBColor(0, 0, 0)

    t2 = title_p.add_run("SKEMA PEMBIAYAAN MESIN PRODUKSI DEDIKASI (DEDICATED CAPEX OFFSET)\n"
                        "DENGAN PENGEMBALIAN AMORTISASI POTONGAN INVOICE 10%\n")
    t2.font.name = "Arial"
    t2.font.size = Pt(11)
    t2.font.bold = True
    t2.font.color.rgb = RGBColor(0, 0, 0)

    t3 = title_p.add_run("Nomor Dokumen: 001/PRP-KCA/CAPEX-MKL/IX/2026 | Nilai Investasi: Rp 200.000.000,-\n")
    t3.font.name = "Arial"
    t3.font.size = Pt(9.5)
    t3.font.italic = True
    t3.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 1 ----------------
    h1 = doc.add_heading("1. LATAR BELAKANG & TUJUAN KEMITRAAN", level=2)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    p1 = doc.add_paragraph(
        "Kebutuhan pasar terhadap produk bahan kimia pembersih industri, commercial laundry, dan hospitality care "
        "(hotel, resto, spa) terus meningkat pesat seiring dengan tuntutan standar higienitas dan efisiensi operasional. "
        "PT Kediri Chemical Abadi hadir sebagai mitra manufaktur yang siap memproduksi produk berkualitas tinggi dengan formulasi "
        "presisi, stabilitas mutu terjamin, dan kapasitas produksi yang dapat diandalkan.\n\n"
        "Melalui proposal ini, PT Kediri Chemical Abadi menawarkan kemitraan eksklusif kepada Mitra Pemodal / Distributor "
        "melalui skema Dedicated CapEx Offset Financing. Skema ini dirancang khusus untuk menciptakan kolaborasi saling menguntungkan "
        "(Win-Win Partnership), di mana Mitra mendanai penambahan mesin produksi khusus, dan modal tersebut dikembalikan penuh secara "
        "otomatis melalui potongan 10% pada setiap transaksi pesanan."
    )
    p1.paragraph_format.line_spacing = 1.15
    for r in p1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 2 ----------------
    h2 = doc.add_heading("2. RINCIAN ALOKASI PENGADAAN MESIN (INVESTASI RP 200.000.000,-)", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    p2 = doc.add_paragraph(
        "Dana pemodalan sebesar Rp 200.000.000,- (Dua Ratus Juta Rupiah) akan dialokasikan 100% secara transparan untuk pengadaan unit "
        "mesin manufaktur modern yang didedikasikan secara prioritas untuk memproduksi lini pesanan Mitra:"
    )
    p2.paragraph_format.line_spacing = 1.15
    for r in p2.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    t_mesin = doc.add_table(rows=1, cols=4)
    t_mesin.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_mesin.rows[0].cells
    m_hdr[0].text = "No"
    m_hdr[1].text = "Uraian Mesin & Spesifikasi Teknis"
    m_hdr[2].text = "Kapasitas / Fungsi"
    m_hdr[3].text = "Estimasi Biaya"
    for c in m_hdr:
        set_cell_shading(c, "1E293B")
        set_cell_margins(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    mesin_items = [
        ("1", "Mesin Mixer Homogenizer SUS316 Double Jacket\n(High-shear mixing, inverter speed control, heater tank)", "1.000 - 2.000 Liter / Batch\n(Sabun, Emulsifier, Detergen)", "Rp 115.000.000"),
        ("2", "Semi-Automatic Liquid Filling Machine 4-Nozzle\n(Pneumatic piston filler presisi 1L s/d 20L + Anti-drip)", "Kapasitas 600 - 800 botol/jam\n(Presisi & Cepat)", "Rp 45.000.000"),
        ("3", "Mesin Continuous Induction Sealer & Capping\n(Penyegel aluminium foil anti-bocor otomatis)", "Penyegelan tutup jerigen & botol", "Rp 25.000.000"),
        ("4", "Sistem Pompa Transfer Kimia, Filter Presisi & Piping\n(Sanitary grade chemical transfer system)", "Sirkulasi bahan baku ke tangki", "Rp 15.000.000"),
        ("", "TOTAL ALOKASI INVESTASI MESIN DEDIKASI", "", "Rp 200.000.000")
    ]
    for no, uraian, kap, nom in mesin_items:
        row = t_mesin.add_row().cells
        row[0].text = no
        row[1].text = uraian
        row[2].text = kap
        row[3].text = nom
        for c in row:
            set_cell_margins(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)
        if "TOTAL" in uraian:
            set_cell_shading(row[0], "F1F5F9")
            set_cell_shading(row[1], "F1F5F9")
            set_cell_shading(row[2], "F1F5F9")
            set_cell_shading(row[3], "F1F5F9")
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 3 ----------------
    h3 = doc.add_heading("3. MEKANISME KERJA & SKEMA PENGEMBALIAN (POTONGAN 10%)", level=2)
    for r in h3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    p3 = doc.add_paragraph(
        "Mekanisme pengembalian investasi dilakukan secara terintegrasi dan otomatis tanpa membebani arus kas kedua belah pihak:\n\n"
        "1. Injeksi Modal: Mitra menyetorkan dana investasi mesin sebesar Rp 200.000.000,- yang dicatat sebagai Titipan Modal Mesin.\n"
        "2. Pengadaan & Instalasi: PT Kediri Chemical Abadi melakukan pengadaan, uji coba (commissioning), dan kalibrasi mesin di pabrik Pagung, Kediri.\n"
        "3. Pemesanan Parsial: Mitra melakukan pesanan berkala (repeat order) sesuai kebutuhan pasar.\n"
        "4. Potongan Faktur 10%: Setiap penerbitan invoice tagihan pesanan, total nilai barang diberikan potongan otomatis 10% untuk melunasi saldo modal mesin.\n"
        "5. Pelunasan Penuh: Setelah akumulasi potongan mencapai Rp 200.000.000,- (setara akumulasi pesanan Rp 2.000.000.000,-), modal mesin dinyatakan Lunas 100%, kepemilikan mesin beralih menjadi aset penuh PT Kediri Chemical Abadi, dan harga pesanan berikutnya kembali normal 100%."
    )
    p3.paragraph_format.line_spacing = 1.15
    for r in p3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Tabel Simulasi
    t_sim = doc.add_table(rows=1, cols=5)
    t_sim.alignment = WD_TABLE_ALIGNMENT.CENTER
    s_hdr = t_sim.rows[0].cells
    s_hdr[0].text = "Transaksi Order"
    s_hdr[1].text = "Nilai Tagihan Bruto"
    s_hdr[2].text = "Kas Dibayar (90%)"
    s_hdr[3].text = "Potongan Modal (10%)"
    s_hdr[4].text = "Sisa Modal Mesin"
    for c in s_hdr:
        set_cell_shading(c, "1E293B")
        set_cell_margins(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    sim_rows = [
        ("Kondisi Awal", "-", "-", "-", "Rp 200.000.000"),
        ("Pesanan Batch 1", "Rp 50.000.000", "Rp 45.000.000", "Rp 5.000.000", "Rp 195.000.000"),
        ("Pesanan Batch 2", "Rp 100.000.000", "Rp 90.000.000", "Rp 10.000.000", "Rp 185.000.000"),
        ("Pesanan Batch 3", "Rp 150.000.000", "Rp 135.000.000", "Rp 15.000.000", "Rp 170.000.000"),
        ("Akumulasi s/d Selesai", "Rp 2.000.000.000", "Rp 1.800.000.000", "Rp 200.000.000", "Rp 0 (LUNAS 100%)")
    ]
    for trx, bruto, kas, pot, sisa in sim_rows:
        row = t_sim.add_row().cells
        row[0].text = trx
        row[1].text = bruto
        row[2].text = kas
        row[3].text = pot
        row[4].text = sisa
        for c in row:
            set_cell_margins(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)
        if "LUNAS" in sisa:
            set_cell_shading(row[0], "E2E8F0")
            set_cell_shading(row[1], "E2E8F0")
            set_cell_shading(row[2], "E2E8F0")
            set_cell_shading(row[3], "E2E8F0")
            set_cell_shading(row[4], "E2E8F0")
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 4 ----------------
    h4 = doc.add_heading("4. KEUNTUNGAN STRATEGIS BAGI MITRA PEMODAL / DISTRIBUTOR", level=2)
    for r in h4.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    p4 = doc.add_paragraph(
        "1. Bebas Beban Operasional Pabrik: Mitra tidak perlu menyewa lahan pabrik, merekrut staf produksi, memikirkan limbah kimia, atau mengurus legalitas manufaktur.\n"
        "2. Pengembalian Modal Otomatis (Auto-ROI): Dana Rp 200 Juta kembali secara pasti di setiap perputaran order tanpa risiko piutang macet.\n"
        "3. Margin Bersaing & Diskon 10%: Mendapatkan keuntungan komersial instan 10% lebih hemat selama masa amortisasi mesin.\n"
        "4. Kapasitas & Prioritas Produksi: Pesanan Mitra mendapatkan antrean utama (fast-track) di lini mesin dedikasi tersebut.\n"
        "5. Fleksibilitas Kemasan & Formula: Mitra bebas menentukan spesifikasi aroma, warna, dan jenis kemasan (1L s/d 20L Jerigen, hingga 10kg Pail)."
    )
    p4.paragraph_format.line_spacing = 1.15
    for r in p4.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 5 ----------------
    h5 = doc.add_heading("5. KATALOG PRODUK UNGGULAN YANG DAPAT DIPRODUKSI", level=2)
    for r in h5.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    
    t_prod = doc.add_table(rows=1, cols=3)
    t_prod.alignment = WD_TABLE_ALIGNMENT.CENTER
    p_hdr = t_prod.rows[0].cells
    p_hdr[0].text = "Kategori Produk"
    p_hdr[1].text = "Nama Produk & Spesifikasi"
    p_hdr[2].text = "Pilihan Kemasan"
    for c in p_hdr:
        set_cell_shading(c, "1E293B")
        set_cell_margins(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    prod_items = [
        ("Laundry Care", "Liquid Detergent Matic (Formula Busa Rendah, Anti Bau Apek)", "1 Liter, 5 Liter, 20 Liter Jerigen"),
        ("Kitchen Care", "Liquid DishWasher (Ekstrak Jeruk Nipis, Daya Angkat Lemak Tinggi)", "1 Liter, 5 Liter, 20 Liter Jerigen"),
        ("Specialty Laundry", "Pencerah Warna (Oxygen Bleach Ramah Serat Kain)", "1 Liter, 5 Liter, 20 Liter Jerigen"),
        ("Specialty Laundry", "Emulsifier (Pengangkat Lemak & Minyak Berat)", "1 Liter, 5 Liter, 20 Liter Jerigen"),
        ("Specialty Laundry", "Bleach Klorin (Pemutih & Disinfektan Sanitasi)", "1 Liter, 5 Liter, 20 Liter Jerigen"),
        ("Spotting Agent", "Rust Tex (Formula Khusus Perontok Karat Tekstil)", "1 Liter, 5 Liter, 20 Liter Jerigen"),
        ("Spotting Agent", "Blood Tex (Formula Enzimatik Noda Darah & Protein)", "1 Liter, 5 Liter, 20 Liter Jerigen"),
        ("Spa Specialty", "Peluruh Noda Minyak SPA (Khusus Massage Oil Towel)", "10 kg Pail Khusus")
    ]
    for kat, nama, kem in prod_items:
        row = t_prod.add_row().cells
        row[0].text = kat
        row[1].text = nama
        row[2].text = kem
        for c in row:
            set_cell_margins(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- LEMBAR PENGESAHAN ----------------
    doc.add_paragraph("\n")
    h6 = doc.add_heading("6. LEMBAR PERSETUJUAN & KESEPAKATAN KERJASAMA", level=2)
    for r in h6.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True
    p_penutup = doc.add_paragraph(
        "Demikian proposal penawaran kerjasama ini kami sampaikan dengan penuh itikad baik dan profesionalisme. "
        "Besar harapan kami untuk dapat menjalin kemitraan jangka panjang yang kokoh dan saling menguntungkan bersama Anda."
    )
    p_penutup.paragraph_format.line_spacing = 1.15
    for r in p_penutup.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("\n")
    t_sign = doc.add_table(rows=3, cols=2)
    t_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t_sign.rows[0].cells[0].text = "PIHAK PERTAMA (PENGELOLA PABRIK)\nPT KEDIRI CHEMICAL ABADI"
    t_sign.rows[0].cells[1].text = "PIHAK KEDUA (MITRA PEMODAL / DISTRIBUTOR)\n"
    
    t_sign.rows[1].cells[0].text = "\n\n\n\n"
    t_sign.rows[1].cells[1].text = "\n\n\n\n"

    t_sign.rows[2].cells[0].text = "YAN EFFENDI / YERIKHO ARFENSIAS E.\nDirektur Utama / General Manager"
    t_sign.rows[2].cells[1].text = "(......................................................)\nNama & Jabatan"

    for r in t_sign.rows:
        for c in r.cells:
            set_cell_margins(c)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_run in p.runs:
                    r_run.font.bold = True
                    r_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(doc_path)
    
    # Duplikat ke folder legalitas
    legal_dir = os.path.join(base_dir, "KCA DOKUMEN/06_MASTER_MANUAL_DAN_LEGALITAS")
    os.makedirs(legal_dir, exist_ok=True)
    legal_path = os.path.join(legal_dir, "PROPOSAL_KERJASAMA_MAKLOON_CAPEX_200JT_KCA.docx")
    doc.save(legal_path)
    
    # Sanitasi Metadata
    sanitize_docx_metadata(doc_path, title="Proposal Kerjasama Makloon & Pemodalan Mesin 200Jt", author="Yerikho Arfensias Effendi")
    sanitize_docx_metadata(legal_path, title="Proposal Kerjasama Makloon & Pemodalan Mesin 200Jt", author="Yerikho Arfensias Effendi")
    
    print(f"SUCCESS_UPDATED: Proposal ISO 9001 (Pure Black Font & Sanitized Metadata) generated at {doc_path}")
    return doc_path

if __name__ == "__main__":
    generate_proposal()

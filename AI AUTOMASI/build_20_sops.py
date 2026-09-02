import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

OUTPUT_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/dokumen_iso9001_pkrt"
os.makedirs(OUTPUT_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PT [NAMA PERUSAHAAN]\nMANUFAKTUR PKRT\nSISTEM ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block(doc, left_title="Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", right_title="Disahkan Oleh,\nDirektur Utama", left_name="( ............................................ )", right_name="( ............................................ )"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    col_widths = [Cm(7.5), Cm(8.0)]
    
    cell_l = table.cell(0, 0)
    cell_l.width = col_widths[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l1 = p_l.add_run(left_title + "\n\n\n\n\n" + left_name + "\nTanggal: ...................................")
    r_l1.font.name = 'Times New Roman'
    r_l1.font.size = Pt(9.5)
    r_l1.font.color.rgb = COLOR_BLACK
    
    cell_r = table.cell(0, 1)
    cell_r.width = col_widths[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r.add_run(right_title + "\n\n\n\n\n" + right_name + "\nTanggal: ...................................")
    r_r1.font.name = 'Times New Roman'
    r_r1.font.size = Pt(9.5)
    r_r1.font.color.rgb = COLOR_BLACK

def build_20_sops():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Buku Kompendium Lengkap 20 Standar Operasional Prosedur (SOP) CPKRTB & ISO 9001:2015", "SOP-KOM-013", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "PENDAHULUAN")
    add_p(doc, "Buku Kompendium ini berisi 20 (dua puluh) Standar Operasional Prosedur (SOP) resmi yang mengatur seluruh aktivitas operasional pabrik manufaktur PKRT PT [Nama Perusahaan], mulai dari fasilitas bangunan, penerimaan gudang, pengolahan kimia, pengawasan mutu laboratorium, hingga pasca-penjualan dan inspeksi audit internal.")

    # -------------------------------------------------------------------------
    # KELOMPOK A: BANGUNAN, FASILITAS & SANITASI (SOP 01 - 04)
    # -------------------------------------------------------------------------
    add_heading_1(doc, "BAGIAN I: PROSEDUR BANGUNAN, FASILITAS & SANITASI")
    
    add_heading_2(doc, "SOP-SAN-001: Prosedur Pembersihan & Sanitasi Ruangan Pabrik")
    add_bullet(doc, "Tujuan: Memastikan seluruh area kerja bebas dari debu, kotoran, dan sisa bahan kimia agar tidak terjadi kontaminasi silang.", "1. ");
    add_bullet(doc, "Penanggung Jawab: Petugas Kebersihan / Sanitasi di bawah pengawasan Kepala Produksi.", "2. ");
    add_bullet(doc, "Prosedur Harian: (a) Sapu dan pel lantai ruang mixing dan filling setiap pagi sebelum produksi dan sore setelah produksi menggunakan larutan disinfektan (Benzalkonium Chloride 0.1% atau Karbol); (b) Bersihkan permukaan meja kerja timbang dan meja filling dengan lap basah bersih berdisinfektan; (c) Buang seluruh sampah padat di tempat penampungan luar pabrik.", "3. ");
    add_bullet(doc, "Prosedur Mingguan: Bersihkan dinding keramik, daun pintu, kaca jendela, dan bersihkan kisi-kisi exhaust fan dari debu yang menempel.", "4. ");
    add_bullet(doc, "Pencatatan: Isi tanggal dan paraf pada Logbook Pembersihan Ruangan.", "5. ");

    add_heading_2(doc, "SOP-EQP-002: Prosedur Pembersihan & Perawatan Mesin/Tangki Mixer")
    add_bullet(doc, "Tujuan: Mencegah percampuran sisa formula produk batch sebelumnya dengan batch baru.", "1. ");
    add_bullet(doc, "Prosedur: (a) Kuras habis sisa cairan di dalam tangki mixer; (b) Bilas dinding tangki dengan air bersih bertekanan; (c) Gosok bagian dalam tangki menggunakan spons halus dan deterjen netral; (d) Bilas ulang dengan air bersih hingga air buangan jernih dan bebas busa; (e) Keringkan bagian dalam tangki dan pasang label 'BERSIH - SIAP PAKAI'.", "2. ");

    add_heading_2(doc, "SOP-HIG-003: Prosedur Higiene Personil & Pemakaian APD")
    add_bullet(doc, "Tujuan: Melindungi pekerja dari paparan uap kimia korosif serta mencegah kontaminasi dari personil ke produk.", "1. ");
    add_bullet(doc, "Prosedur: (a) Personil wajib mencuci tangan dengan sabun antiseptik selama minimal 20 detik; (b) Wajib mengenakan pakaian kerja/jas lab, masker medis/respirator, tutup kepala (hairnet), sarung tangan nitril/karet, dan sepatu tertutup/boots; (c) Dilarang mengenakan cincin/perhiasan dan dilarang makan/minum/merokok di area produksi.", "2. ");

    add_heading_2(doc, "SOP-PST-004: Prosedur Pengendalian Hama (Pest Control)")
    add_bullet(doc, "Tujuan: Mencegah masuknya serangga, lalat, kecoa, dan hewan pengerat (tikus) ke dalam area pabrik.", "1. ");
    add_bullet(doc, "Prosedur: (a) Pasang insect fly catcher elektrik pada area pintu masuk; (b) Pasang perangkap tikus mekanis (glue trap/box trap) di sepanjang dinding luar dan sudut gudang; (c) Periksa perangkap setiap hari Senin pagi, bersihkan dan ganti lem perekat bila kotor; (d) Catat hasil temuan hama pada Logbook Pest Control.", "2. ");

    # -------------------------------------------------------------------------
    # KELOMPOK B: GUDANG & PENGADAAN BAHAN (SOP 05 - 08)
    # -------------------------------------------------------------------------
    add_heading_1(doc, "BAGIAN II: PROSEDUR GUDANG & PENGADAAN BAHAN")
    
    add_heading_2(doc, "SOP-GUD-005: Prosedur Penerimaan & Pemeriksaan Bahan Masuk")
    add_bullet(doc, "Tujuan: Memastikan bahan baku dan bahan kemas yang diterima sesuai spesifikasi pesanan dan bermutu baik.", "1. ");
    add_bullet(doc, "Prosedur: (a) Periksa surat jalan dan kecocokan Certificate of Analysis (CoA) dari supplier; (b) Periksa kondisi fisik drum/jerigen (tidak bocor, tidak berkarat, segel utuh); (c) Pasang label kuning 'DIKARANTINA' dan laporkan ke QC untuk pengambilan sampel uji; (d) Setelah QC meluluskan, ganti label menjadi 'DILULUSKAN' (hijau) dan catat di Kartu Stok.", "2. ");

    add_heading_2(doc, "SOP-GUD-006: Prosedur Penyimpanan Bahan Kimia Berbahaya/Korosif")
    add_bullet(doc, "Tujuan: Menjamin penyimpanan bahan kimia keras (Sodium Hypochlorite, LABSA, Asam/Basa Kuat) aman dan tidak merusak lingkungan.", "1. ");
    add_bullet(doc, "Prosedur: (a) Simpan bahan korosif di area gudang berventilasi khusus terpisah dari bahan parfum/pewarna; (b) Letakkan wadah di atas palet plastik dengan wadah penampung tumpahan (spill pallet); (c) Sediakan pasir/serbuk gergaji dan eye washer darurat di dekat area penyimpanan korosif.", "2. ");

    add_heading_2(doc, "SOP-GUD-007: Prosedur Pengeluaran Bahan Metode FIFO / FEFO")
    add_bullet(doc, "Tujuan: Menghindari pengendapan bahan baku yang kedaluwarsa atau rusak akibat penyimpanan terlalu lama.", "1. ");
    add_bullet(doc, "Prosedur: (a) Bahan baku yang masuk lebih awal (FIFO) atau yang tanggal kedaluwarsanya lebih dekat (FEFO) wajib dikeluarkan terlebih dahulu; (b) Tempel label tanggal masuk dan tanggal expired pada setiap drum bahan; (c) Petugas gudang memotong saldo pada Kartu Stok saat bahan diserahkan ke bagian produksi.", "2. ");

    add_heading_2(doc, "SOP-GUD-008: Prosedur Penyimpanan & Pengiriman Produk Jadi")
    add_bullet(doc, "Tujuan: Menjaga integritas kemasan produk jadi sebelum dan selama proses distribusi ke pelanggan/distributor.", "1. ");
    add_bullet(doc, "Prosedur: (a) Simpan karton boks produk jadi di atas palet kayu/plastik dengan tumpukan maksimal 5 tingkat karton; (b) Gudang barang jadi wajib terlindung dari sinar matahari langsung; (c) Saat pengiriman, terbitkan Surat Jalan resmi yang mencantumkan nama produk, nomor batch, jumlah, dan lampiran CoA dari PJT.", "2. ");

    # -------------------------------------------------------------------------
    # KELOMPOK C: OPERASIONAL PRODUKSI & PENGEMASAN (SOP 09 - 12)
    # -------------------------------------------------------------------------
    add_heading_1(doc, "BAGIAN III: PROSEDUR PRODUKSI & PENGEMASAN")
    
    add_heading_2(doc, "SOP-PRD-009: Prosedur Penimbangan Bahan Baku")
    add_bullet(doc, "Tujuan: Menjamin ketepatan takaran bahan kimia sesuai Master Formula yang disahkan Kemenkes.", "1. ");
    add_bullet(doc, "Prosedur: (a) Lakukan kalibrasi harian timbangan digital; (b) Timbang bahan baku satu per satu menggunakan wadah penimbang yang bersih dan kering; (c) Tempel label 'SUDAH DITIMBANG' memuat nama bahan, berat aktual, dan nomor batch target; (d) Catat hasil timbang aktual pada Catatan Pengolahan Bets (Batch Record).", "2. ");

    add_heading_2(doc, "SOP-PRD-010: Prosedur Proses Pencampuran (Mixing) Sabun & Kimia Pembersih")
    add_bullet(doc, "Tujuan: Menghasilkan larutan pembersih yang homogen, stabil, dan memenuhi spesifikasi viskositas.", "1. ");
    add_bullet(doc, "Prosedur: (a) Masukkan air baku ke dalam tangki mixer; (b) Masukkan surfaktan (SLES/LABSA) dan nyalakan motor pengaduk pada kecepatan rendah untuk mencegah pembentukan busa berlebih; (c) Larutkan pengawet dan pewarna secara terpisah sebelum dimasukkan ke tangki; (d) Masukkan parfum dan larutan pengatur kekentalan (NaCl) perlahan hingga viskositas tercapai; (e) Diamkan larutan selama 12–24 jam untuk degasifikasi gelembung udara.", "2. ");

    add_heading_2(doc, "SOP-PRD-011: Prosedur Pengisian (Filling) & Penutupan Kemasan")
    add_bullet(doc, "Tujuan: Memastikan volume isi botol tepat dan penutupan kemasan kedap serta anti-bocor.", "1. ");
    add_bullet(doc, "Prosedur: (a) Kalibrasi nozel mesin pengisi sesuai volume target (misal 450 mL); (b) Isi cairan ke dalam botol bersih; (c) Kencangkan tutup botol dengan torsi penutupan yang rapat; (d) Lakukan uji tekan/balik pada 5 botol sampel per batch untuk memastikan tidak ada rembesan cairan.", "2. ");

    add_heading_2(doc, "SOP-PRD-012: Prosedur Pelabelan Kemasan & Pengemasan ke Karton")
    add_bullet(doc, "Tujuan: Menjamin penempelan label stiker rapi, memuat data legalitas valid, dan terlindung dalam karton.", "1. ");
    add_bullet(doc, "Prosedur: (a) Periksa stiker label: pastikan tercantum Nomor PKD Kemenkes, nama pabrik/distributor, komposisi bahan, nomor batch, dan tanggal kedaluwarsa; (b) Tempel stiker pada posisi tengah botol secara simetris tanpa kerutan; (c) Susun botol ke dalam karton boks sekunder, segel dengan lakban rapat, dan beri tanda arah panah 'THIS SIDE UP'.", "2. ");

    # -------------------------------------------------------------------------
    # KELOMPOK D: QUALITY CONTROL & LABORATORIUM (SOP 13 - 16)
    # -------------------------------------------------------------------------
    add_heading_1(doc, "BAGIAN IV: PROSEDUR QUALITY CONTROL (QC) & LAB")
    
    add_heading_2(doc, "SOP-QC-013: Prosedur Pengambilan Sampel Uji (Sampling Protocol)")
    add_bullet(doc, "Tujuan: Memperoleh sampel uji yang representatif dari bahan baku dan produk jadi.", "1. ");
    add_bullet(doc, "Prosedur: (a) Gunakan wadah sampling bersih dan pipet steril; (b) Ambil sampel bahan baku dari 10% jumlah drum yang datang; (c) Ambil sampel in-process 250 mL dari tangki mixing; (d) Ambil 3 botol sampel produk jadi pada awal, tengah, dan akhir proses pengemasan.", "2. ");

    add_heading_2(doc, "SOP-QC-014: Prosedur Pengujian Kualitas Produk Jadi di Laboratorium")
    add_bullet(doc, "Tujuan: Menentukan status kelulusan mutu produk berdasarkan spesifikasi teknis Kemenkes.", "1. ");
    add_bullet(doc, "Prosedur: (a) Uji Organoleptik: amati warna, kejernihan, dan bau sediaan; (b) Uji pH: celupkan elektroda pH meter terkalibrasi ke dalam sampel suhu 25°C, catat nilai pH (standar sabun: 6.00–8.00; pemutih: 11.00–12.50); (c) Uji Bobot Jenis: timbang sampel dalam piknometer pada suhu 25°C; (d) Catat seluruh hasil pada Lembar Pengujian QC.", "2. ");

    add_heading_2(doc, "SOP-QC-015: Prosedur Pengelolaan Sampel Pertinggal (Retained Sample)")
    add_bullet(doc, "Tujuan: Menyimpan arsip fisik produk dari setiap nomor batch untuk kebutuhan investigasi jika terjadi keluhan komsumen.", "1. ");
    add_bullet(doc, "Prosedur: (a) Ambil 2 kemasan jadi lengkap dengan label resmi dari setiap batch yang diproduksi; (b) Simpan di lemari Sampel Pertinggal pada suhu kamar terkontrol; (c) Simpan selama minimal masa kedaluwarsa ditambah 1 tahun (total minimal 3 tahun); (d) Musnahkan sampel yang telah habis masa retensinya dan buat Berita Acara Pemusnahan.", "2. ");

    add_heading_2(doc, "SOP-QC-016: Prosedur Penerbitan Sertifikat Analisis (CoA)")
    add_bullet(doc, "Tujuan: Menerbitkan dokumen sertifikat resmi jaminan mutu produk yang disahkan oleh PJT.", "1. ");
    add_bullet(doc, "Prosedur: (a) Petugas QC menyerahkan Lembar Hasil Uji QC yang telah lolos spesifikasi kepada PJT; (b) PJT memeriksa dan memverifikasi data hasil uji; (c) PJT menandatangani blanko Certificate of Analysis (CoA) resmi; (d) Berikan salinan CoA kepada bagian ekspedisi untuk diserahkan kepada pembeli/distributor.", "2. ");

    # -------------------------------------------------------------------------
    # KELOMPOK E: PASCA-PRODUKSI, RECALL & AUDIT MUTU (SOP 17 - 20)
    # -------------------------------------------------------------------------
    add_heading_1(doc, "BAGIAN V: PROSEDUR PASCA-PRODUKSI, RECALL & AUDIT MUTU")
    
    add_heading_2(doc, "SOP-MKT-017: Prosedur Penanganan Keluhan Pelanggan (Complaint Handling)")
    add_bullet(doc, "Tujuan: Merespon, menginvestigasi, dan menyelesaikan setiap keluhan mutu produk dari pelanggan secara cepat dan tuntas.", "1. ");
    add_bullet(doc, "Prosedur: (a) Catat keluhan pada Formulir Keluhan Pelanggan dalam waktu 1x24 jam; (b) PJT mengambil sampel pertinggal dari batch yang dikeluhkan untuk diuji ulang di lab; (c) Tentukan penyebab masalah (kesalahan produksi, penyimpanan di toko, atau kemasan rusak saat ekspedisi); (d) Berikan surat tanggapan resmi dan lakukan penggantian barang jika terbukti kesalahan produksi.", "2. ");

    add_heading_2(doc, "SOP-REG-018: Prosedur Penarikan Kembali Produk dari Peredaran (Product Recall)")
    add_bullet(doc, "Tujuan: Melakukan penarikan produk secara efektif dari pasar apabila ditemukan cacat mutu berbahaya atau atas perintah Kemenkes RI.", "1. ");
    add_bullet(doc, "Prosedur: (a) Direktur membentuk Tim Penarikan Produk yang diketuai oleh PJT; (b) Telusuri data distribusi batch terkait dari arsip Surat Jalan; (c) Terbitkan Surat Pemberitahuan Penarikan Produk (Recall) kepada seluruh distributor/toko dalam waktu 2x24 jam; (d) Karantina seluruh produk yang ditarik di area bertanda 'BARANG DITOLAK/RECALL'; (e) Kirim laporan resmi pelaksanaan recall ke Direktorat Pengawasan PKRT Kemenkes RI.", "2. ");

    add_heading_2(doc, "SOP-PRD-019: Prosedur Penanganan Produk Tidak Sesuai (Reject / Rework)")
    add_bullet(doc, "Tujuan: Mengendalikan produk atau bahan baku yang tidak memenuhi spesifikasi agar tidak terkirim ke pasar.", "1. ");
    add_bullet(doc, "Prosedur: (a) Pasang label merah 'DITOLAK / REJECT' pada wadah atau palet yang tidak lolos QC; (b) PJT melakukan evaluasi: apakah cairan bisa dilakukan pengerjaan ulang (rework) dengan penyesuaian pH/viskositas, atau harus dimusnahkan; (c) Jika di-rework, buat Catatan Pengolahan Rework tersendiri; (d) Jika dimusnahkan, netralkan cairan sebelum dibuang ke IPAL dan buat Berita Acara Pemusnahan.", "2. ");

    add_heading_2(doc, "SOP-QMS-020: Prosedur Audit Internal & Inspeksi Diri Standar ISO 9001 & CPKRTB")
    add_bullet(doc, "Tujuan: Memastikan seluruh sistem mutu dan prosedur operasional pabrik berjalan konsisten dan siap menghadapi audit Kemenkes/Dinkes.", "1. ");
    add_bullet(doc, "Prosedur: (a) Audit internal dilaksanakan minimal 2 kali dalam setahun (setiap 6 bulan); (b) Tim Auditor Internal memeriksa seluruh area (Gudang, Produksi, QC, Dokumen PJT, Sanitasi); (c) Setiap temuan ketidaksesuaian dicatat pada Formulir Laporan CAPA; (d) Lakukan verifikasi tindakan perbaikan maksimal 30 hari kalender setelah audit; (e) Laporkan hasil audit dalam Rapat Tinjauan Manajemen (Management Review).", "2. ");

    add_signature_block(doc, "Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", "Disahkan Oleh,\nDirektur Utama PT [Nama Perusahaan]")
    
    file_path = os.path.join(OUTPUT_DIR, "13_BUKU_KOMPENDIUM_LENGKAP_20_SOP_CPKRTB_ISO9001.docx")
    doc.save(file_path)
    print(f"Generated: {file_path}")

if __name__ == "__main__":
    build_20_sops()

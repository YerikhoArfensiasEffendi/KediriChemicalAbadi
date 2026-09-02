import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

BASE_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/KCA DOKUMEN"

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PT KCA CHEMICAL\nMANUFAKTUR PKRT\nSISTEM ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block_2(doc, left_title="Disiapkan Oleh,\nPenanggung Jawab Teknis (PJT)", right_title="Disahkan Oleh,\nDirektur Utama", left_name="( ............................................ )", right_name="( ............................................ )"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    col_widths = [Cm(7.5), Cm(8.0)]
    
    cell_l = table.cell(0, 0)
    cell_l.width = col_widths[0]
    p_l = cell_l.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l1 = p_l.add_run(left_title + "\n\n\n\n\n" + left_name + "\nTanggal: ...................................")
    r_l1.font.name = 'Times New Roman'
    r_l1.font.size = Pt(9.5)
    r_l1.font.color.rgb = COLOR_BLACK
    
    cell_r = table.cell(0, 1)
    cell_r.width = col_widths[1]
    p_r = cell_r.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r.add_run(right_title + "\n\n\n\n\n" + right_name + "\nTanggal: ...................................")
    r_r1.font.name = 'Times New Roman'
    r_r1.font.size = Pt(9.5)
    r_r1.font.color.rgb = COLOR_BLACK

# ==============================================================================
# FOLDER 07: FORMULIR DAN LABEL STATUS SIAP CETAK
# ==============================================================================
def create_folder_07():
    folder_path = os.path.join(BASE_DIR, "07_FORMULIR_DAN_LABEL_STATUS_SIAP_CETAK")
    os.makedirs(folder_path, exist_ok=True)
    
    # FRM-01: BPR Sabun Cuci Piring
    doc = setup_iso_document()
    add_iso_header_box(doc, "Catatan Pengolahan Bets (Batch Processing Record) - Sabun Cuci Piring", "FRM-PRD-01A", "00", "19 Agustus 2026")
    add_p(doc, "Nama Produk       : Sabun Cuci Piring Cair          Besar Bets    : 1.000 Liter\nVarian / Aroma    : Jeruk Nipis                     Nomor Bets    : ............................\nTanggal Mulai     : ............................    Tanggal Selesai: ............................\nNo. Izin Edar PKD : KEMENKES RI PKD [No. PKD]       Tgl Kedaluwarsa: ............................")
    
    add_heading_1(doc, "1. PEMERIKSAAN KESIAPAN RUANGAN (LINE CLEARANCE)")
    add_p(doc, "[  ] Ruang mixing bersih & bebas sisa batch lalu    [  ] Tangki mixer berlabel BERSIH\n[  ] Timbangan digital terkalibrasi (Tara: 0.00)    [  ] Operator memakai APD lengkap\nPetugas Pemeriksa (Kepala Produksi): (...................................) Paraf: [........]")
    
    add_heading_1(doc, "2. PENIMBANGAN BAHAN BAKU")
    headers_t = ["No", "Nama Bahan Kimia", "Formula Standar (1000L)", "Hasil Timbang Aktual", "Nomor Batch Supplier", "Paraf Penimbang", "Paraf QC"]
    data_t = [
        ["1", "Air Demineral / Bersih", "820,0 Kg", "........ Kg", "....................", "[........]", "[........]"],
        ["2", "SLES (Texapon 70%)", "100,0 Kg", "........ Kg", "....................", "[........]", "[........]"],
        ["3", "LABSA (Sulfonat)", "40,0 Kg", "........ Kg", "....................", "[........]", "[........]"],
        ["4", "NaCl (Garam Pengental)", "25,0 Kg", "........ Kg", "....................", "[........]", "[........]"],
        ["5", "Fragrance Lime Oil", "5,0 Kg", "........ Kg", "....................", "[........]", "[........]"],
        ["6", "DMDM Hydantoin (Pengawet)", "2,0 Kg", "........ Kg", "....................", "[........]", "[........]"],
        ["7", "Pewarna Hijau Foodgrade", "0,5 Kg", "........ Kg", "....................", "[........]", "[........]"]
    ]
    add_styled_table(doc, headers_t, data_t, [Cm(0.8), Cm(3.8), Cm(2.4), Cm(2.2), Cm(3.0), Cm(1.6), Cm(1.6)])
    
    add_heading_1(doc, "3. TAHAPAN PROSES PENCAMPURAN (MIXING)")
    add_p(doc, "1. Masukkan Air Baku (800L) ke tangki mixer.                       Jam: ........ Paraf: [......]\n"
               "2. Masukkan SLES dan LABSA, putar kecepatan rendah hingga larut.   Jam: ........ Paraf: [......]\n"
               "3. Larutkan pewarna & pengawet, tuangkan ke tangki mixer.          Jam: ........ Paraf: [......]\n"
               "4. Masukkan Parfum Jeruk Nipis, aduk rata homogen.                Jam: ........ Paraf: [......]\n"
               "5. Masukkan larutan garam NaCl perlahan hingga viskositas pas.     Jam: ........ Paraf: [......]\n"
               "6. Diamkan larutan selama 12-24 jam untuk menghilangkan busa.     Jam: ........ Paraf: [......]")
    
    add_heading_1(doc, "4. HASIL PEMERIKSAAN QC & PELULUSAN BETS")
    add_p(doc, "Bentuk Fisik: [  ] Cairan Kental Jernih    Warna: [  ] Hijau Transparan    Bau: [  ] Khas Jeruk Nipis\n"
               "Nilai pH    : ......... (Standar: 6.0-8.0)  Bobot Jenis: ......... g/mL     Status: [  ] LULUS  [  ] REJECT")
    add_signature_block_2(doc, "Penanggung Jawab Teknis (PJT)", "Kepala Produksi")
    doc.save(os.path.join(folder_path, "FRM-01_Catatan_Pengolahan_Bets_BPR_Sabun_Cuci_Piring.docx"))
    
    # FRM-08: Surat Jalan Pengiriman (Delivery Order)
    doc8 = setup_iso_document()
    add_iso_header_box(doc8, "Surat Jalan Pengiriman Produk Jadi (Delivery Order)", "FRM-GUD-06A", "00", "19 Agustus 2026")
    add_p(doc8, "Nomor Surat Jalan : SJ/PKRT/[Bulan]/[Tahun]/[No]        Tanggal Pengiriman: ............................\n"
                "Kepada Yth.       : PT [Nama Mitra / Distributor]       Kendaraan / No. Pol: ............................\n"
                "Alamat Penerima   : ................................... Nama Pengemudi   : ............................")
    
    headers_sj = ["No", "Nama Produk & Varian", "No. Izin Edar PKD", "Nomor Bets", "Bentuk Kemasan", "Jumlah (Koli/Karton)", "Total Botol/Pcs", "Keterangan QC"]
    data_sj = [
        ["1", "Sabun Cuci Piring Lime 450mL", "KEMENKES RI PKD [No PKD]", "SCP-260819-01", "Karton (24 Btl)", "50 Karton", "1.200 Botol", "CoA Dilampirkan"],
        ["2", "Pembersih Lantai Pine 5L", "KEMENKES RI PKD [No PKD]", "PL-260820-01", "Jerigen 5 Liter", "20 Jerigen", "20 Jerigen", "CoA Dilampirkan"],
        ["3", "Pemutih Pakaian 5.25% 5L", "KEMENKES RI PKD [No PKD]", "BL-260821-01", "Jerigen 5 Liter", "20 Jerigen", "20 Jerigen", "CoA Dilampirkan"]
    ]
    add_styled_table(doc8, headers_sj, data_sj, [Cm(0.8), Cm(3.6), Cm(2.8), Cm(2.2), Cm(2.4), Cm(2.0), Cm(2.0), Cm(2.0)])
    add_p(doc8, "Catatan: Barang telah diperiksa dalam kondisi baik, segel utuh, tidak bocor, dan disertai Certificate of Analysis (CoA) asli dari PJT.")
    
    table_ttd = doc8.add_table(rows=1, cols=3)
    table_ttd.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_ttd, color="FFFFFF", sz="0", val="none")
    ttds = [("Penerima / PT Distributor", "( .................................... )"), ("Pengemudi / Ekspedisi", "( .................................... )"), ("Petugas Gudang PT KCA", "( .................................... )")]
    for i, (t, n) in enumerate(ttds):
        c = table_ttd.cell(0, i)
        c.width = Cm(5.2)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(t + "\n\n\n\n\n" + n + "\nTanggal: .....................")
    doc8.save(os.path.join(folder_path, "FRM-08_Surat_Jalan_Pengiriman_Delivery_Order.docx"))

    # FRM-10: Template Label Status Barang
    doc10 = setup_iso_document()
    add_iso_header_box(doc10, "Template Label Status Barang (Karantina, Diluluskan, Ditolak, Siap Pakai)", "FRM-LBL-10", "00", "19 Agustus 2026")
    add_heading_1(doc10, "LABEL 1: STATUS DIKARANTINA (WARNA KUNING)")
    add_p(doc10, "┌────────────────────────────────────────────────────────────┐\n"
                 "│                    STATUS: DIKARANTINA                     │\n"
                 "│                 (MENUNGGU HASIL UJI QC)                    │\n"
                 "├────────────────────────────────────────────────────────────┤\n"
                 "│ Nama Bahan / Produk : .................................... │\n"
                 "│ Nomor Bets / Lot    : .................................... │\n"
                 "│ Tanggal Masuk       : .................................... │\n"
                 "│ Jumlah Wadah        : ........... Drum / Jerigen / Karton  │\n"
                 "│ Petugas Gudang      : (..................) Paraf: [......] │\n"
                 "└────────────────────────────────────────────────────────────┘")
    
    add_heading_1(doc10, "LABEL 2: STATUS DILULUSKAN (WARNA HIJAU)")
    add_p(doc10, "┌────────────────────────────────────────────────────────────┐\n"
                 "│                    STATUS: DILULUSKAN                      │\n"
                 "│              (MEMENUHI SYARAT MUTU CPKRTB)                 │\n"
                 "├────────────────────────────────────────────────────────────┤\n"
                 "│ Nama Bahan / Produk : .................................... │\n"
                 "│ Nomor Bets / Lot    : .................................... │\n"
                 "│ Tanggal Lulus Uji   : .................................... │\n"
                 "│ Tanggal Kedaluwarsa : .................................... │\n"
                 "│ Penanggung Jwb QC   : (..................) Paraf: [......] │\n"
                 "└────────────────────────────────────────────────────────────┘")

    add_heading_1(doc10, "LABEL 3: STATUS DITOLAK / REJECT (WARNA MERAH)")
    add_p(doc10, "┌────────────────────────────────────────────────────────────┐\n"
                 "│                     STATUS: DITOLAK                        │\n"
                 "│               (TIDAK MEMENUHI SPESIFIKASI)                 │\n"
                 "├────────────────────────────────────────────────────────────┤\n"
                 "│ Nama Bahan / Produk : .................................... │\n"
                 "│ Nomor Bets / Lot    : .................................... │\n"
                 "│ Alasan Penolakan    : .................................... │\n"
                 "│ Tindakan Disposisi  : [ ] Retur Pemasok  [ ] Pemusnahan    │\n"
                 "│ Pengesahan PJT      : (..................) Paraf: [......] │\n"
                 "└────────────────────────────────────────────────────────────┘")
    doc10.save(os.path.join(folder_path, "FRM-10_Template_Label_Status_Barang_Karantina_Lolos_Reject.docx"))
    print("Created folder 07 templates.")

# ==============================================================================
# FOLDER 08: LEGALITAS HR, PJT & ADMINISTRASI MAKLOON
# ==============================================================================
def create_folder_08():
    folder_path = os.path.join(BASE_DIR, "08_LEGALITAS_HR_PJT_DAN_ADMINISTRASI_MAKLOON")
    os.makedirs(folder_path, exist_ok=True)
    
    # LEG-01: SK Direktur Penunjukan PJT
    doc1 = setup_iso_document()
    add_iso_header_box(doc1, "Surat Keputusan Direktur tentang Pengangkatan Penanggung Jawab Teknis (PJT)", "SK-DIR-001", "00", "19 Agustus 2026")
    add_p(doc1, "SURAT KEPUTUSAN DIREKTUR UTAMA PT KCA CHEMICAL\nNomor: 005/SK-DIR/KCA/VIII/2026\nTentang:\nPENGANGKATAN PENANGGUNG JAWAB TEKNIS (PJT) INDUSTRI PERBEKALAN KESEHATAN RUMAH TANGGA (PKRT)")
    
    add_heading_1(doc1, "MENIMBANG:")
    add_bullet(doc1, "Bahwa dalam rangka memenuhi ketentuan Permenkes No. 14 Tahun 2021 dan standar CPKRTB, industri PKRT wajib memiliki Penanggung Jawab Teknis yang memiliki kompetensi di bidang Farmasi/Kimia.", "a. ");
    add_bullet(doc1, "Bahwa nama yang tercantum di bawah ini dinilai cakap, memiliki kualifikasi pendidikan yang sesuai, dan bersedia menjalankan kewajiban sebagai PJT.", "b. ");
    
    add_heading_1(doc1, "MEMUTUSKAN:")
    add_bullet(doc1, "Mengangkat Saudara/i: [Nama Lengkap PJT, Gelar S.Farm / S.Si / A.Md.Farm] sebagai PENANGGUNG JAWAB TEKNIS (PJT) PT KCA CHEMICAL terhitung sejak tanggal 19 Agustus 2026.", "PERTAMA: ");
    add_bullet(doc1, "PJT bertanggung jawab penuh atas pemenuhan Cara Pembuatan PKRT yang Baik (CPKRTB), pengesahan Master Formula, registrasi Izin Edar Kemenkes RI PKD, pengawasan mutu laboratorium, dan pelulusan produk jadi (Batch Release).", "KEDUA: ");
    add_bullet(doc1, "Surat Keputusan ini berlaku sejak tanggal ditetapkan dan apabila di kemudian hari terdapat kekeliruan akan dilakukan perbaikan sebagaimana mestinya.", "KETIGA: ");
    
    add_signature_block_2(doc1, "Diterima dan Disanggupi Oleh,\nPenanggung Jawab Teknis (PJT)", "Ditetapkan di Kediri,\nDirektur Utama PT KCA Chemical")
    doc1.save(os.path.join(folder_path, "LEG-01_SK_Direktur_Penunjukan_PJT_dan_Struktur_Organisasi.docx"))

    # LEG-02: Surat Pernyataan PJT Bekerja Penuh Waktu
    doc2 = setup_iso_document()
    add_iso_header_box(doc2, "Surat Pernyataan Penanggung Jawab Teknis Bekerja Penuh Waktu (Full-Time)", "LEG-PJT-002", "00", "19 Agustus 2026")
    add_p(doc2, "SURAT PERNYATAAN PENANGGUNG JAWAB TEKNIS (PJT)\nBEKERJA PENUH WAKTU (FULL TIME)")
    add_p(doc2, "Yang bertanda tangan di bawah ini:\n"
                "Nama Lengkap         : [Nama Lengkap PJT]\n"
                "Tempat, Tanggal Lahir: [Tempat, Tgl Lahir]\n"
                "Pendidikan Terakhir  : [S1 Farmasi / D3 Farmasi / S1 Kimia / S1 Teknik Kimia]\n"
                "Nomor STRTTK / SIPA  : [Nomor STRTTK/SIPA jika ada]\n"
                "Alamat KTP           : [Alamat Lengkap KTP]\n"
                "Nomor Telepon / HP   : [Nomor Kontak HP]\n\n"
                "Dengan ini menyatakan dengan sebenar-benarnya bahwa saya:\n"
                "1. Bersedia dan bertindak sebagai PENANGGUNG JAWAB TEKNIS (PJT) pada sarana industri PKRT PT KCA CHEMICAL yang beralamat di [Alamat Pabrik].\n"
                "2. Bekerja secara PENUH WAKTU (FULL TIME) dan tidak bekerja sebagai Penanggung Jawab Teknis pada sarana produksi atau distribusi farmasi/alat kesehatan/PKRT lain.\n"
                "3. Bertanggung jawab penuh terhadap aspek teknis formulasi, pengawasan mutu, dan kepatuhan standar CPKRTB Kementerian Kesehatan RI.\n\n"
                "Demikian surat pernyataan ini saya buat dengan sadar, tanpa paksaan dari pihak manapun, dan bermaterai cukup untuk dipergunakan sebagaimana mestinya.")
    
    add_signature_block_2(doc2, "Yang Membuat Pernyataan,\nPenanggung Jawab Teknis (PJT)\n\n\n( Materai Rp 10.000 )", "Mengetahui,\nDirektur Utama PT KCA Chemical")
    doc2.save(os.path.join(folder_path, "LEG-02_Surat_Pernyataan_PJT_Bekerja_Penuh_Waktu_FullTime.docx"))

    # LEG-03: NDA Kerahasiaan Formula Makloon
    doc3 = setup_iso_document()
    add_iso_header_box(doc3, "Perjanjian Kerahasiaan Informasi dan Formula (Non-Disclosure Agreement - NDA)", "NDA-MKT-003", "00", "19 Agustus 2026")
    add_p(doc3, "PERJANJIAN KERAHASIAAN INFORMASI DAN RAHASIA DAGANG (NON-DISCLOSURE AGREEMENT)\nNomor: 012/NDA/KCA-MKT/VIII/2026")
    add_p(doc3, "Perjanjian Kerahasiaan ini dibuat antara PT KCA CHEMICAL (Pihak Pabrik Manufaktur) dan PT [Nama Mitra Distributor] (Pihak Pemilik Merek/Klien Makloon).\n\n"
                "PASAL 1: DEFINISI INFORMASI RAHASIA\n"
                "Informasi Rahasia mencakup komposisi kimia, Master Formula, metode pencampuran (mixing), data biaya produksi, strategi pemasaran, dan data pelanggan yang dipertukarkan antara PARA PIHAK.\n\n"
                "PASAL 2: KEWAJIBAN KERAHASIAAN\n"
                "1. PARA PIHAK sepakat untuk menjaga kerahasiaan seluruh informasi dan tidak akan membocorkan, mempublikasikan, atau mengalihkan informasi formula kepada pihak ketiga manapun tanpa persetujuan tertulis.\n"
                "2. Kewajiban kerahasiaan ini berlaku selama masa kerja sama makloon berlangsung dan tetap mengikat hingga 5 (lima) tahun setelah kerja sama berakhir.\n\n"
                "PASAL 3: SANKSI DAN GANTI RUGI\n"
                "Pelanggaran terhadap perjanjian kerahasiaan ini memberikan hak kepada pihak yang dirugikan untuk menuntut ganti rugi material dan immaterial sesuai peraturan perundang-undangan Republik Indonesia (UU Rahasia Dagang No. 30 Tahun 2000).")
    add_signature_block_2(doc3, "PIHAK PABRIK,\nPT KCA CHEMICAL\n\n\n( Materai Rp 10.000 )", "PIHAK KLIEN / MITRA,\nPT [NAMA MITRA DISTRIBUTOR]\n\n\n( Materai Rp 10.000 )")
    doc3.save(os.path.join(folder_path, "LEG-03_Surat_Perjanjian_Kerahasiaan_Formula_NDA_Makloon.docx"))

    # LEG-04: Surat Kuasa Pendaftaran Izin Edar PKD dari Mitra
    doc4 = setup_iso_document()
    add_iso_header_box(doc4, "Surat Kuasa Pendaftaran Izin Edar Kemenkes RI PKD (Letter of Authorization)", "LOA-REG-004", "00", "19 Agustus 2026")
    add_p(doc4, "SURAT KUASA PENDAFTARAN IZIN EDAR PKRT KEMENKES RI\nNomor: [No.Kuasa/Mitra/Bulan/Tahun]\n\n"
                "Yang bertanda tangan di bawah ini:\n"
                "Nama Direktur Mitra   : [Nama Direktur Pemilik Merek]\n"
                "Nama Perusahaan Mitra : PT [Nama Mitra Distributor]\n"
                "Pemilik Merek Dagang  : Merek \"[Nama Merek]\" (Nomor Agenda/Sertifikat DJKI: [........])\n\n"
                "Dengan ini MEMBERIKAN KUASA PENUH kepada:\n"
                "Nama Perusahaan Pabrik: PT KCA CHEMICAL\n"
                "Alamat Pabrik         : [Alamat Pabrik Anda]\n"
                "Nama PJT Pelaksana    : [Nama PJT PT KCA]\n\n"
                "Untuk mendaftarkan, mengurus, dan mewakili permohonan Izin Edar Perbekalan Kesehatan Rumah Tangga (KEMENKES RI PKD) atas produk Merek \"[Nama Merek]\" melalui sistem e-Farmalkes Kementerian Kesehatan RI hingga terbit Nomor Izin Edar resmi.")
    add_signature_block_2(doc4, "Pemberi Kuasa (Pemilik Merek),\nPT [Nama Mitra Distributor]\n\n\n( Materai Rp 10.000 )", "Penerima Kuasa (Pabrik),\nPT KCA CHEMICAL\n\n\n( Materai Rp 10.000 )")
    doc4.save(os.path.join(folder_path, "LEG-04_Surat_Kuasa_Pendaftaran_Izin_Edar_PKD_dari_Mitra.docx"))
    print("Created folder 08 legal & HR templates.")

if __name__ == "__main__":
    create_folder_07()
    create_folder_08()
    print("\nALL ADDITIONAL ESSENTIAL FILES CREATED IN 'KCA DOKUMEN'!")

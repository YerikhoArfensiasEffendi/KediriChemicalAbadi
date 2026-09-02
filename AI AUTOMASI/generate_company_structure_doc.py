import os
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from kca_doc_humanizer import sanitize_docx_metadata

charts_dir = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/Keuangan/Proposal/charts"
os.makedirs(charts_dir, exist_ok=True)
org_chart_png = os.path.join(charts_dir, "chart5_kca_org_structure.png")

# 1. RENDER VISUAL ORGANIZATIONAL CHART
def render_org_chart():
    fig, ax = plt.subplots(figsize=(10, 6.2), dpi=300)
    ax.axis('off')
    
    # Style helper
    def draw_box(x, y, w, h, title, name, role_desc, bg_color="#0E2A47", border_color="#0E2A47", text_white=True):
        rect = patches.FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.1", 
                                      edgecolor=border_color, facecolor=bg_color, linewidth=1.8)
        ax.add_patch(rect)
        t_col = "white" if text_white else "#0F172A"
        sub_col = "#E2E8F0" if text_white else "#334155"
        ax.text(x + w/2, y + h - 0.28, title, color=t_col, fontsize=8.5, fontweight='bold', ha='center', va='center')
        ax.text(x + w/2, y + h/2 - 0.05, name, color=t_col, fontsize=9.5, fontweight='bold', ha='center', va='center')
        ax.text(x + w/2, y + 0.25, role_desc, color=sub_col, fontsize=7.5, ha='center', va='center', style='italic')

    # Line helper
    def draw_connector(x1, y1, x2, y2, color="#0E2A47"):
        ax.plot([x1, x2], [y1, y2], color=color, linewidth=2)

    # 1. Direktur Utama (Top)
    draw_box(3.5, 5.0, 3.0, 1.1, "DIREKTUR UTAMA / OWNER", "YAN EFFENDI", "Kebijakan Strategis & Pengesahan Anggaran", "#0E2A47")
    
    # Line down to GM
    draw_connector(5.0, 5.0, 5.0, 4.4)
    
    # 2. General Manager / Finance & Ops (Middle)
    draw_box(3.2, 3.3, 3.6, 1.1, "GENERAL MANAGER & OPERATIONAL CONTROL", "YERIKHO ARFENSIAS EFFENDI", "Manajemen Operasional, Keuangan & ISO 9001", "#1E3A8A")

    # Line down to Department branches
    draw_connector(5.0, 3.3, 5.0, 2.7)
    draw_connector(1.2, 2.7, 8.8, 2.7) # Horizontal bus bar
    
    # Drop lines to 4 Depts
    draw_connector(1.2, 2.7, 1.2, 2.2)
    draw_connector(3.7, 2.7, 3.7, 2.2)
    draw_connector(6.3, 2.7, 6.3, 2.2)
    draw_connector(8.8, 2.7, 8.8, 2.2)
    
    # 3. Four Departments (Bottom)
    draw_box(0.1, 0.8, 2.2, 1.4, "PENANGGUNG JAWAB TEKNIS", "PJT / APOTEKER / QC", "Kendali Mutu, Formulasi,\nUji Lab & CoA Resmi", "#F8FAFC", "#1E3A8A", False)
    draw_box(2.6, 0.8, 2.2, 1.4, "PRODUKSI & MIXING", "OPERATOR PRODUKSI", "Penimbangan Bahan,\nMixing, Filling & Sealing", "#F8FAFC", "#1E3A8A", False)
    draw_box(5.2, 0.8, 2.2, 1.4, "LOGISTIK & GUDANG", "TENAGA OPERASIONAL", "Gudang FIFO/FEFO,\nBongkar Muat & Delivery", "#F8FAFC", "#1E3A8A", False)
    draw_box(7.7, 0.8, 2.2, 1.4, "MARKETING & MAKLOON", "SALES & DISTRIBUSI", "Layanan Distributor,\nOrder B2B & Faktur Invoice", "#F8FAFC", "#1E3A8A", False)

    ax.set_xlim(-0.2, 10.2)
    ax.set_ylim(0.4, 6.5)
    fig.suptitle("BAGAN STRUKTUR ORGANISASI RESMI PT KEDIRI CHEMICAL ABADI", fontsize=12, fontweight='bold', color='#0E2A47', y=0.98)
    
    plt.tight_layout()
    plt.savefig(org_chart_png, bbox_inches='tight', dpi=300)
    plt.close()
    print("Saved Org Chart PNG:", org_chart_png)

def set_cell_shd(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_mrg(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_doc():
    render_org_chart()
    
    base_dir = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical"
    doc_paths = [
        os.path.join(base_dir, "KCA DOKUMEN/09_DOKUMEN_DIREKSI_DAN_MANAJEMEN_PUNCAK/MGT-07_Struktur_Organisasi_dan_Uraian_Jabatan_Resmi_KCA.docx"),
        "/Users/arthur/Documents/KCA DOKUMEN/09_DOKUMEN_DIREKSI_DAN_MANAJEMEN_PUNCAK/MGT-07_Struktur_Organisasi_dan_Uraian_Jabatan_Resmi_KCA.docx"
    ]

    for p in doc_paths:
        os.makedirs(os.path.dirname(p), exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # ---------------- KOP SURAT ----------------
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k1 = kop.add_run("PT KEDIRI CHEMICAL ABADI\n")
    k1.font.name = "Arial"
    k1.font.size = Pt(16)
    k1.font.bold = True
    k1.font.color.rgb = RGBColor(0, 0, 0)

    k2 = kop.add_run("SISTEM MANAJEMEN MUTU ISO 9001:2015 & TATA KELOLA PERUSAHAAN\n")
    k2.font.name = "Arial"
    k2.font.size = Pt(10)
    k2.font.bold = True
    k2.font.color.rgb = RGBColor(0, 0, 0)

    k3 = kop.add_run("Pabrik: RT.1/RW.6, Pagung, Kec. Semen, Kabupaten Kediri, Jawa Timur 64161\n"
                    "Hotline/WhatsApp: 0822-4400-6699 | Email: kdrchemicals@gmail.com\n")
    k3.font.name = "Arial"
    k3.font.size = Pt(9)
    k3.font.color.rgb = RGBColor(0, 0, 0)

    p_line = doc.add_paragraph("═" * 58)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_line.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- JUDUL ----------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title_p.add_run("STRUKTUR ORGANISASI & URAIAN TUGAS JABATAN (JOB DESCRIPTION)\n")
    t1.font.name = "Arial"
    t1.font.size = Pt(13.5)
    t1.font.bold = True
    t1.font.color.rgb = RGBColor(0, 0, 0)

    t2 = title_p.add_run("Nomor Dokumen: MGT-07/KCA-ORG/VIII/2026 | Status: Dokumen Terkendali | Rev: 03\n")
    t2.font.name = "Arial"
    t2.font.size = Pt(9.5)
    t2.font.italic = True
    t2.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 1: BAGAN STRUKTUR ORGANISASI ----------------
    h1 = doc.add_heading("1. BAGAN STRUKTUR ORGANISASI RESMI", level=2)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p1 = doc.add_paragraph(
        "Bagan struktur organisasi di bawah ini menetapkan garis komando, wewenang, dan tanggung jawab fungsional "
        "seluruh jajaran manajemen dan operasional di lingkungan PT Kediri Chemical Abadi:"
    )
    p1.paragraph_format.line_spacing = 1.15
    for r in p1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # EMBED ORG CHART IMAGE
    if os.path.exists(org_chart_png):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(6)
        doc.add_picture(org_chart_png, width=Inches(6.4))
        p_cap = doc.add_paragraph("Gambar 1: Bagan Hierarki dan Garis Tanggung Jawab PT Kediri Chemical Abadi")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(10)
        for r in p_cap.runs:
            r.font.size = Pt(8.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 2: URAIAN TUGAS & WEWENANG (JOB DESCRIPTIONS) ----------------
    doc.add_page_break()
    h2 = doc.add_heading("2. URAIAN TUGAS, TANGGUNG JAWAB & WEWENANG JABATAN", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    # 1. Direktur Utama
    doc.add_heading("A. DIREKTUR UTAMA (TOP MANAGEMENT / OWNER)", level=3)
    for r in doc.paragraphs[-1].runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_dir = doc.add_paragraph(
        "Pejabat: Bpk. Yan Effendi\n"
        "Tanggung Jawab Utama:\n"
        "1. Menetapkan visi, misi, arah kebijakan strategis, dan sasaran mutu jangka panjang perusahaan.\n"
        "2. Mengotorisasi keputusan investasi modal besar, pengadaan mesin/fasilitas baru (CapEx), dan restrukturisasi kewajiban bisnis.\n"
        "3. Menandatangani dokumen legalitas resmi, perjanjian kerjasama strategis (MoU/Kontrak Makloon), dan laporan keuangan tahunan.\n"
        "4. Mewakili perusahaan dalam hubungan hukum dengan instansi pemerintah, perbankan, dan mitra bisnis utama.\n"
        "Wewenang:\n"
        "• Otorisasi penuh atas struktur kepemilikan, perubahan anggaran dasar, dan pembagian dividen perusahaan."
    )
    p_dir.paragraph_format.line_spacing = 1.15
    for r in p_dir.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 2. General Manager
    doc.add_heading("B. GENERAL MANAGER / OPERATIONAL & FINANCE CONTROL", level=3)
    for r in doc.paragraphs[-1].runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_gm = doc.add_paragraph(
        "Pejabat: Bpk. Yerikho Arfensias Effendi\n"
        "Tanggung Jawab Utama:\n"
        "1. Memimpin, mengendalikan, dan mengawasi seluruh operasional harian pabrik, produksi, gudang, dan logistik delivery.\n"
        "2. Mengelola administrasi keuangan, arus kas (cash flow), budgeting bulanan, dan rekonsiliasi pembukuan akuntansi.\n"
        "3. Bertindak sebagai Management Representative (MR) dalam penerapan dan pemeliharaan Sistem Manajemen Mutu ISO 9001:2015 dan standar CPKRTB.\n"
        "4. Mengendalikan rantai pasok (supply chain), negosiasi pengadaan bahan baku kimia, serta mengawasi penagihan invoice piutang usaha.\n"
        "5. Menyetujui Purchase Order (PO), Sales Order (SO), Delivery Order (DO), serta pengesahan payroll gaji karyawan.\n"
        "Wewenang:\n"
        "• Pengambilan keputusan operasional harian, rotasi penugasan staf, persetujuan belanja rutin, dan validasi dokumen mutu pabrik."
    )
    p_gm.paragraph_format.line_spacing = 1.15
    for r in p_gm.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 3. Penanggung Jawab Teknis (PJT)
    doc.add_heading("C. PENANGGUNG JAWAB TEKNIS (PJT) / QA & LABORATORIUM", level=3)
    for r in doc.paragraphs[-1].runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_pjt = doc.add_paragraph(
        "Kualifikasi: Tenaga Teknis Kefarmasian / Kimia (PJT Resmi Terdaftar Kemenkes)\n"
        "Tanggung Jawab Utama:\n"
        "1. Menjamin seluruh proses formulasi dan produksi memenuhi standar Cara Pembuatan PKRT yang Baik (CPKRTB).\n"
        "2. Melakukan pengujian mutu laboratorium (uji organoleptik, pH, bobot jenis, kadar zat aktif, dan uji stabilitas).\n"
        "3. Menerbitkan dan menandatangani Sertifikat Analisis (Certificate of Analysis / CoA) resmi untuk setiap bets produk rilis.\n"
        "4. Mengelola sampel pertinggal (retained sample) dan menangani investigasi teknis apabila terdapat keluhan mutu pelanggan."
    )
    p_pjt.paragraph_format.line_spacing = 1.15
    for r in p_pjt.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 4. Operator Produksi
    doc.add_heading("D. OPERATOR PRODUKSI & MIXING (TENAGA KERJA TETAP)", level=3)
    for r in doc.paragraphs[-1].runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_prd = doc.add_paragraph(
        "Status: Karyawan Tetap Pabrik\n"
        "Tanggung Jawab Utama:\n"
        "1. Melaksanakan penimbangan bahan baku kimia secara teliti sesuai Catatan Pengolahan Bets (BPR).\n"
        "2. Mengoperasikan mesin tangki mixer homogenizer dan mengawasi tahapan pelarutan serta pengadukan.\n"
        "3. Menjalankan proses pengisian (filling), penutupan (capping), penyegelan induksi (induction sealing), dan pelabelan kemasan.\n"
        "4. Menjaga kebersihan, sanitasi ruangan produksi, serta pemeliharaan rutin mesin-mesin pabrik."
    )
    p_prd.paragraph_format.line_spacing = 1.15
    for r in p_prd.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # 5. Tenaga Operasional & Logistik
    doc.add_heading("E. TENAGA OPERASIONAL, GUDANG & LOGISTIK (FREELANCE / LAPANGAN)", level=3)
    for r in doc.paragraphs[-1].runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_ops = doc.add_paragraph(
        "Status: Tenaga Operasional & Delivery\n"
        "Tanggung Jawab Utama:\n"
        "1. Melaksanakan penerimaan, pengecekan fisik, dan penataan bahan baku masuk di gudang sesuai sistem FIFO/FEFO.\n"
        "2. Mengelola penyimpanan produk jadi, penyusunan karton pada palet, dan pengecekan segel kemasan siap kirim.\n"
        "3. Menjalankan aktivitas pengiriman barang (delivery) ke lokasi customer/distributor menggunakan armada operasional.\n"
        "4. Memastikan kelengkapan Surat Jalan (Delivery Order) bertanda tangan penerima saat serah terima barang."
    )
    p_ops.paragraph_format.line_spacing = 1.15
    for r in p_ops.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 3: MATRIKS KOMPETENSI & DISTRIBUSI WEWENANG ----------------
    h3 = doc.add_heading("3. MATRIKS WEWENANG & PERSETUJUAN DOKUMEN", level=2)
    for r in h3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    t_mat = doc.add_table(rows=1, cols=4)
    t_mat.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_mat.rows[0].cells
    m_hdr[0].text = "Jenis Dokumen / Transaksi"
    m_hdr[1].text = "Disusun / Diajukan"
    m_hdr[2].text = "Diverifikasi / Disetujui"
    m_hdr[3].text = "Disahkan (Otorisasi Akhir)"
    for c in m_hdr:
        set_cell_shd(c, "1E293B")
        set_cell_mrg(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    mat_rows = [
        ("Kebijakan Mutu & Anggaran Tahunan", "General Manager", "General Manager", "Direktur Utama"),
        ("Kontrak Kerjasama Makloon & MoU", "General Manager", "General Manager", "Direktur Utama"),
        ("Laporan Keuangan Bulanan & P&L", "General Manager", "General Manager", "Direktur Utama"),
        ("Purchase Order (PO) Bahan Baku", "Tenaga Gudang / Ops", "General Manager", "General Manager"),
        ("Surat Jalan (DO) & Invoice Penjualan", "Marketing / Admin", "General Manager", "General Manager"),
        ("Sertifikat Analisis (CoA) Produk", "Staf QC / Analis", "PJT Apoteker", "PJT Apoteker"),
        ("Catatan Bets Produksi (BPR)", "Operator Produksi", "PJT Apoteker", "General Manager")
    ]
    for dok, susun, verif, sah in mat_rows:
        row = t_mat.add_row().cells
        row[0].text = dok
        row[1].text = susun
        row[2].text = verif
        row[3].text = sah
        for c in row:
            set_cell_mrg(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- LEMBAR PENGESAHAN ----------------
    doc.add_paragraph("\n")
    h4 = doc.add_heading("4. LEMBAR PENGESAHAN DOKUMEN RESMI", level=2)
    for r in h4.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_sah = doc.add_paragraph(
        "Dokumen Struktur Organisasi dan Uraian Tugas Jabatan ini berlaku efektif sejak tanggal ditetapkan dan "
        "mengikat seluruh personil di lingkungan PT Kediri Chemical Abadi."
    )
    p_sah.paragraph_format.line_spacing = 1.15
    for r in p_sah.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph("\n")
    t_sign = doc.add_table(rows=3, cols=2)
    t_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t_sign.rows[0].cells[0].text = "Disiapkan Oleh:\nGENERAL MANAGER / MANAGEMENT REPRESENTATIVE"
    t_sign.rows[0].cells[1].text = "Disetujui & Disahkan Oleh:\nDIREKTUR UTAMA"
    
    t_sign.rows[1].cells[0].text = "\n\n\n\n"
    t_sign.rows[1].cells[1].text = "\n\n\n\n"

    t_sign.rows[2].cells[0].text = "YERIKHO ARFENSIAS EFFENDI\nTanggal: 14 Agustus 2026"
    t_sign.rows[2].cells[1].text = "YAN EFFENDI\nTanggal: 18 Agustus 2026"

    for r in t_sign.rows:
        for c in r.cells:
            set_cell_mrg(c)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_run in p.runs:
                    r_run.font.bold = True
                    r_run.font.color.rgb = RGBColor(0, 0, 0)

    for p_out in doc_paths:
        doc.save(p_out)
        sanitize_docx_metadata(p_out, title="Struktur Organisasi dan Uraian Jabatan Resmi KCA", author="Yerikho Arfensias Effendi")
        print("Generated and Sanitized:", p_out)

if __name__ == "__main__":
    generate_doc()

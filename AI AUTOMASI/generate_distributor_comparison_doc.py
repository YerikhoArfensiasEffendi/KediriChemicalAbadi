import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from kca_doc_humanizer import sanitize_docx_metadata

def set_cell_shading(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_comparison_doc():
    base_dir = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical"
    output_dir = os.path.join(base_dir, "Keuangan/Proposal")
    os.makedirs(output_dir, exist_ok=True)
    doc_path = os.path.join(output_dir, "KOMPARASI_FINANSIAL_DISTRIBUTOR_REGULER_VS_DEDICATED_KCA.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # ---------------- KOP SURAT RESMI ----------------
    kop = doc.add_paragraph()
    kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    k1 = kop.add_run("PT KEDIRI CHEMICAL ABADI\n")
    k1.font.name = "Arial"
    k1.font.size = Pt(16)
    k1.font.bold = True
    k1.font.color.rgb = RGBColor(0, 0, 0)

    k2 = kop.add_run("ANALISIS STRATEGIS & KOMPARASI FINANSIAL DISTRIBUSI KIMIA\n")
    k2.font.name = "Arial"
    k2.font.size = Pt(10)
    k2.font.bold = True
    k2.font.color.rgb = RGBColor(0, 0, 0)

    k3 = kop.add_run("Pabrik: RT.1/RW.6, Pagung, Kec. Semen, Kabupaten Kediri, Jawa Timur 64161\n"
                    "Hotline/WhatsApp: 0822-4400-6699 | Email: kdrchemicals@gmail.com\n")
    k3.font.name = "Arial"
    k3.font.size = Pt(9)
    k3.font.color.rgb = RGBColor(0, 0, 0)

    p_line = doc.add_paragraph("═" * 58)
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_line.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- JUDUL KOMPARASI ----------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t1 = title_p.add_run("STUDI KOMPARASI FINANSIAL & OPERASIONAL DISTRIBUTOR\n")
    t1.font.name = "Arial"
    t1.font.size = Pt(13.5)
    t1.font.bold = True
    t1.font.color.rgb = RGBColor(0, 0, 0)

    t2 = title_p.add_run("SKENARIO A (DISTRIBUSI REGULER TANPA MODAL MESIN)\n"
                        "VS SKENARIO B (KEMITRAAN DEDICATED CAPEX DENGAN DISKON INVOICE 10%)\n")
    t2.font.name = "Arial"
    t2.font.size = Pt(11)
    t2.font.bold = True
    t2.font.color.rgb = RGBColor(0, 0, 0)

    t3 = title_p.add_run("Nomor Dokumen: 003/CMP-KCA/REG-VS-CAPEX/IX/2026 | Basis Volume: 200.000 Unit\n")
    t3.font.name = "Arial"
    t3.font.size = Pt(9.5)
    t3.font.italic = True
    t3.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- RINGKASAN INTISARI ----------------
    h_exec = doc.add_heading("INTISARI PERBANDINGAN STRATEGIS", level=2)
    for r in h_exec.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p_exec = doc.add_paragraph(
        "Dokumen ini menyajikan perbandingan head-to-head antara dua model bisnis bagi Distributor dalam memasarkan produk "
        "PT Kediri Chemical Abadi dengan basis volume yang sama yaitu 200.000 unit:\n"
        "• Skenario A (Distributor Reguler): Tanpa modal setor awal Rp 200 Juta, membeli di harga normal pabrik Rp 10.000/unit.\n"
        "• Skenario B (Mitra Dedicated CapEx): Menyetor modal mesin Rp 200 Juta di awal, membeli di harga diskon Rp 9.000/unit (potongan 10%)."
    )
    p_exec.paragraph_format.line_spacing = 1.15
    for r in p_exec.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- TABEL UTAMA PERBANDINGAN ----------------
    h1 = doc.add_heading("1. TABEL MATRIKS KOMPARASI FINANSIAL HEAD-TO-HEAD", level=2)
    for r in h1.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    t_main = doc.add_table(rows=1, cols=4)
    t_main.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_main.rows[0].cells
    m_hdr[0].text = "Parameter Finansial"
    m_hdr[1].text = "Skenario A\n(Distributor Reguler)"
    m_hdr[2].text = "Skenario B\n(Mitra Mesin Dedikasi)"
    m_hdr[3].text = "Analisis Keunggulan Skenario B"
    for c in m_hdr:
        set_cell_shading(c, "1E293B")
        set_cell_margins(c)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.bold = True

    comp_rows = [
        ("Modal Setor Awal Mesin", "Rp 0", "Rp 200.000.000", "Modal mesin di awal untuk booking kapasitas pabrik"),
        ("Harga Beli per Unit dari KCA", "Rp 10.000 / unit", "Rp 9.000 / unit", "Lebih hemat Rp 1.000 / unit (10% Diskon Invoice)"),
        ("Kas Keluar Belanja (200.000 Unit)", "Rp 2.000.000.000", "Rp 1.800.000.000", "Kas belanja barang LEBIH HEMAT Rp 200.000.000"),
        ("TOTAL MODAL KELUAR (Mesin + Belanja)", "Rp 2.000.000.000", "Rp 2.000.000.000", "TOTAL MODAL YANG KELUAR SAMA PERSIS!"),
        ("Harga Jual ke Pasar (Laundry/Hotel)", "Rp 15.000 / unit", "Rp 15.000 / unit", "Harga jual rata-rata di pasar"),
        ("Total Kas Masuk dari Pasar", "Rp 3.000.000.000", "Rp 3.000.000.000", "Total penerimaan kas dari konsumen"),
        ("Pengembalian Modal Mesin Awal", "Rp 0 (Tidak Ada)", "Rp 200.000.000", "100% Modal Pokok Mesin Kembali Utuh"),
        ("LABA BERSIH OPERASIONAL", "Rp 1.000.000.000", "Rp 1.000.000.000", "Profit bersih dari margin penjualan barang"),
        ("Margin Keuntungan per Unit Beli", "50,00%", "66,67%", "Lonjakan Efisiensi Margin +16,67% Murni"),
        ("Beban Kas per Batch (10.000 Unit)", "Rp 100.000.000 / order", "Rp 90.000.000 / order", "Arus kas harian LEBIH RINGAN Rp 10 Juta/order")
    ]
    for param, ska, skb, anal in comp_rows:
        row = t_main.add_row().cells
        row[0].text = param
        row[1].text = ska
        row[2].text = skb
        row[3].text = anal
        for c in row:
            set_cell_margins(c)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor(0, 0, 0)
        if "TOTAL MODAL KELUAR" in param or "LABA BERSIH" in param or "Margin Keuntungan" in param:
            set_cell_shading(row[0], "F1F5F9")
            set_cell_shading(row[1], "F1F5F9")
            set_cell_shading(row[2], "F1F5F9")
            set_cell_shading(row[3], "F1F5F9")
            for c in row:
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 2 ----------------
    h2 = doc.add_heading("2. PERBEDAAN STRATEGIS & DAYA SAING DI LAPANGAN", level=2)
    for r in h2.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p2 = doc.add_paragraph(
        "Meskipun secara nominal laba kotor pada volume 200.000 unit sama-sama Rp 1 Miliar, posisi strategis Skenario B "
        "jauh lebih unggul dan mendominasi pasar karena faktor-faktor berikut:\n\n"
        "1. Kekuatan Banting Harga (Pricing Warfare Dominance):\n"
        "   • Di Skenario A, harga modal beli distributor adalah Rp 10.000. Jika kompetitor menjual di harga Rp 12.000, margin distributor hanya Rp 2.000 (sangat tipis).\n"
        "   • Di Skenario B, harga modal beli distributor hanya Rp 9.000. Distributor dapat menjual di harga Rp 12.000 untuk merebut kontrak hotel/laundry besar, dan tetap mengantongi laba tebal Rp 3.000 (33,33%).\n\n"
        "2. Beban Arus Kas Harian Jauh Lebih Ringan (Working Capital Relief):\n"
        "   • Di Skenario A, setiap kali distributor melakukan restock 10.000 botol, kas tunai yang harus disetor ke pabrik adalah Rp 100 Juta.\n"
        "   • Di Skenario B, distributor hanya perlu menyetor Rp 90 Juta tunai. Sisa Rp 10 Juta langsung memotong modal mesin, sehingga sisa kas kerja dapat dipakai untuk biaya operasional armada sales.\n\n"
        "3. Jaminan Antrean Produksi & Formula Custom Eksklusif:\n"
        "   • Di Skenario A, distributor bersaing dengan pembeli lain. Jika kapasitas pabrik penuh, pesanan bisa tertunda.\n"
        "   • Di Skenario B, lini mesin homogenizer & filling SUS316 didedikasikan secara prioritas untuk distributor, dengan kebebasan menentukan aroma, warna, dan formulasi eksklusif."
    )
    p2.paragraph_format.line_spacing = 1.15
    for r in p2.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- BAB 3 ----------------
    h3 = doc.add_heading("3. KESIMPULAN REKOMENDASI UNTUK DISTRIBUTOR", level=2)
    for r in h3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.bold = True

    p3 = doc.add_paragraph(
        "Kesimpulan Finansial:\n"
        "Skenario B (Kemitraan Mesin Dedikasi) mengubah modal pasif Rp 200 Juta menjadi instrumen pengurang biaya operasional "
        "yang menghasilkan pengembalian modal 100% plus laba bersih Rp 1.000.000.000,- dengan proteksi pasar eksklusif. "
        "Skema ini memberikan keamanan bisnis dan daya saing harga yang tidak dapat ditandingi oleh distributor reguler manapun."
    )
    p3.paragraph_format.line_spacing = 1.15
    for r in p3.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    # ---------------- LEMBAR PENGESAHAN ----------------
    doc.add_paragraph("\n")
    t_sign = doc.add_table(rows=3, cols=2)
    t_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    t_sign.rows[0].cells[0].text = "DISETUJUI OLEH:\nPT KEDIRI CHEMICAL ABADI"
    t_sign.rows[0].cells[1].text = "DITERIMA & DIKONFIRMASI:\nMITRA DISTRIBUTOR UTAMA"
    
    t_sign.rows[1].cells[0].text = "\n\n\n\n"
    t_sign.rows[1].cells[1].text = "\n\n\n\n"

    t_sign.rows[2].cells[0].text = "YAN EFFENDI / YERIKHO ARFENSIAS E.\nDirektur Utama / General Manager"
    t_sign.rows[2].cells[1].text = "(......................................................)\nNama & Jabatan"

    for r in t_sign.rows:
        for c in r.cells:
            set_cell_margins(c)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r_run in p.runs:
                    r_run.font.bold = True
                    r_run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(doc_path)
    
    # Salin ke arsip legalitas
    legal_dir = os.path.join(base_dir, "KCA DOKUMEN/06_MASTER_MANUAL_DAN_LEGALITAS")
    os.makedirs(legal_dir, exist_ok=True)
    legal_path = os.path.join(legal_dir, "KOMPARASI_FINANSIAL_DISTRIBUTOR_REGULER_VS_DEDICATED_KCA.docx")
    doc.save(legal_path)
    
    # Sanitasi Metadata
    sanitize_docx_metadata(doc_path, title="Komparasi Finansial Distributor Reguler vs Dedicated KCA", author="Yerikho Arfensias Effendi")
    sanitize_docx_metadata(legal_path, title="Komparasi Finansial Distributor Reguler vs Dedicated KCA", author="Yerikho Arfensias Effendi")
    
    print(f"SUCCESS: Comparison Document generated at {doc_path} and {legal_path}")
    return doc_path

if __name__ == "__main__":
    create_comparison_doc()

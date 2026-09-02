import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

BASE_DIR = "/Users/arthur/Documents/WebsiteKCA/kediri-chemical/KCA DOKUMEN"
MGT_DIR = os.path.join(BASE_DIR, "09_DOKUMEN_DIREKSI_DAN_MANAJEMEN_PUNCAK")
os.makedirs(MGT_DIR, exist_ok=True)

COLOR_BLACK = RGBColor(0, 0, 0)
HEX_BLACK = "000000"
HEX_HEADER_BG = "D9D9D9"
HEX_ROW_ALT = "F7F7F7"
HEX_BORDER = "000000"

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_table_borders(table, color=HEX_BORDER, sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def setup_iso_document():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.top_margin = Cm(3.0)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_BLACK
    style.paragraph_format.line_spacing = 1.3
    style.paragraph_format.space_after = Pt(4)
    return doc

def add_iso_header_box(doc, doc_title, doc_no, rev_no="00", eff_date="19 Agustus 2026"):
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="6")
    
    col_widths = [Cm(3.8), Cm(6.2), Cm(2.5), Cm(3.0)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
    
    c_logo = table.cell(0, 0)
    c_logo.merge(table.cell(2, 0))
    p_logo = c_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("PT KCA CHEMICAL\nDIREKSI & MANAJEMEN\nSISTEM ISO 9001:2015")
    r_logo.font.name = 'Times New Roman'
    r_logo.font.size = Pt(9.5)
    r_logo.font.bold = True
    r_logo.font.color.rgb = COLOR_BLACK
    
    c_title = table.cell(0, 1)
    c_title.merge(table.cell(2, 1))
    p_title = c_title.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(doc_title.upper())
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(10.5)
    r_title.font.bold = True
    r_title.font.color.rgb = COLOR_BLACK
    
    metadata = [
        ("No. Dokumen", doc_no),
        ("No. Revisi / Tgl", f"Rev. {rev_no} / {eff_date}"),
        ("Status Dokumen", "TERKENDALI - DIREKSI")
    ]
    
    for row_idx, (label, val) in enumerate(metadata):
        c_lbl = table.cell(row_idx, 2)
        p_lbl = c_lbl.paragraphs[0]
        r_lbl = p_lbl.add_run(label)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = COLOR_BLACK
        
        c_val = table.cell(row_idx, 3)
        p_val = c_val.paragraphs[0]
        r_val = p_val.add_run(val)
        r_val.font.size = Pt(9)
        r_val.font.color.rgb = COLOR_BLACK
        
    doc.add_paragraph()

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_BLACK
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.bold = True
        r_pre.font.color.rgb = COLOR_BLACK
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.color.rgb = COLOR_BLACK
    return p

def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_BLACK
    return p

def add_styled_table(doc, headers, data, col_widths=None, alignments=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=HEX_BORDER, sz="4")
    
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_background(cell, HEX_HEADER_BG)
        set_cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_BLACK
        
    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx + 1]
        bg_color = "FFFFFF" if r_idx % 2 == 0 else HEX_ROW_ALT
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            if bg_color != "FFFFFF":
                set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            
            if alignments and c_idx < len(alignments):
                p.alignment = alignments[c_idx]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.font.color.rgb = COLOR_BLACK
            
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = w
                    
    doc.add_paragraph()
    return table

def add_signature_block_single(doc, title="Ditetapkan dan Disahkan Oleh,\nDirektur Utama PT KCA Chemical", name="( ............................................ )"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="FFFFFF", sz="0", val="none")
    cell = table.cell(0, 0)
    cell.width = Cm(15.5)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title + "\n\n\n\n\n" + name + "\nTanggal: ...................................")
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.font.color.rgb = COLOR_BLACK

# ==============================================================================
# DOC 1: KEBIJAKAN MUTU & SASARAN MUTU KPI DIREKSI
# ==============================================================================
def create_mgt_01():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Kebijakan Mutu & Sasaran Mutu Perusahaan (Quality Policy & Objectives ISO 9001)", "MGT-DIR-001", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. PERNYATAAN KEBIJAKAN MUTU PERUSAHAAN (ISO 9001 KLAUSUL 5.2)")
    add_p(doc, "PT KCA CHEMICAL berkomitmen menjadi produsen makloon Perbekalan Kesehatan Rumah Tangga (PKRT) yang unggul, terpercaya, dan patuh hukum di Indonesia dengan menerapkan 4 Pilar Keunggulan:")
    add_bullet(doc, "Memastikan 100% proses manufaktur memenuhi pedoman Cara Pembuatan PKRT yang Baik (CPKRTB) Kemenkes RI dan Sistem Manajemen Mutu ISO 9001:2015.", "1. Kepatuhan Regulasi & Mutu: ");
    add_bullet(doc, "Menjamin kerahasiaan formula, ketepatan waktu serah terima pesanan (On-Time Delivery), dan legalitas izin edar resmi KEMENKES RI PKD.", "2. Kepuasan Pelanggan & Mitra: ");
    add_bullet(doc, "Menerapkan sistem keselamatan kerja kimia (K3) yang ketat dan ramah lingkungan.", "3. Keselamatan Kerja & Lingkungan: ");
    add_bullet(doc, "Melakukan perbaikan berkelanjutan secara sistematis pada seluruh proses produksi, efisiensi biaya, dan peningkatan kompetensi sumber daya manusia.", "4. Perbaikan Berkelanjutan (Continual Improvement): ");

    add_heading_1(doc, "2. MATRIKS SASARAN MUTU & KPI PABRIK (ISO 9001 KLAUSUL 6.2)")
    headers_kpi = ["No", "Departemen", "Indikator Kinerja Utama (KPI)", "Target Tahunan", "Metode Monitoring", "PIC"]
    data_kpi = [
        ["1", "Manajemen / Direksi", "Pertumbuhan Volume Produksi Makloon", "Min. 10.000 Liter / Bln", "Laporan Penjualan Bulanan", "Direktur"],
        ["2", "PJT / Regulasi", "Keberhasilan Terbit Izin Edar PKD Baru", "100% Sesuai Timeline", "Status Aplikasi e-Farmalkes", "PJT"],
        ["3", "Produksi", "Tingkat Produk Lolos Pertama (First Pass Yield)", "Min. 98,5 %", "Rekapitulasi Batch Record", "Ka. Prod"],
        ["4", "Quality Control (QC)", "Keluhan Mutu Konsumen (Customer Complaint)", "Maksimal < 1 % per batch", "Logbook Keluhan Pelanggan", "QC / PJT"],
        ["5", "Gudang & Logistik", "Akurasi Stok Bahan Baku & Barang Jadi", "Min. 99,0 %", "Stock Opname Bulanan", "Ka. Gudang"],
        ["6", "HR & K3", "Zero Accident (Nol Kecelakaan Kerja Berat)", "0 Kasus / Tahun", "Laporan Insiden K3 Harian", "Koord. K3"]
    ]
    add_styled_table(doc, headers_kpi, data_kpi, [Cm(0.8), Cm(3.2), Cm(5.0), Cm(2.8), Cm(3.0), Cm(1.6)])

    add_signature_block_single(doc)
    doc.save(os.path.join(MGT_DIR, "MGT-01_Kebijakan_Mutu_dan_Sasaran_Mutu_KPI_Direksi.docx"))
    print("Created MGT-01")

# ==============================================================================
# DOC 2: MANAJEMEN RISIKO & ANALISIS BISNIS MAKLOON
# ==============================================================================
def create_mgt_02():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Manajemen Risiko Bisnis & Analisis Operasional Pabrik PKRT (ISO 9001 Klausul 6.1)", "MGT-DIR-002", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. TUJUAN")
    add_p(doc, "Mengidentifikasi, menganalisis, dan memitigasi seluruh risiko operasional, kepatuhan hukum, finansial, dan teknis manufaktur yang dapat menghambat pencapaian target usaha PT KCA Chemical.")

    add_heading_1(doc, "2. REGISTER IDENTIFIKASI RISIKO & RENCANA MITIGASI DIREKSI")
    headers_rsk = ["No", "Kategori Risiko", "Deskripsi Potensi Bahaya / Kejadian", "Tingkat Dampak", "Rencana Tindakan Mitigasi Direksi", "Penanggung Jawab"]
    data_rsk = [
        ["1", "Risiko Regulasi", "Izin PKD terlambat terbit karena dokumen mutu lab tidak lengkap", "Tinggi (High)", "Merekrut PJT kompeten, bekerja sama dengan lab terakreditasi KAN rujukan, dan audit kelengkapan formula sebelum submit.", "PJT & Direktur"],
        ["2", "Risiko HKI / Merek", "Klien makloon bersengketa merek dagang dengan pihak ketiga", "Tinggi (High)", "Wajib verifikasi sertifikat DJKI Kemenkumham dan meminta Surat Pernyataan Kepemilikan Merek bermaterai sebelum kontrak.", "Marketing & Legal"],
        ["3", "Risiko Mutu Produk", "Cairan sabun mengalami pemisahan fase/keruh setelah 3 bulan", "Sedang (Medium)", "Wajib uji stabilitas dipercepat (Accelerated Stability Test) dan simpan Sampel Pertinggal (Retained Sample) per batch.", "PJT & QC"],
        ["4", "Risiko K3 Kimia", "Kebocoran uap klorin atau tumpahan cairan asam di gudang", "Tinggi (High)", "Penyediaan area berventilasi khusus korosif, penampung tumpahan (spill pallet), APD respirator, dan eye washer.", "Ka. Prod & K3"],
        ["5", "Risiko Finansial", "Gagal bayar atau penundaan pelunasan oleh pihak distributor", "Sedang (Medium)", "Menerapkan sistem pembayaran DP 50% di awal dan pelunasan 50% wajib sebelum barang dimuat ke truk ekspedisi.", "Direktur / Keuangan"]
    ]
    add_styled_table(doc, headers_rsk, data_rsk, [Cm(0.8), Cm(2.4), Cm(4.2), Cm(2.2), Cm(5.0), Cm(1.8)])

    add_signature_block_single(doc)
    doc.save(os.path.join(MGT_DIR, "MGT-02_Manajemen_Risiko_dan_Mitigasi_Bisnis_Pabrik.docx"))
    print("Created MGT-02")

# ==============================================================================
# DOC 3: NOTULEN RAPAT TINJAUAN MANAJEMEN (MANAGEMENT REVIEW)
# ==============================================================================
def create_mgt_03():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Format Notulen Rapat Tinjauan Manajemen (Management Review ISO 9001 Klausul 9.3)", "MGT-DIR-003", "00", "19 Agustus 2026")
    
    add_p(doc, "NOTULEN RAPAT TINJAUAN MANAJEMEN (MANAGEMENT REVIEW MEETING)\n"
                "Hari / Tanggal  : .......................................   Waktu: 09:00 - 12:00 WIB\n"
                "Tempat          : Ruang Rapat PT KCA Chemical             Pimpinan Rapat: Direktur Utama\n"
                "Peserta Hadir   : Direktur, PJT, Kepala Produksi, QC Lab, Kepala Gudang, Marketing")
    
    add_heading_1(doc, "AGENDA PEMBAHASAN WAJIB (ISO 9001 KLAUSUL 9.3):")
    add_bullet(doc, "Tinjauan tindak lanjut hasil rapat tinjauan manajemen periode sebelumnya.", "1. ");
    add_bullet(doc, "Perubahan isu internal dan eksternal yang relevan dengan bisnis makloon PKRT.", "2. ");
    add_bullet(doc, "Kinerja sistem manajemen mutu (Pencapaian Sasaran Mutu / KPI tiap bagian).", "3. ");
    add_bullet(doc, "Hasil audit internal CPKRTB dan status tindakan korektif (CAPA).", "4. ");
    add_bullet(doc, "Umpan balik pelanggan dan keluhan konsumen selama 6 bulan terakhir.", "5. ");
    add_bullet(doc, "Kecukupan alokasi sumber daya (anggaran mesin, personil PJT, bahan baku).", "6. ");
    add_bullet(doc, "Peluang peningkatan berkelanjutan dan rencana pengembangan produk baru.", "7. ");

    add_heading_1(doc, "MATRIKS KEPUTUSAN & RENCANA TINDAK LANJUT DIREKSI:")
    headers_dec = ["No", "Topik Pembahasan & Temuan", "Keputusan / Tindakan Perbaikan", "Alokasi Anggaran / Sumber Daya", "Batas Waktu (Deadline)", "PIC"]
    data_dec = [
        ["1", "Kapasitas mixing 1000L mulai penuh", "Pengadaan 1 unit tangki mixer HDPE 2000L tambahan", "Disetujui Direktur (Rp 15 Juta)", "Bulan ke-3", "Ka. Prod"],
        ["2", "Registrasi PKD varian baru", "Pendaftaran e-Reg varian Pembersih Lantai Lavender", "Anggaran Uji Lab & PNBP Kemenkes", "Bulan ke-2", "PJT"],
        ["3", "Peningkatan sanitasi gudang", "Pemasangan 2 unit exhaust fan dinding tambahan", "Disetujui Direktur (Rp 2 Juta)", "Minggu ke-2", "Teknisi"]
    ]
    add_styled_table(doc, headers_dec, data_dec, [Cm(0.8), Cm(3.8), Cm(4.2), Cm(3.5), Cm(2.2), Cm(1.8)])

    add_signature_block_single(doc)
    doc.save(os.path.join(MGT_DIR, "MGT-03_Notulen_Rapat_Tinjauan_Manajemen_ISO9001.docx"))
    print("Created MGT-03")

# ==============================================================================
# DOC 4: PERATURAN PERUSAHAAN & TATA TERTIB PABRIK
# ==============================================================================
def create_mgt_04():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Peraturan Perusahaan, Kode Etik & Tata Tertib Karyawan Pabrik", "MGT-DIR-004", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "BAB I: KETENTUAN HARI & JAM KERJA")
    add_bullet(doc, "Hari kerja operasional pabrik adalah hari Senin sampai dengan Jumat (08:00 - 16:30 WIB) dan Sabtu (08:00 - 13:00 WIB).", "1.1 ");
    add_bullet(doc, "Karyawan wajib melakukan presensi kehadiran sebelum jam kerja dimulai dan dilarang meninggalkan area kerja tanpa izin pimpinan.", "1.2 ");

    add_heading_1(doc, "BAB II: TATA TERTIB KESELAMATAN & HIGIENE PABRIK")
    add_bullet(doc, "Seluruh karyawan wajib mematuhi standar CPKRTB: dilarang merokok, makan, minum, atau meludah di area produksi dan gudang.", "2.1 ");
    add_bullet(doc, "Wajib mengenakan Alat Pelindung Diri (APD) lengkap yang telah disediakan perusahaan selama berada di ruang kerja.", "2.2 ");
    add_bullet(doc, "Karyawan wajib menjaga kebersihan mesin, alat, meja kerja, dan merapikan peralatan setelah selesai digunakan.", "2.3 ");

    add_heading_1(doc, "BAB III: KERAHASIAAN FORMULA & INFORMASI PERUSAHAAN")
    add_bullet(doc, "Seluruh data formula kimia, takaran bahan, daftar supplier, dan data klien makloon adalah RAHASIA DAGANG PERUSAHAAN yang dilindungi hukum.", "3.1 ");
    add_bullet(doc, "Karyawan dilarang keras menyalin, memfoto, membawa keluar, atau membocorkan data formula kepada pihak luar manapun.", "3.2 ");
    add_bullet(doc, "Pelanggaran terhadap kerahasiaan formula akan dikenakan Pemutusan Hubungan Kerja (PHK) seketika dan tuntutan pidana sesuai UU Rahasia Dagang No. 30 Tahun 2000.", "3.3 ");

    add_heading_1(doc, "BAB IV: TINGKAT SANKSI & SURAT PERINGATAN (SP)")
    add_p(doc, "Pelanggaran terhadap tata tertib akan dikenakan sanksi bertingkat:\n"
               "1. Teguran Lisan / Tertulis (Pelanggaran ringan seperti terlambat atau tidak rapi).\n"
               "2. Surat Peringatan Pertama (SP 1) - Masa berlaku 6 bulan.\n"
               "3. Surat Peringatan Kedua (SP 2) - Pengulangan pelanggaran.\n"
               "4. Surat Peringatan Ketiga (SP 3) / PHK (Pelanggaran berat seperti pencurian bahan, perkelahian, merokok di area kimia, atau pembocoran formula).")

    add_signature_block_single(doc)
    doc.save(os.path.join(MGT_DIR, "MGT-04_Peraturan_Perusahaan_dan_Tata_Tertib_Karyawan.docx"))
    print("Created MGT-04")

# ==============================================================================
# DOC 5: SURAT PERJANJIAN PENGUASAAN LAHAN / BANGUNAN PABRIK
# ==============================================================================
def create_mgt_05():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Surat Perjanjian Sewa-Menyewa / Penguasaan Bangunan Sarana Pabrik", "MGT-DIR-005", "00", "19 Agustus 2026")
    
    add_p(doc, "SURAT PERJANJIAN SEWA-MENYEWA BANGUNAN SARANA INDUSTRI PABRIK\nNomor: 008/SWA-PBR/KCA/VIII/2026\n\n"
                "Pada hari ini, [Hari], tanggal [Tanggal] bulan [Bulan] tahun [Tahun], bertempat di [Kota], telah disepakati perjanjian sewa-menyewa oleh dan antara:\n"
                "1. [Nama Pemilik Lahan/Bangunan], bertindak selaku PEMILIK BANGUNAN (selanjutnya disebut PIHAK PERTAMA).\n"
                "2. [Nama Direktur Anda], selaku Direktur Utama bertindak untuk dan atas nama PT KCA CHEMICAL (selanjutnya disebut PIHAK KEDUA).\n\n"
                "PARA PIHAK sepakat untuk mengikatkan diri dalam perjanjian sewa-menyewa dengan ketentuan sebagai berikut:\n\n"
                "PASAL 1: OBJEK SEWA & PERUNTUKAN\n"
                "PIHAK PERTAMA menyewakan kepada PIHAK KEDUA sebidang tanah dan bangunan seluas [.....] m² yang terletak di [Alamat Lengkap Lokasi Pabrik] untuk dipergunakan secara sah sebagai Sarana Industri Manufaktur Perbekalan Kesehatan Rumah Tangga (PKRT) dan Gudang.\n\n"
                "PASAL 2: JANGKA WAKTU SEWA\n"
                "Sewa-menyewa ini dilangsungkan untuk jangka waktu [3 (tiga) / 5 (lima)] tahun terhitung sejak tanggal [Tgl Mulai] sampai dengan [Tgl Berakhir] guna memenuhi persyaratan perizinan berusaha sarana industri Kementerian Kesehatan RI.\n\n"
                "PASAL 3: BIAYA SEWA DAN PEMBAYARAN\n"
                "Harga sewa yang disepakati adalah sebesar Rp [Nominal] per tahun atau total sebesar Rp [Total] untuk seluruh masa sewa yang dibayarkan lunas dengan bukti kuitansi sah.")

    add_signature_block_single(doc, "Disetujui Oleh Kedua Belah Pihak,\n\nPIHAK PERTAMA (Pemilik),                         PIHAK KEDUA (Direktur PT KCA),\n\n\n\n( Materai Rp 10.000 )                            ( Materai Rp 10.000 )\n[ Nama Pemilik Bangunan ]                        [ Nama Direktur Anda ]")
    doc.save(os.path.join(MGT_DIR, "MGT-05_Surat_Perjanjian_Sewa_Bangunan_Pabrik.docx"))
    print("Created MGT-05")

# ==============================================================================
# DOC 6: EXECUTIVE FINANCIAL DASHBOARD & PERENCANAAN BIAYA
# ==============================================================================
def create_mgt_06():
    doc = setup_iso_document()
    add_iso_header_box(doc, "Format Laporan Eksekutif Keuangan, HPP & Utilisasi Kapasitas Pabrik", "MGT-DIR-006", "00", "19 Agustus 2026")
    
    add_heading_1(doc, "1. STRUKTUR ESTIMASI HARGA POKOK PRODUKSI (HPP) JASA MAKLOON")
    add_p(doc, "Berikut adalah kalkulasi struktur biaya per kemasan untuk penentuan margin keuntungan makloon Direksi:")
    
    headers_hpp = ["Komponen Biaya", "Sabun Cuci Piring 450mL", "Pembersih Lantai 5 Liter", "Pemutih Pakaian 5 Liter", "Keterangan"]
    data_hpp = [
        ["Bahan Baku Kimia (Surfaktan/Klorin)", "Rp 1.850", "Rp 12.500", "Rp 11.200", "Formula Kemenkes"],
        ["Kemasan Botol/Jerigen + Tutup", "Rp 1.100", "Rp 8.500", "Rp 8.500", "HDPE Tebal Anti-Bocor"],
        ["Stiker Label & Karton Boks", "Rp 350", "Rp 1.800", "Rp 1.500", "Label Cromo / Vynil"],
        ["Biaya Operasional / Listrik / Air", "Rp 200", "Rp 1.200", "Rp 1.100", "Overhead Pabrik"],
        ["Upah Tenaga Kerja & PJT per Unit", "Rp 300", "Rp 1.500", "Rp 1.500", "Direct Labor & QA"],
        ["TOTAL HPP RIIL", "Rp 3.800", "Rp 25.500", "Rp 23.800", "Biaya Modal Dasar"],
        ["HARGA JUAL MAKLOON (FULL SERVICE)", "Rp 4.500 - Rp 5.000", "Rp 30.000 - Rp 34.000", "Rp 28.000 - Rp 32.000", "Penawaran ke Klien"],
        ["MARGIN LABA KOTOR DIREKSI", "Rp 700 - Rp 1.200 (20-30%)", "Rp 4.500 - Rp 8.500 (25-33%)", "Rp 4.200 - Rp 8.200 (25-35%)", "Profit per Kemasan"]
    ]
    add_styled_table(doc, headers_hpp, data_hpp, [Cm(4.5), Cm(3.2), Cm(3.2), Cm(3.2), Cm(2.4)])

    add_heading_1(doc, "2. FORMULIR REKAPITULASI OMSET & UTILISASI BULANAN DIREKSI")
    headers_mon = ["Bulan", "Klien Makloon (PT Mitra)", "Produk yang Diproduksi", "Total Volume (Liter)", "Total Omset Makloon", "Status Pelunasan"]
    data_mon = [
        ["Bulan 1", "PT Mitra Distributor A", "Sabun Cuci Piring 450 mL", "5.000 Liter (11.000 Btl)", "Rp 49.500.000", "Lunas 100%"],
        ["Bulan 1", "PT Berkah Abadi", "Pemutih Pakaian 5 Liter", "2.500 Liter (500 Jerigen)", "Rp 15.000.000", "Lunas 100%"],
        ["Bulan 2", "................................", "................................", ".................... Liter", "Rp ....................", "[  ] DP  [  ] Lunas"]
    ]
    add_styled_table(doc, headers_mon, data_mon, [Cm(2.0), Cm(3.8), Cm(4.2), Cm(3.2), Cm(2.8), Cm(2.0)])

    add_signature_block_single(doc)
    doc.save(os.path.join(MGT_DIR, "MGT-06_Format_Laporan_Keuangan_dan_Kalkulasi_HPP_Direksi.docx"))
    print("Created MGT-06")

if __name__ == "__main__":
    create_mgt_01()
    create_mgt_02()
    create_mgt_03()
    create_mgt_04()
    create_mgt_05()
    create_mgt_06()
    print("\nALL 6 EXECUTIVE MANAGEMENT DOCUMENTS SUCCESSFULLY CREATED!")
